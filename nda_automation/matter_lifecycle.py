"""Repository matter lifecycle operations.

This module is the deeper matter lifecycle interface above the persistence
adapters. Callers use it to express domain transitions, while the existing
``MatterRepository`` adapters keep owning storage details.
"""
from __future__ import annotations

import logging
import threading
import json
from collections.abc import Callable
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any

from . import artifact_service, workflow
from .artifact_registry import ArtifactRegistryError
from .matter_repository import MatterRepository

logger = logging.getLogger(__name__)

BackgroundRunner = Callable[[Callable[[], None]], None]
MAX_REDLINE_DRAFT_ITEMS = 200

# Recipient-facing disclosure for PDF-source redline sends. The attached Word
# document is a pdf2docx RECONSTRUCTION of the counterparty's executed PDF, so its
# layout/tables/formatting may differ from the original. The sender already sees a
# caveat in the app; this is the SAME honest signal carried to the counterparty who
# opens the file. Kept verbatim so tests/UI can assert on it.
PDF_RECONSTRUCTION_RECIPIENT_CAVEAT = (
    "Note: the attached Word document was reconstructed from a PDF and may differ "
    "in formatting (layout, tables, spacing) from the original. The tracked changes "
    "reflect the proposed edits; please review the wording rather than the layout."
)


def _prepend_reconstruction_caveat(body: str | None) -> str:
    """Prepend the PDF-reconstruction caveat to an outgoing email body.

    Additive and minimal: it only prefixes a clearly-marked disclosure line and
    leaves the caller-supplied body intact below it. Never touches recipient,
    subject, or any addressing/auth.
    """
    caveat = PDF_RECONSTRUCTION_RECIPIENT_CAVEAT
    existing = (body or "").strip("\n")
    if not existing:
        return caveat
    # Idempotent: don't double-stamp if the caveat is already present.
    if PDF_RECONSTRUCTION_RECIPIENT_CAVEAT in existing:
        return body or ""
    return f"{caveat}\n\n{existing}"


# Short cover-note stamped into the reconstructed DOCX itself so the disclosure
# survives even if the email body is stripped/forwarded. Kept short and verbatim.
PDF_RECONSTRUCTION_DOCX_COVER_NOTE = (
    "RECONSTRUCTED DOCUMENT - This Word file was rebuilt from a PDF and may differ "
    "in formatting from the original; review the wording, not the layout."
)


def _stamp_docx_reconstruction_notice(docx_bytes: bytes) -> bytes:
    """Best-effort: prepend a one-line reconstruction cover note to a DOCX body.

    Inserts a single paragraph as the FIRST child of ``<w:body>`` and never touches
    existing runs (tracked-change insertions/deletions are preserved), so the
    upstream redline-coverage gate stays satisfied. FAILS OPEN: any error returns
    the original bytes unchanged so a disclosure failure can never block a send.
    Only ever called for PDF-reconstruction sends.
    """

    try:
        import io

        from docx import Document  # noqa: PLC0415
        from docx.oxml.ns import qn  # noqa: PLC0415
        from docx.shared import Pt, RGBColor  # noqa: PLC0415

        document = Document(io.BytesIO(docx_bytes))
        body = document.element.body
        # A fresh empty <w:p> hosting the note; placed at the top of the body so it
        # reads as a cover line above the reconstructed content.
        paragraph = document.add_paragraph()
        run = paragraph.add_run(PDF_RECONSTRUCTION_DOCX_COVER_NOTE)
        run.bold = True
        try:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
        except Exception:
            # Styling is cosmetic; never let it abort the stamp.
            pass
        note_p = paragraph._p
        note_p.getparent().remove(note_p)
        # Insert before the first existing block (sectPr stays last automatically).
        first_block = None
        for child in body:
            if child.tag in (qn("w:p"), qn("w:tbl")):
                first_block = child
                break
        if first_block is not None:
            first_block.addprevious(note_p)
        else:
            sect_pr = body.find(qn("w:sectPr"))
            if sect_pr is not None:
                sect_pr.addprevious(note_p)
            else:
                body.append(note_p)

        out = io.BytesIO()
        document.save(out)
        return out.getvalue()
    except Exception:
        # Fail open: the email-body caveat is the guaranteed disclosure; a DOCX
        # stamping failure must never block or corrupt an otherwise-valid send.
        return docx_bytes


class MatterLifecycleError(RuntimeError):
    pass


class MatterNotFoundError(MatterLifecycleError):
    pass


class MatterApprovalBlockedError(MatterLifecycleError):
    def __init__(self, blocks: list[dict[str, Any]], resolution: dict[str, Any]) -> None:
        super().__init__("NDA cannot be approved yet.")
        self.blocks = blocks
        self.resolution = resolution


class RedlineDraftError(MatterLifecycleError):
    pass


class MatterReviewUnavailableError(MatterLifecycleError):
    pass


class MatterSendBlockedError(MatterLifecycleError):
    pass


class MatterDeliveryError(MatterLifecycleError):
    pass


def _matter_review_block_resolved(matter: dict[str, Any]) -> bool:
    """Has a human resolved the review/fail block so the redline can go out?

    The send gate now blocks BOTH needs-review and unresolved-fail (check) state
    (see review_state.result_requires_human_review). Both are cleared the same
    way: a human engages the matter. ``human_reviewed`` is the canonical
    reviewed flag (set via the board's "mark reviewed" toggle); a recorded
    approval (``status == "approved"`` / ``approved_at``) is a strictly stronger
    sign-off and also clears it, so an approved fail-state matter is never
    permanently wedged at the send step.
    """
    if matter.get("human_reviewed"):
        return True
    if str(matter.get("status") or "").strip().lower() == "approved":
        return True
    return bool(matter.get("approved_at"))


@dataclass(frozen=True)
class ReviewRefreshResult:
    matter: dict[str, Any]
    was_stale: bool
    had_redline_draft: bool
    refresh_attempted: bool = True


@dataclass(frozen=True)
class MatterApprovalResult:
    matter: dict[str, Any]
    approved_at: str
    approver: str
    timeline_event: dict[str, Any]
    resolution: dict[str, Any]


@dataclass(frozen=True)
class MatterRedlineSendResult:
    matter: dict[str, Any]
    filename: str
    sent: dict[str, Any]
    # True when the sent Word file was reconstructed from a PDF source (best-effort,
    # not faithful original formatting). The route surfaces this caveat to the operator.
    reconstructed_from_pdf: bool = False


@dataclass(frozen=True)
class MatterDocumentSendResult:
    matter: dict[str, Any]
    filename: str
    sent: dict[str, Any]


def run_in_daemon_thread(work: Callable[[], None]) -> None:
    """Run lifecycle follow-up work off the request path."""
    thread = threading.Thread(target=work, name="drive-auto-intake", daemon=True)
    thread.start()


class RepositoryMatterLifecycle:
    """Lifecycle transitions for Repository matters.

    The persistence seam stays intentionally small in the caller: route and
    intake modules ask for lifecycle transitions, not individual storage writes.
    """

    def __init__(self, repository: MatterRepository) -> None:
        self._repository = repository

    def complete_intake(
        self,
        matter: dict[str, Any],
        *,
        owner_user_id: str = "",
        drive_sync_runner: BackgroundRunner = run_in_daemon_thread,
    ) -> None:
        """Run the fail-soft hooks that complete a newly created matter intake."""
        self._register_original_artifact(matter, owner_user_id=owner_user_id)
        self._record_intake_timeline(matter, owner_user_id=owner_user_id)
        self._auto_sync_drive(matter, owner_user_id=owner_user_id, runner=drive_sync_runner)

    def refresh_review(
        self,
        matter: dict[str, Any],
        *,
        review_engine_func: Callable[..., dict[str, Any]] | None = None,
        review_staleness_func: Callable[[object], bool] | None = None,
    ) -> ReviewRefreshResult:
        from .checker import EvidenceProvenanceError, ParagraphAlignmentError, PlaybookTemplateError
        from .review_engine import ActiveReviewEngineError, review_nda_with_active_engine
        from .review_staleness import review_result_is_stale
        from .triage import triage_review_result

        review_engine = review_engine_func or review_nda_with_active_engine
        staleness_check = review_staleness_func or review_result_is_stale
        was_stale = staleness_check(matter.get("review_result"))
        had_redline_draft = isinstance(matter.get("redline_draft"), dict)
        # Snapshot the matter's updated_at BEFORE the (multi-second) AI runs. The
        # guarded store writer below compares against it: if any write (a human
        # mark-reviewed or a saved redline draft) lands during the AI window, the
        # updated_at will have moved and those human edits are preserved instead of
        # being silently reset/popped.
        expected_updated_at = str(matter.get("updated_at") or "")
        if not was_stale:
            return ReviewRefreshResult(matter=matter, was_stale=False, had_redline_draft=had_redline_draft)

        extracted_text = str(matter.get("extracted_text") or "")
        if not extracted_text.strip():
            return ReviewRefreshResult(matter=matter, was_stale=True, had_redline_draft=had_redline_draft)

        paragraphs = self._original_docx_paragraphs(matter)
        try:
            review_result = review_engine(extracted_text, paragraphs=paragraphs)
        except ParagraphAlignmentError:
            if paragraphs is None:
                return ReviewRefreshResult(matter=matter, was_stale=True, had_redline_draft=had_redline_draft)
            try:
                review_result = review_engine(extracted_text)
            except (ActiveReviewEngineError, EvidenceProvenanceError, ParagraphAlignmentError, PlaybookTemplateError, ValueError):
                return ReviewRefreshResult(matter=matter, was_stale=True, had_redline_draft=had_redline_draft)
        except (ActiveReviewEngineError, EvidenceProvenanceError, PlaybookTemplateError, ValueError):
            return ReviewRefreshResult(matter=matter, was_stale=True, had_redline_draft=had_redline_draft)

        triage = triage_review_result(review_result)
        matter_id = str(matter.get("id") or "")
        owner_user_id = str(matter.get("owner_user_id") or "")
        # PERSIST-POINT RIGHT OF WAY for the user-initiated refresh too. We do NOT
        # hard-error a deliberate Refresh click (that reads as broken); the refresh
        # still runs and returns a fresh review. But its store WRITE stands back so
        # a concurrent foreground generate's save wins the single global store lock.
        # Bounded + fail-open: the write always lands, just a beat later if a
        # generate is mid-save. (The verifier itself also defers while a generate is
        # active, so the refresh's AI burst is light during that window.)
        from . import generation_priority  # noqa: PLC0415 - keep the dep light/local.

        generation_priority.yield_store_to_generation()
        updated_matter = self._repository.refresh_matter_review(
            matter_id,
            review_result,
            triage,
            expected_updated_at=expected_updated_at,
            owner_user_id=owner_user_id,
        )
        if updated_matter is None:
            updated_matter = {
                **matter,
                "review_result": review_result,
                **triage,
                "human_reviewed": False,
            }
            updated_matter.pop("redline_draft", None)
        return ReviewRefreshResult(
            matter=updated_matter,
            was_stale=True,
            had_redline_draft=had_redline_draft,
        )

    def save_redline_draft(
        self,
        matter_id: str,
        raw_redline_draft: object,
        *,
        owner_user_id: str = "",
    ) -> dict[str, Any]:
        if raw_redline_draft is None:
            draft = None
        elif isinstance(raw_redline_draft, dict):
            draft = clean_redline_draft(raw_redline_draft)
        else:
            raise RedlineDraftError("Redline draft must be an object or null.")
        matter = self._repository.update_redline_draft(matter_id, draft, owner_user_id=owner_user_id)
        if matter is None:
            raise MatterNotFoundError("NDA not found.")
        return matter

    def approve_matter(
        self,
        matter_id: str,
        *,
        actor: str,
        owner_user_id: str = "",
    ) -> MatterApprovalResult:
        from . import approval

        matter = self._repository.get_matter(matter_id, owner_user_id=owner_user_id)
        if matter is None:
            raise MatterNotFoundError("NDA not found.")
        blocks = approval.approval_blocks(matter)
        if blocks:
            raise MatterApprovalBlockedError(blocks, approval.resolution_summary(matter))

        approved_at = datetime.now(timezone.utc).isoformat()
        timeline_event = approval.approval_timeline_event(actor=actor)
        updated_matter = self._repository.record_matter_approval(
            matter_id,
            approver=actor,
            approved_at=approved_at,
            timeline_event=timeline_event,
            owner_user_id=owner_user_id,
        )
        if updated_matter is None:
            raise MatterNotFoundError("NDA not found.")
        return MatterApprovalResult(
            matter=updated_matter,
            approved_at=approved_at,
            approver=actor,
            timeline_event=timeline_event,
            resolution=approval.resolution_summary(updated_matter),
        )

    def send_redline(
        self,
        matter_id: str,
        payload: dict[str, Any],
        *,
        owner_user_id: str = "",
        token_owner_user_id: str = "",
        to: str | None = None,
        confirmed_recipient: str | None = None,
        subject: str | None = None,
        body: str | None = None,
    ) -> MatterRedlineSendResult:
        from . import app_settings, gmail_integration, matter_view, redline_export_service, source_document_policy

        matter = self._repository.get_matter(matter_id, owner_user_id=owner_user_id)
        if matter is None:
            raise MatterNotFoundError("NDA not found.")
        if not to and not gmail_integration.matter_reply_recipient(matter):
            raise MatterDeliveryError("NDA does not have a valid reply recipient email address.")
        if not confirmed_recipient:
            raise MatterDeliveryError("Confirm the outbound recipient email address before sending.")
        if matter_view.matter_needs_human_review(matter) and not _matter_review_block_resolved(matter):
            raise MatterSendBlockedError("NDA needs human review before a redline can be sent.")
        if not app_settings.gmail_role_enabled("outbound"):
            raise MatterSendBlockedError("Gmail outbound is disabled in Admin.")

        validate_kwargs = {
            "to": to,
            "confirmed_recipient": confirmed_recipient,
        }
        if token_owner_user_id:
            validate_kwargs["owner_user_id"] = token_owner_user_id
        gmail_integration.validate_outbound_send_ready(matter, **validate_kwargs)

        redline_export = redline_export_service.build_matter_redline(
            matter_id,
            payload,
            repository=self._repository,
            owner_user_id=owner_user_id,
        )
        send_matter = self._repository.get_matter(matter_id, owner_user_id=owner_user_id)
        if send_matter is None:
            raise MatterNotFoundError("NDA not found.")
        if matter_view.matter_needs_human_review(send_matter) and not _matter_review_block_resolved(send_matter):
            raise MatterSendBlockedError("NDA needs human review before a redline can be sent.")

        # For PDF-source matters the redline is reconstructed from the PDF (the
        # export stamps the X-PDF-DOCX-Reconstruction header), so the sent Word file
        # is best-effort, not faithful original formatting. Determine this BEFORE the
        # send so the recipient-facing caveat can be injected into the outgoing body.
        # Prefer the per-export marker; fall back to the matter-source predicate.
        reconstructed_from_pdf = bool(
            (redline_export.headers and redline_export.headers.get("X-PDF-DOCX-Reconstruction"))
            or source_document_policy.matter_source_is_pdf(send_matter)
        )

        # DOCX-source sends are byte- and body-unchanged; only PDF-reconstruction
        # sends get the recipient disclosure (email body + best-effort DOCX cover note).
        outbound_body = body
        outbound_attachment = redline_export.data
        if reconstructed_from_pdf:
            outbound_body = _prepend_reconstruction_caveat(body)
            outbound_attachment = _stamp_docx_reconstruction_notice(redline_export.data)

        send_kwargs = {
            "body": outbound_body,
            "confirmed_recipient": confirmed_recipient,
            "subject": subject,
            "to": to,
        }
        if token_owner_user_id:
            send_kwargs["owner_user_id"] = token_owner_user_id
        sent = gmail_integration.send_redline_email(
            send_matter,
            outbound_attachment,
            redline_export.filename,
            **send_kwargs,
        )
        updated_matter = self.record_sent_delivery(
            matter_id,
            sent,
            filename=redline_export.filename,
            owner_user_id=owner_user_id,
        )
        # Capture the emailed document as a SENT lifecycle artifact (the exact
        # bytes that went out, plus the resolved recipient). Best-effort: the
        # hook stub is a no-op today, and a hook failure must never undo a send
        # that already succeeded.
        self._capture_sent_artifact(
            matter_id,
            sent_bytes=outbound_attachment,
            filename=redline_export.filename,
            recipient=str(sent.get("to") or confirmed_recipient or ""),
            owner_user_id=owner_user_id,
        )
        return MatterRedlineSendResult(
            matter=updated_matter,
            filename=redline_export.filename,
            sent=sent,
            reconstructed_from_pdf=reconstructed_from_pdf,
        )

    def send_document(
        self,
        *,
        filename: str,
        document_bytes: bytes,
        recipient: str,
        subject: str,
        body: str | None = None,
        owner_user_id: str = "",
        token_owner_user_id: str = "",
    ) -> MatterDocumentSendResult:
        from . import gmail_integration

        transient_matter = {"subject": subject, "reply_to": recipient}
        send_kwargs = {"body": body, "subject": subject, "to": recipient}
        if token_owner_user_id:
            send_kwargs["owner_user_id"] = token_owner_user_id
        sent = gmail_integration.send_redline_email(
            transient_matter,
            document_bytes,
            filename,
            **send_kwargs,
        )
        matter = self._repository.create_matter(
            source_filename=filename,
            document_bytes=document_bytes,
            extracted_text="",
            review_result={},
            triage=send_document_triage(),
            source_type="send_document",
            board_column="sent",
            intake_metadata=send_document_metadata(filename, recipient, subject),
            owner_user_id=owner_user_id,
        )
        matter_id = str(matter.get("id") or "")
        updated_matter = self.record_sent_delivery(
            matter_id,
            sent,
            filename=filename,
            owner_user_id=owner_user_id,
        )
        # Capture the emailed document as a SENT lifecycle artifact (the exact
        # bytes that went out, plus the resolved recipient), like the redline
        # path. Best-effort + guarded: the send already succeeded, so artifact
        # capture must never undo it.
        self._capture_sent_artifact(
            matter_id,
            sent_bytes=document_bytes,
            filename=filename,
            recipient=str(sent.get("to") or recipient or ""),
            owner_user_id=owner_user_id,
        )
        return MatterDocumentSendResult(matter=updated_matter, filename=filename, sent=sent)

    def record_sent_delivery(
        self,
        matter_id: str,
        sent: dict[str, Any],
        *,
        filename: str,
        owner_user_id: str = "",
    ) -> dict[str, Any]:
        updated_matter = self._repository.update_matter_fields(
            matter_id,
            {
                "board_column": "sent",
                "last_outbound_account": sent.get("outbound_account", ""),
                "last_outbound_at": sent.get("sent_at", ""),
                "last_outbound_filename": filename,
                "last_outbound_message_id": sent.get("message_id", ""),
                "last_outbound_subject": sent.get("subject", ""),
                "last_outbound_thread_id": sent.get("thread_id", ""),
                "last_outbound_to": sent.get("to", ""),
                "status": "active",
            },
            owner_user_id=owner_user_id,
        )
        if updated_matter is None:
            raise MatterNotFoundError("NDA not found.")
        self._stamp_sent_timeline(updated_matter, sent, owner_user_id=owner_user_id)
        return updated_matter

    def _capture_sent_artifact(
        self,
        matter_id: str,
        *,
        sent_bytes: bytes,
        filename: str,
        recipient: str,
        owner_user_id: str = "",
    ) -> None:
        """Register the emailed document as a SENT lifecycle artifact (best-effort).

        Delegates to the ``lifecycle_sent`` hook module (a safe no-op stub until
        the hook agent implements it). Guarded so a hook failure never undoes a
        delivery that already succeeded.
        """
        try:
            from . import lifecycle_sent

            lifecycle_sent.capture_sent_artifact(
                self._repository,
                matter_id,
                owner_user_id,
                sent_bytes,
                filename,
                recipient,
            )
        except Exception:
            # The send already succeeded; SENT-artifact capture is additive.
            pass

    def _register_original_artifact(self, matter: dict[str, Any], *, owner_user_id: str = "") -> None:
        if matter.get("_existing_gmail_duplicate"):
            return
        try:
            artifact_service.backfill_matter(
                matter,
                repository=self._repository,
                owner_user_id=owner_user_id,
            )
        except (ArtifactRegistryError, OSError, KeyError, ValueError):
            # The matter is already persisted; backfill can recover later.
            pass

    def _record_intake_timeline(self, matter: dict[str, Any], *, owner_user_id: str = "") -> None:
        if not isinstance(matter, dict) or matter.get("_existing_gmail_duplicate"):
            return
        matter_id = str(matter.get("id") or "")
        if not matter_id:
            return
        state = workflow.workflow_state(matter)
        try:
            self._repository.append_timeline_event(
                matter_id,
                workflow.build_timeline_event(
                    workflow.EVENT_CREATED,
                    phase=workflow.PHASE_INTAKE,
                    status=workflow.STATUS_EXTRACTED,
                    actor="system",
                    detail=str(matter.get("source_type") or ""),
                ),
                owner_user_id=owner_user_id,
            )
            event_type = (
                workflow.EVENT_FLAGGED_FOR_HUMAN
                if state["status"] == workflow.STATUS_AWAITING_HUMAN
                else workflow.EVENT_REVIEW_COMPLETED
            )
            self._repository.append_timeline_event(
                matter_id,
                workflow.build_timeline_event(
                    event_type,
                    phase=state["phase"],
                    status=state["status"],
                    actor="system",
                ),
                owner_user_id=owner_user_id,
            )
        except Exception:
            # Timeline stamping must not break a successful intake.
            return

    def _auto_sync_drive(
        self,
        matter: dict[str, Any],
        *,
        owner_user_id: str = "",
        runner: BackgroundRunner = run_in_daemon_thread,
    ) -> None:
        if not isinstance(matter, dict) or matter.get("_existing_gmail_duplicate"):
            return
        matter_id = str(matter.get("id") or "")
        if not matter_id:
            return

        from . import app_settings, drive_integration, telemetry

        try:
            connected = drive_integration.drive_connected(owner_user_id)
            auto_intake = app_settings.drive_auto_intake_enabled()
            # Master pause gate: an explicitly-paused Drive stops all activity.
            active = app_settings.drive_active()
        except Exception:
            telemetry.increment("drive_auto_intake_skipped")
            return
        if not connected or not active or not auto_intake:
            telemetry.increment("drive_auto_intake_skipped")
            return

        try:
            root_folder_id = str(app_settings.drive_settings().get("folder_id") or "")
        except Exception:
            root_folder_id = ""

        def _work() -> None:
            self._perform_drive_sync(
                matter_id,
                owner_user_id=owner_user_id,
                root_folder_id=root_folder_id,
            )

        try:
            runner(_work)
        except Exception:
            telemetry.increment("drive_auto_intake_failed")

    def _perform_drive_sync(
        self,
        matter_id: str,
        *,
        owner_user_id: str = "",
        root_folder_id: str = "",
    ) -> None:
        from . import drive_integration, telemetry

        try:
            matter = self._repository.get_matter(matter_id, owner_user_id=owner_user_id)
            if not isinstance(matter, dict):
                telemetry.increment("drive_auto_intake_failed")
                return
            synced_at = datetime.now(timezone.utc).isoformat()
            synced = drive_integration.sync_matter_folder(
                matter=matter,
                matter_id=matter_id,
                owner_user_id=owner_user_id,
                root_folder_id=root_folder_id,
                synced_at=synced_at,
            )
            self._repository.update_matter_fields(
                matter_id,
                {
                    "drive": {
                        "matter_folder_id": synced["matter_folder_id"],
                        "matter_folder_url": synced["matter_folder_url"],
                        "synced_at": synced_at,
                        "artifacts": synced["artifacts"],
                    }
                },
                owner_user_id=owner_user_id,
            )
            telemetry.increment("drive_auto_intake_synced")
            telemetry.increment("drive_files_synced", amount=int(synced.get("synced_count") or 0))
        except Exception as error:
            telemetry.increment("drive_auto_intake_failed")
            # Daemon-thread failure was previously counter-only (invisible in the
            # process log). Exception CLASS only -- no filenames/folder names/content.
            logger.warning(
                "Drive auto-intake sync failed for matter %s: %s",
                matter_id,
                error.__class__.__name__,
            )

    def _stamp_sent_timeline(self, matter: dict[str, Any], sent: dict[str, Any], *, owner_user_id: str = "") -> None:
        matter_id = str(matter.get("id") or "")
        if not matter_id:
            return
        try:
            self._repository.append_timeline_event(
                matter_id,
                workflow.build_timeline_event(
                    workflow.EVENT_SENT,
                    phase=workflow.PHASE_SENT,
                    status=workflow.STATUS_SENT_AWAITING_COUNTERPARTY,
                    actor=str(sent.get("outbound_account") or "system"),
                    detail=str(sent.get("to") or ""),
                ),
                owner_user_id=owner_user_id,
            )
        except Exception:
            return

    def _original_docx_paragraphs(self, matter: dict[str, Any]) -> list[dict[str, Any]] | None:
        source_filename = str(matter.get("source_filename") or matter.get("stored_filename") or "")
        if not source_filename.lower().endswith(".docx"):
            return None
        source_bytes = self._repository.get_source_document_bytes(matter)
        if not source_bytes:
            return None
        try:
            from .checker import ParagraphAlignmentError
            from .docx_text import DocxExtractionError, extract_docx_paragraphs
            from .review_document import align_document_paragraphs

            rich = extract_docx_paragraphs(source_bytes)
            source_text = "\n\n".join(str(paragraph.get("text", "")) for paragraph in rich)
            return align_document_paragraphs(rich, source_text)
        except (DocxExtractionError, ParagraphAlignmentError, ValueError, OSError):
            return None


def clean_redline_draft(draft: dict[str, Any]) -> dict[str, Any]:
    from . import export_service

    manual_redlines = [
        cleaned
        for cleaned in (
            export_service.clean_manual_export_redline(redline)
            for redline in clean_dict_list(draft.get("manual_redline_edits"))
        )
        if cleaned is not None
    ]
    cleaned = {
        "clause_decisions": clean_bool_map(draft.get("clause_decisions")),
        "redline_decisions": clean_bool_map(draft.get("redline_decisions")),
        "template_selections": clean_text_map(draft.get("template_selections")),
        "reviewed_clause_ids": clean_bool_map(draft.get("reviewed_clause_ids")),
        "export_redline_edits": clean_dict_list(draft.get("export_redline_edits")),
        "manual_redline_edits": manual_redlines,
        "review_comments": export_service.clean_review_comments(draft.get("review_comments")),
        "saved_at": datetime.now(timezone.utc).isoformat(),
    }
    cleaned["summary"] = {
        "included_redline_count": len(cleaned["export_redline_edits"]),
        "manual_redline_count": len(cleaned["manual_redline_edits"]),
        "review_comment_count": len(cleaned["review_comments"]),
    }
    return cleaned


def clean_bool_map(value: object) -> dict[str, bool]:
    if not isinstance(value, dict):
        return {}
    cleaned = {}
    for key, item in list(value.items())[:MAX_REDLINE_DRAFT_ITEMS]:
        key = str(key).strip()[:120]
        if key:
            cleaned[key] = bool(item)
    return cleaned


def clean_text_map(value: object) -> dict[str, str]:
    if not isinstance(value, dict):
        return {}
    cleaned = {}
    for key, item in list(value.items())[:MAX_REDLINE_DRAFT_ITEMS]:
        key = str(key).strip()[:120]
        item = str(item).strip()[:240]
        if key and item:
            cleaned[key] = item
    return cleaned


def clean_dict_list(value: object) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    cleaned = []
    for item in value[:MAX_REDLINE_DRAFT_ITEMS]:
        if not isinstance(item, dict):
            continue
        cleaned.append(json.loads(json.dumps(item)))
    return cleaned


def ai_first_review_store_metadata(
    ai_first_review_result: dict[str, Any],
    *,
    started_at: str,
    completed_at: str,
) -> dict[str, object]:
    result_metadata = ai_first_review_result.get("ai_first_review")
    if not isinstance(result_metadata, dict):
        result_metadata = {}
    return {
        "status": str(result_metadata.get("status") or "completed"),
        "mode": str(result_metadata.get("mode") or "ai_first_assessor"),
        "provider": str(result_metadata.get("provider") or ""),
        "model": str(result_metadata.get("model") or ""),
        "review_mode": str(ai_first_review_result.get("review_mode") or ""),
        "review_engine_version": ai_first_review_result.get("review_engine_version"),
        "started_at": started_at,
        "completed_at": completed_at,
        "requirements_passed": int(ai_first_review_result.get("requirements_passed") or 0),
        "requirements_needs_review": int(ai_first_review_result.get("requirements_needs_review") or 0),
        "requirements_failed": int(ai_first_review_result.get("requirements_failed") or 0),
    }


def send_document_metadata(filename: str, recipient: str, subject: str) -> dict[str, str]:
    from pathlib import Path

    return {
        "sender": recipient,
        "reply_to": recipient,
        "subject": subject,
        "received_at": datetime.now(timezone.utc).isoformat(),
        "message_snippet": f"Sent {Path(str(filename or '')).name or 'document'} to {recipient}.",
        "attachment_filename": filename,
    }


def send_document_triage() -> dict[str, Any]:
    return {
        "triage_status": "sent",
        "next_action": "Document sent",
        "issue_count": 0,
        "requirements_passed": 0,
        "requirements_needs_review": 0,
        "requirements_failed": 0,
    }


def complete_intake(
    matter: dict[str, Any],
    *,
    repository: MatterRepository,
    owner_user_id: str = "",
    drive_sync_runner: BackgroundRunner = run_in_daemon_thread,
) -> None:
    RepositoryMatterLifecycle(repository).complete_intake(
        matter,
        owner_user_id=owner_user_id,
        drive_sync_runner=drive_sync_runner,
    )
