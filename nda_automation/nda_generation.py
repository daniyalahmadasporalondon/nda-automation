"""Generate a company NDA from the Generic NDA template + the Playbook.

The generation engine fills the company's Generic NDA template (the structural
frame and boilerplate) with deal variables, and aligns the *substantive*
clauses to the live Playbook — the Playbook is the authority on clause wording,
the template is the authority on structure. The output is a ``.docx`` plus a
machine-readable manifest describing every fill, so the result can be verified
deterministically and saved as a tracked artifact.

Division of authority (confirmed in the Phase-1 mapping):

* **Template** owns the frame: recitals, the boilerplate clauses (NO OBLIGATION,
  USE & NON-DISCLOSURE, COPIES, IP, REMEDIES, CONFIRMATIONS, NO WARRANTIES,
  RETURN, WAIVER/entire-agreement, SEVERABILITY), the party roles (Aspora is
  always the SECOND party, the counterparty the FIRST), and the signature block.
* **Playbook** owns the substantive positions. Two template clauses drift off
  the Playbook and are realigned at generation time:
    - **Term and survival**: the template caps at "two (2) years" with no
      survival carve-out; the Playbook caps the term at ``max_term_years`` and
      requires a trade-secret / legal / data-protection survival carve-out.
    - **Confidential Information exclusions**: the template lists three carve-outs
      and omits the Playbook's "independently developed without use of
      Confidential Information" exclusion.
* **non_circumvention** is a *prohibited* Playbook position. The template omits
  it by design; generation must never introduce one.
* **Governing law** fills from the chosen entity's approved position.

The deterministic core fills variables and realigns the two clauses so the
output passes the Playbook with zero failures. An optional :class:`ClauseAdapter`
seam lets an AI adapt phrasing to the deal *without* changing which position the
clause takes — adaptation is constrained to the slots, never the substance.
"""

from __future__ import annotations

import datetime as _dt
import os
import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Mapping, Protocol

from docx import Document
from docx.document import Document as DocxDocument

from . import telemetry
from .checks.common import _year_count_label
from .playbook_runtime import ActivePlaybookBundle

# The tracked template asset (the company's Generic NDA). Resolved relative to
# this module so it works from any worktree / install location.
TEMPLATE_PATH = Path(__file__).resolve().parent / "templates" / "generic_nda.docx"

# NDA posture. v1 builds the mutual path (the template is mutual). One-way is a
# possible add-on; it is declared here so callers and manifests can name it, but
# only ``mutual`` is implemented.
NDA_TYPE_MUTUAL = "mutual"
NDA_TYPE_ONE_WAY = "one_way"

# Playbook clause ids whose substantive wording this engine realigns.
CLAUSE_TERM = "term_and_survival"
CLAUSE_CONFIDENTIAL = "confidential_information"
CLAUSE_GOVERNING_LAW = "governing_law"

# Kill-switch for the generation prose-polish AI. The AI clause adapter only
# rephrases already-on-position clause text (it never changes which position a
# clause takes), so disabling it yields the same Playbook-compliant document via
# the pure deterministic path — just instantly and with zero OpenRouter calls.
# DEFAULTS TO ENABLED when unset/blank, so existing behaviour and tests are
# preserved; set ``NDA_GENERATION_AI_ENABLED=false`` (or 0/no/off) to force the
# deterministic-only path. Live on prod as the reliability default.
GENERATION_AI_ENABLED_ENV = "NDA_GENERATION_AI_ENABLED"


def generation_ai_enabled() -> bool:
    """Whether the generation clause-adapter AI may be built.

    Default-ENABLED: returns ``True`` unless the env flag is explicitly set to a
    falsey value (``false``/``0``/``no``/``off``, case-insensitive). An unset or
    blank value preserves the original AI-on behaviour. This is the single seam
    that decides whether ``generate_nda_for_entity`` constructs the clause
    adapter at all — when it returns ``False`` no OpenRouter client is built and
    no AI call is made, so generation runs the pure deterministic Playbook path.
    """

    raw = os.environ.get(GENERATION_AI_ENABLED_ENV, "").strip().lower()
    if not raw:
        return True
    return raw not in {"false", "0", "no", "off"}

# --------------------------------------------------------------------------- #
# DocuSign signature anchors
# --------------------------------------------------------------------------- #
# DocuSign does NOT auto-detect signature spots — the envelope tells it where to
# drop each signer's signHere / dateSigned tab via an *anchor string* it searches
# for in the document. For a generated NDA we own the template, so we plant a
# DISTINCT, per-party anchor token on each party's signature line and route that
# party's tabs to it. This is the reliable case the send-for-signature flow relies
# on (``docusign_workflow`` imports these exact constants, so what is written and
# what is anchored can never drift).
#
# Why explicit tokens rather than reusing existing block text: after the signature
# block is normalised, BOTH party cells are structurally identical ("For <party>",
# "By: ___", "Name: ...", ...), and the only distinguishing text — the party's
# legal name — also appears earlier in the document (the party-introduction
# recital). DocuSign anchors to the FIRST occurrence of a string, so reusing the
# legal name would land the field on the recital, not the signature line, and the
# two cells could not be told apart. The tokens below are unique, appear once each,
# and sit exactly on the signature line.
#
# Token shape: backslash-delimited (``\sig_party_aspora\``) — the conventional
# DocuSign anchor-marker form. It carries no square brackets (so it never trips the
# leftover-placeholder guard, which only flags ``[...]`` template slots) and
# matches none of the prohibited-position patterns (so the ship gate leaves it
# alone). The tokens are designed to be hidden at send time (white/small font) but
# are searchable text in the .docx so DocuSign can find them.
SIGNATURE_ANCHOR_ASPORA = "\\sig_party_aspora\\"
SIGNATURE_ANCHOR_COUNTERPARTY = "\\sig_party_counterparty\\"

# The signer ``role`` (see ``docusign_workflow``) each anchor belongs to, so the
# workflow can map "this recipient is the Aspora signatory / the counterparty" to
# the right anchor without re-encoding the token strings.
SIGNATURE_ANCHOR_BY_ROLE = {
    "aspora": SIGNATURE_ANCHOR_ASPORA,
    "counterparty": SIGNATURE_ANCHOR_COUNTERPARTY,
}

# --------------------------------------------------------------------------- #
# Untrusted free-text validation (purpose / business_description)
# --------------------------------------------------------------------------- #
# intake.purpose and intake.business_description are free text from the caller
# that gets filled verbatim into the recital / [BUSINESS DESCRIPTION] slot. That
# makes them a surface for TWO distinct kinds of bad input, and we handle them
# DIFFERENTLY:
#
#   1. INJECTION (the ``drafter_instruction`` family + a one-way *posture* ask):
#      text trying to manipulate the drafter / AI ("ignore previous instructions",
#      "note to drafter", "add a clause", "make this one-way"). This is NOT
#      legitimate business prose the user expects to keep, so we never let it reach
#      the document AND we surface a clear validation error rather than silently
#      rewriting it. The patterns are an injection *defence*, so they stay
#      hardcoded here (they are not a Playbook legal position).
#
#   2. PROHIBITED LEGAL POSITION (non_compete, non_solicit, non_circumvention,
#      exclusivity, ip_assignment, penalty, perpetual_confidentiality, ...): a real
#      off-position legal restraint asked for in the recital. These are sourced
#      from the PLAYBOOK (the single source of truth) via
#      ``prohibited_positions.PROHIBITED_POSITION_PATTERNS`` -- no divergent
#      hardcoded legal-position list lives here any more. We FLAG-AND-SURFACE: a
#      clear 400 naming the field and the position so the user can rephrase, rather
#      than silently substituting safe boilerplate (which let the signed NDA recite
#      something different from what the user typed).
#
# Both kinds raise :class:`FreeTextValidationError` (an ``NdaGenerationError``
# subclass the generate route already turns into a 400). The audit trail is kept
# via a telemetry counter per kind and the flagged field/family carried on the
# raised exception -- generation aborts before a manifest exists, so there is no
# silent rewrite to record on the document.

# INJECTION patterns -- hardcoded here because they are an injection defence, not a
# Playbook legal position. ``drafter_instruction`` is prompt-injection; ``one_way``
# is a posture-manipulation ask (the engine only emits the mutual template, so a
# free-text "make this one-way" is an attempt to steer the document, not a clause
# the Playbook models).
_INJECTION_FREE_TEXT_PATTERNS: tuple[tuple[str, "re.Pattern[str]"], ...] = (
    ("drafter_instruction", re.compile(r"ignore (?:all )?(?:prior|previous) instructions|note to drafter|you are drafting|add a clause|do not mention", re.IGNORECASE)),
    ("one_way", re.compile(r"one-?way nda|binding only (?:on|the) (?:the )?receiving party|only (?:the )?receiving party|we receive only", re.IGNORECASE)),
)

# Plain-language names for the flagged families, used in the user-facing error so
# the message is actionable rather than a raw label.
_FREE_TEXT_FAMILY_LABELS: dict[str, str] = {
    "drafter_instruction": "drafter instruction",
    "one_way": "one-way posture request",
    "non_compete": "non-compete",
    "non_solicit": "non-solicit",
    "non_circumvention": "non-circumvention",
    "exclusivity": "exclusivity",
    "ip_assignment": "IP assignment",
    "penalty": "penalty / liquidated damages",
    "perpetual_confidentiality": "perpetual confidentiality",
    "auto_renew_lock": "auto-renewal lock",
}


def _family_label(label: str) -> str:
    return _FREE_TEXT_FAMILY_LABELS.get(label, label.replace("_", " "))


def detect_injection(text: str) -> str:
    """Return the first injection family label found in ``text``, or "".

    Injection = a drafter instruction or a one-way posture ask. Hardcoded defence,
    deliberately kept separate from the Playbook legal-position scan."""
    clean = str(text or "").strip()
    if not clean:
        return ""
    for label, pattern in _INJECTION_FREE_TEXT_PATTERNS:
        if pattern.search(clean):
            return label
    return ""


def detect_prohibited_position(text: str) -> str:
    """Return the first prohibited legal-position family in ``text``, or "".

    Sourced from the Playbook via ``prohibited_positions`` so this scan and the
    ship gate / adapter guard can never drift apart."""
    clean = str(text or "").strip()
    if not clean:
        return ""
    from .prohibited_positions import first_prohibited_position  # noqa: PLC0415

    return first_prohibited_position(clean)


class NdaGenerationError(ValueError):
    """Raised when generation inputs are invalid or the template is malformed."""


class FreeTextValidationError(NdaGenerationError):
    """Raised when a free-text intake field (purpose / business_description) carries
    an injection attempt or a prohibited legal position.

    A subclass of :class:`NdaGenerationError` so the generate route already turns it
    into a 400 ``{"error": ...}`` with a clear, field-scoped message -- the user is
    told their wording was rejected and why, instead of having it silently rewritten.
    Carries the flagged ``field_name``, ``family`` (the raw label) and ``kind``
    (``"injection"`` | ``"position"``) for telemetry / audit.
    """

    def __init__(self, message: str, *, field_name: str = "", family: str = "", kind: str = "") -> None:
        super().__init__(message)
        self.field_name = field_name
        self.family = family
        self.kind = kind


def _validate_free_text(value: str, *, field_name: str) -> None:
    """Validate one free-text field; raise if it is unsafe to fill verbatim.

    Raises :class:`FreeTextValidationError` (-> 400) when the field carries an
    injection attempt or a prohibited legal position, recording an audit telemetry
    counter for the flagged kind. Clean input returns ``None``. Identity fields are
    NOT scanned here (they are structured slots, validated by
    :func:`validate_intake_identity_fields` and the entity gate)."""

    text = str(value or "").strip()
    if not text:
        return

    injection = detect_injection(text)
    if injection:
        telemetry.increment("generate_nda_free_text_injection_blocked")
        raise FreeTextValidationError(
            f"The {field_name} field contains content that cannot be included "
            f"({_family_label(injection)}). Please rephrase using plain business language.",
            field_name=field_name,
            family=injection,
            kind="injection",
        )

    position = detect_prohibited_position(text)
    if position:
        telemetry.increment("generate_nda_free_text_position_flagged")
        raise FreeTextValidationError(
            f"The {field_name} field describes a prohibited position "
            f"({_family_label(position)}). Please revise the wording or remove it.",
            field_name=field_name,
            family=position,
            kind="position",
        )


def _validate_intake_free_text(intake: "CounterpartyIntake") -> None:
    """Validate the free-text intake fields before they are filled into the document.

    FLAG-AND-SURFACE (replacing the old silent-replace): the two free-text slots
    that flow into prose (purpose, business description) are scanned for an injection
    attempt or a prohibited legal position, and a hit raises
    :class:`FreeTextValidationError` (a 400 to the caller) naming the field and the
    flagged family -- the user's typed wording is NEVER silently substituted, so a
    generated/signed NDA can never recite something different from what they typed.
    Identity fields (names/addresses) are filled into structured slots and checked by
    the entity gate, not the position scan. Returns ``None`` when the intake is clean;
    the audit trail for a flag is the telemetry counter + the raised exception's
    ``field_name`` / ``family`` (no document/manifest is produced for a flagged
    intake)."""

    # purpose first so its error wins when both are dirty (it is the primary recital
    # field). Each raises on the first flagged family.
    for value, field_name in (
        (intake.purpose, "purpose"),
        (intake.business_description, "business_description"),
    ):
        _validate_free_text(value, field_name=field_name)


# The counterparty identity fields the caller supplies via intake. Unlike the
# free-text recital fields (purpose / business_description, which are sanitised),
# these are filled verbatim into the document's structured identity slots and so
# must NOT be altered — a name/address/jurisdiction is a legal value. They are
# validated and *rejected* (not rewritten) when they carry a square bracket,
# because the engine fills template tokens of the form ``[GOVERNING LAW]`` /
# ``[COMPANY NAME]`` and a bracket in an identity value either collides with a
# real fill token (tripping the fail-closed leftover-placeholder guard with an
# opaque error) or silently lands stray ``[...]`` text in a signed NDA.
_IDENTITY_INTAKE_FIELDS: tuple[tuple[str, str], ...] = (
    ("company_name", "company name"),
    ("registered_office", "registered office"),
    ("jurisdiction_of_incorporation", "jurisdiction of incorporation"),
)


def validate_intake_identity_fields(intake: "CounterpartyIntake") -> None:
    """Reject counterparty identity values that conflict with the fill markers.

    The document is filled by substituting square-bracket template tokens (e.g.
    ``[COMPANY NAME]``, ``[GOVERNING LAW]``). A square bracket in a user-supplied
    identity value therefore either collides with one of those markers — which
    used to surface as the opaque ``"Generated NDA still contains unfilled
    placeholders: ..."`` failure with no hint of which field caused it — or leaves
    stray bracketed text in a signed legal document. We REJECT such values rather
    than silently sanitising them (a name/address is a legal value the engine must
    never alter), raising a clear, field-scoped :class:`NdaGenerationError` that
    names the offending field and the reason. The first offending field is
    reported (the caller fixes one field per attempt).
    """

    for attr, label in _IDENTITY_INTAKE_FIELDS:
        value = str(getattr(intake, attr, "") or "")
        if "[" in value or "]" in value:
            bad = "[" if "[" in value else "]"
            raise NdaGenerationError(
                f"The {label} ({attr}) contains a square bracket '{bad}', which conflicts "
                f"with the NDA template's fill markers (e.g. [GOVERNING LAW]). Remove the "
                f"bracketed text from {attr} and try again."
            )


@dataclass(frozen=True)
class EntityParty:
    """The Aspora signing entity (the SECOND party) for this NDA.

    Mirrors the fields the engine consumes from an ``entity_registry`` bundle —
    constructed via :func:`entity_party_from_bundle` so the registry stays the
    single source of entity truth.
    """

    legal_name: str
    registered_office: str
    jurisdiction_of_incorporation: str
    governing_law_value: str
    forum: str
    # Empty signatory name/title means "no signer assigned" — the signature block
    # renders blank fill-lines (underscores) for signing, never bracketed text.
    signatory_name: str = ""
    signatory_title: str = ""
    entity_id: str = ""
    # The registry address id actually written into ``registered_office`` (the
    # user's pick, or the default). Recorded on the manifest for provenance parity
    # with the governing-law override.
    registered_office_address_id: str = ""


@dataclass(frozen=True)
class CounterpartyIntake:
    """The counterparty (FIRST party) + deal variables, from draft-ui intake."""

    company_name: str
    registered_office: str
    jurisdiction_of_incorporation: str
    business_description: str
    purpose: str
    term_years: int = 2
    nda_type: str = NDA_TYPE_MUTUAL
    agreement_date: _dt.date | None = None


@dataclass
class GenerationManifest:
    """Ground-truth record of what was filled — for verification + provenance.

    The verifier (gen-verify) diffs its structural / entity assertions against
    this instead of re-parsing the prose, so a fill is checked against intent
    rather than against a regex over the output.
    """

    entity_id: str
    entity_legal_name: str
    counterparty_name: str
    nda_type: str
    term_years: int
    agreement_date: str
    governing_law_value: str
    forum: str
    # The registry address id whose lines were written as the entity's registered
    # office. Records the user's address pick (or the default) for provenance parity
    # with the governing-law override, so gen-verify / audit can confirm WHICH office
    # was used rather than re-parsing the prose.
    entity_address_id: str = ""
    slot_fills: dict[str, str] = field(default_factory=dict)
    clause_alignments: list[str] = field(default_factory=list)
    # Free-text intake fields whose untrusted content was neutralised before it
    # could reach the document, e.g. "purpose: removed prohibited position
    # (non_solicit)". Empty in the normal case; auditable when an injection was
    # caught (the UI can surface that the purpose/description was sanitised).
    sanitized_fields: list[str] = field(default_factory=list)
    # Governing-law provenance. ``governing_law_value`` above is always the
    # EFFECTIVE law written into the clause; ``governing_law_option_id`` is its
    # Playbook approved-option id (the join key). When the user picked a non-default
    # (but still Playbook-approved) law, ``governing_law_overridden`` is True and
    # ``entity_default_governing_law_value`` records the entity's registry default,
    # so gen-verify validates against the INTENDED law rather than flagging it as a
    # mismatch.
    governing_law_option_id: str = ""
    governing_law_overridden: bool = False
    entity_default_governing_law_value: str = ""

    def to_dict(self) -> dict[str, Any]:
        return {
            "entity_id": self.entity_id,
            "entity_legal_name": self.entity_legal_name,
            "counterparty_name": self.counterparty_name,
            "nda_type": self.nda_type,
            "term_years": self.term_years,
            "agreement_date": self.agreement_date,
            "governing_law_value": self.governing_law_value,
            "forum": self.forum,
            "entity_address_id": self.entity_address_id,
            "slot_fills": dict(self.slot_fills),
            "clause_alignments": list(self.clause_alignments),
            "sanitized_fields": list(self.sanitized_fields),
            "governing_law_option_id": self.governing_law_option_id,
            "governing_law_overridden": self.governing_law_overridden,
            "entity_default_governing_law_value": self.entity_default_governing_law_value,
        }


@dataclass
class GenerationResult:
    """The generated NDA bytes + its manifest."""

    docx_bytes: bytes
    manifest: GenerationManifest


class ClauseAdapter(Protocol):
    """Optional AI seam: adapt clause phrasing to the deal, on-position only.

    ``adapt(clause_id, playbook_text, context)`` returns adapted text for the
    given Playbook clause. The deterministic core supplies ``playbook_text`` as
    the authoritative position; an adapter may polish phrasing for the deal but
    must keep the position. When no adapter is supplied the Playbook text is used
    verbatim, which is why the engine is fully testable offline.
    """

    def adapt(self, clause_id: str, playbook_text: str, context: Mapping[str, Any]) -> str: ...


def entity_party_from_bundle(
    bundle: Mapping[str, Any],
    playbook: Mapping[str, Any],
    *,
    governing_law_option_id: str = "",
    address_id: str = "",
) -> EntityParty:
    """Build an :class:`EntityParty` from an ``entity_registry`` bundle.

    Resolves the governing-law *value* through the Playbook so the slot text
    always matches an approved option: the bundle carries a
    ``governing_law.playbook_option_id`` which joins onto the Playbook's
    ``governing_law.rules.approved_options[].id``. A bundle whose option is not
    approved is rejected — generation never emits an off-position governing law.

    ``governing_law_option_id`` overrides the entity's default law with another
    Playbook-approved option (the product lets the user pick a different — still
    approved — governing law). An unapproved override is rejected. When the law is
    overridden, the forum tracks the overridden law (rather than the entity's
    default courts) so the draft never pairs one jurisdiction's law with another's
    forum.

    ``address_id`` selects which registry address fills the entity's registered
    office (the product lets the user pick a non-default office). A blank id falls
    back to the default-flagged address; an unknown id is rejected (same guard shape
    as the governing-law override). The chosen id is carried on the returned
    :class:`EntityParty` so the manifest can record it for provenance.
    """

    legal_name = str(bundle.get("legal_name") or "").strip()
    if not legal_name:
        raise NdaGenerationError("Entity bundle is missing legal_name.")

    approved = _approved_governing_law_options(playbook)
    default_option_id = str((bundle.get("governing_law") or {}).get("playbook_option_id") or "").strip()
    if default_option_id not in approved:
        raise NdaGenerationError(
            f"Entity governing-law option {default_option_id!r} is not an approved Playbook option "
            f"(approved: {sorted(approved)})."
        )

    override_id = str(governing_law_option_id or "").strip()
    overridden = bool(override_id) and override_id != default_option_id
    if override_id and override_id not in approved:
        raise NdaGenerationError(
            f"Governing-law override {override_id!r} is not an approved Playbook option "
            f"(approved: {sorted(approved)})."
        )
    option_id = override_id or default_option_id
    governing_law_value = approved[option_id]

    address, chosen_address_id = select_bundle_address(bundle, address_id)
    signatory = bundle.get("signatory") or {}
    # The entity's registered forum is correct for its default law. When the law is
    # overridden to a different option, derive the forum that goes WITH the chosen
    # option (each approved option's proper forum is registered against the entity
    # that defaults to it, e.g. england_and_wales -> "Courts of England and Wales"),
    # so we never pair one jurisdiction's law with another's courts. If that lookup
    # is unavailable, fall back to the law value rather than the wrong forum.
    if not overridden:
        forum = str(bundle.get("jurisdiction") or "").strip() or governing_law_value
    else:
        forum = _forum_for_option_id(option_id)
    forum = _require_court_forum(forum, option_id, governing_law_value)

    return EntityParty(
        legal_name=legal_name,
        registered_office=address,
        jurisdiction_of_incorporation=_incorporation_jurisdiction(bundle, governing_law_value),
        governing_law_value=governing_law_value,
        forum=forum,
        # The registry ships placeholder signatory strings ("[Authorised Signatory]"
        # / "[Title]") when no real signer is assigned; treat those as unassigned so
        # the signature block renders clean blank fill-lines, not bracketed text.
        signatory_name=_real_or_blank(signatory.get("name")),
        signatory_title=_real_or_blank(signatory.get("title")),
        entity_id=str(bundle.get("id") or "").strip(),
        registered_office_address_id=chosen_address_id,
    )


# The proper COURT/VENUE that goes with each Playbook-approved governing-law
# option id. This is the canonical fallback when no registry entity defaults to
# an overridden option (so the registry lookup in ``_forum_for_option_id``
# returns nothing). Without it the old code wrote the bare LAW NAME (e.g. "DIFC",
# "Delaware") into the forum/submission clause -- a jurisdiction string that is
# NOT a court. Keyed by the approved-option id from playbook.json.
_COURT_FOR_OPTION_ID: dict[str, str] = {
    "india": "the courts of Mumbai, India",
    "delaware": "the state and federal courts located in the State of Delaware",
    "england_and_wales": "the courts of England and Wales",
    "difc": "the DIFC Courts, Dubai",
    "ontario_canada": "the courts of the Province of Ontario, Canada",
}


def _forum_for_option_id(option_id: str) -> str:
    """The proper court/forum for a governing-law option.

    The Playbook approved_options carry only id/label/value (no forum). Resolution
    order:

    1. The registry: each entity registers a ``jurisdiction`` that goes WITH its
       default governing-law option (e.g. england_and_wales -> "Courts of England
       and Wales"), so an overridden option's forum is taken from whichever
       registry entity defaults to that option.
    2. The canonical :data:`_COURT_FOR_OPTION_ID` court map -- used when NO registry
       entity defaults to the option (the gap that previously made this return ""
       and let the caller write the bare law name as the forum).

    Returns "" only for an option id we have no court for at all, which the caller
    (:func:`_require_court_forum`) turns into a hard refusal rather than writing a
    non-court venue into a signed NDA."""

    try:
        from . import entity_registry  # noqa: PLC0415

        for bundle in entity_registry.list_entities():
            bundle_option = str((bundle.get("governing_law") or {}).get("playbook_option_id") or "").strip()
            if bundle_option == option_id:
                forum = str(bundle.get("jurisdiction") or "").strip()
                if forum:
                    return forum
    except Exception:  # noqa: BLE001 - registry absent/!importable -> fall to court map
        pass
    return _COURT_FOR_OPTION_ID.get(option_id, "")


def _require_court_forum(forum: str, option_id: str, governing_law_value: str) -> str:
    """Hard gate: refuse to write a non-court venue into the NDA.

    The forum/submission clause must name a COURT, not a bare jurisdiction/law
    label. The old fallback ``forum = _forum_for_option_id(...) or
    governing_law_value`` wrote the LAW NAME (e.g. "DIFC", "Delaware") whenever no
    forum resolved -- a value that is not a court. The manifest flagged it but the
    doc still rendered/sent/signed. We now REFUSE generation rather than emit such
    a venue: a resolved court is required, and a resolved value that merely echoes
    the law name is rejected as non-court."""

    forum = str(forum or "").strip()
    if not forum or forum == governing_law_value.strip():
        raise NdaGenerationError(
            f"Refusing to generate: no court/venue is defined for governing-law option "
            f"{option_id!r} (law {governing_law_value!r}). Writing the law name as the "
            "forum would put a non-court jurisdiction into a signed NDA. Register a "
            "forum for this option (entity registry jurisdiction or the court map)."
        )
    return forum


def _real_or_blank(value: object) -> str:
    """Return a real signatory value, or "" if it is empty or a bracketed placeholder.

    The registry uses ``"[Authorised Signatory]"`` / ``"[Title]"`` as placeholders
    for an unassigned signer. We map those (and any ``[...]`` form) to an empty
    string so the signature block renders blank fill-lines for signing.
    """

    text = str(value or "").strip()
    if not text or (text.startswith("[") and text.endswith("]")):
        return ""
    return text


def _incorporation_jurisdiction(bundle: Mapping[str, Any], governing_law_value: str) -> str:
    """The entity's jurisdiction of incorporation.

    Consumes the registry's explicit ``incorporation_jurisdiction`` field when
    present (the authoritative value, e.g. "Delaware, United States" for a
    Delaware LLC). Falls back to the governing-law value, which the registry
    confirms matches incorporation for all current entities.
    """

    explicit = str(bundle.get("incorporation_jurisdiction") or "").strip()
    return explicit or governing_law_value


def generate_nda(
    entity: EntityParty,
    intake: CounterpartyIntake,
    *,
    playbook: Mapping[str, Any],
    template_path: Path | str = TEMPLATE_PATH,
    clause_adapter: ClauseAdapter | None = None,
) -> GenerationResult:
    """Generate the NDA ``.docx`` + manifest from the template and Playbook.

    This is the core seam: ``generate_nda(entity, intake) -> (docx, manifest)``.
    Saving the result as a tracked artifact is a separate, optional step
    (:func:`save_generated_nda`) so this function has no storage dependency and
    runs fully offline.
    """

    if intake.nda_type != NDA_TYPE_MUTUAL:
        # The template is mutual; one-way support is not implemented in v1.
        raise NdaGenerationError(
            f"nda_type {intake.nda_type!r} is not supported yet; only {NDA_TYPE_MUTUAL!r} is implemented."
        )

    # Reject counterparty identity values that collide with the fill markers
    # BEFORE any fill runs, with a clear field-scoped error (replacing the opaque
    # leftover-placeholder failure). Identity values are never altered.
    validate_intake_identity_fields(intake)

    term_years = _resolve_term_years(intake.term_years, playbook)
    agreement_date = intake.agreement_date or _dt.date.today()

    document = _load_template(template_path)

    manifest = GenerationManifest(
        entity_id=entity.entity_id,
        entity_legal_name=entity.legal_name,
        counterparty_name=intake.company_name,
        nda_type=intake.nda_type,
        term_years=term_years,
        agreement_date=agreement_date.isoformat(),
        governing_law_value=entity.governing_law_value,
        forum=entity.forum,
        entity_address_id=entity.registered_office_address_id,
    )

    # Validate untrusted free text BEFORE it reaches either the document fills or
    # the AI adapter context, so an injected drafter-instruction / one-way ask or a
    # prohibited legal position can never land in the recital or steer the adapter.
    # A hit raises FreeTextValidationError (-> 400) rather than silently rewriting
    # the user's wording; nothing is filled and no document is produced.
    _validate_intake_free_text(intake)

    _fill_variable_slots(document, entity, intake, agreement_date, manifest)

    # When the AI clause adapter is active, adapt the three clauses CONCURRENTLY
    # rather than back-to-back. Each adaptation is independent and individually
    # guarded (a failure/timeout/drift degrades to that clause's deterministic
    # Playbook base text), so overlapping the network waits cuts the AI portion of
    # a generate from ~the sum of three calls to ~the slowest single call — and the
    # tightened per-call timeout bounds that. With no adapter this resolves to an
    # empty map and each clause uses its deterministic base text exactly as before.
    adapted = _prefetch_clause_adaptations(playbook, clause_adapter, intake)
    _align_mutuality(
        document, playbook, clause_adapter, intake, manifest,
        adapted_text=adapted.get("mutuality"),
    )
    _align_confidential_information(
        document, playbook, clause_adapter, intake, manifest,
        adapted_text=adapted.get(CLAUSE_CONFIDENTIAL),
    )
    _align_term_and_survival(
        document, playbook, term_years, clause_adapter, intake, manifest,
        adapted_text=adapted.get(CLAUSE_TERM),
    )

    # Unassigned signatories render as blank fill-lines (not bracketed text), so
    # no bracketed values are ever legitimately retained — the guard is strict.
    _assert_no_unfilled_placeholders(document, allowed=set())

    return GenerationResult(docx_bytes=_document_bytes(document), manifest=manifest)


def save_generated_nda(
    result: GenerationResult,
    matter_id: str,
    *,
    add_artifact: Callable[..., Any] | None = None,
    repository: Any | None = None,
    based_on_artifact_id: str = "",
    owner_user_id: str = "",
) -> Any:
    """Persist a generated NDA artifact (actor=entity, role=generated).

    The actor is the entity id slug (e.g. ``aspora_technology``) — stable and
    short, so the auto-generated filename stays clean; the full legal name is
    preserved on the manifest metadata. The role is ``generated``;
    ``based_on_artifact_id`` records lineage when the generation derives from an
    existing matter artifact (e.g. the original), and is optional because a
    generated NDA is often the matter's first document.

    ``add_artifact`` defaults to the live ``artifact_service.add_artifact`` but
    stays injectable so the engine can be tested without the artifact registry.
    The import is deferred so importing this module never requires artifact-spine.
    """

    if add_artifact is None:
        from .artifact_service import add_artifact as add_artifact  # noqa: PLC0415

    actor = result.manifest.entity_id or result.manifest.entity_legal_name or "aspora"
    return add_artifact(
        matter_id,
        source="generated",
        actor=actor,
        role="generated",
        document_bytes=result.docx_bytes,
        based_on_artifact_id=based_on_artifact_id,
        repository=repository,
        owner_user_id=owner_user_id,
        metadata={"generation": result.manifest.to_dict()},
    )


def generate_nda_for_entity(
    entity_id: str,
    intake: CounterpartyIntake,
    *,
    playbook: Mapping[str, Any] | None = None,
    playbook_bundle: ActivePlaybookBundle | None = None,
    clause_adapter: ClauseAdapter | None = None,
    use_ai: bool = True,
    use_frozen: bool = False,
    governing_law_override: str = "",
    address_id: str = "",
) -> GenerationResult:
    """Resolve the entity from the registry and generate the NDA (no save).

    The single convenience entry the HTTP route, gen-verify, and tests share:
    ``entity_id`` + intake -> ``GenerationResult``. Resolves the entity via the
    live ``entity_registry`` and builds the AI clause adapter when ``use_ai`` and
    no explicit adapter is given (Playbook-bounded; falls back to deterministic
    with no API key). The live ``entity_registry`` import is deferred so this
    module imports cleanly where the registry is absent.

    ``use_frozen`` selects the repeatable golden-fixture adapter instead of the
    live (non-deterministic) AI adapter: same guarded AI codepath, but replaying
    recorded on-position clause text, so gen-verify can gate the AI-shaped output
    deterministically. An explicit ``clause_adapter`` overrides both.

    ``governing_law_override`` is a Playbook governing-law option id; when given
    and different from the entity default, the draft uses that (still approved)
    law and the manifest records the override provenance (``governing_law_overridden``
    + ``entity_default_governing_law_value``) so gen-verify validates the chosen
    law rather than flagging an entity mismatch.

    ``address_id`` is the registry address id the user picked for the entity's
    registered office; blank uses the entity default, an unknown id is rejected (see
    :func:`entity_party_from_bundle`). The chosen id lands on the manifest
    (``entity_address_id``).
    """

    from . import entity_registry  # noqa: PLC0415

    if playbook is None and playbook_bundle is not None:
        playbook = playbook_bundle.playbook
    if playbook is None:
        from .checker import load_playbook  # noqa: PLC0415

        playbook = load_playbook()

    if clause_adapter is None and use_frozen:
        from .nda_generation_ai import build_frozen_clause_adapter  # noqa: PLC0415

        clause_adapter = build_frozen_clause_adapter()
    elif clause_adapter is None and use_ai and generation_ai_enabled():
        # Build the live AI clause adapter only when the kill-switch is on (the
        # default). With NDA_GENERATION_AI_ENABLED=false we skip construction
        # entirely: no OpenRouter client, no network call — ``clause_adapter``
        # stays ``None`` and generation runs the pure deterministic Playbook
        # path. An explicit ``clause_adapter`` (tests) and ``use_frozen`` (the
        # gen-verify golden fixture) are unaffected — neither calls OpenRouter.
        from .nda_generation_ai import build_clause_adapter  # noqa: PLC0415

        clause_adapter = build_clause_adapter()

    bundle = entity_registry.get_entity(entity_id)
    if bundle is None:
        raise NdaGenerationError(f"Unknown signing entity {entity_id!r}.")

    approved = _approved_governing_law_options(playbook)
    default_option_id = str((bundle.get("governing_law") or {}).get("playbook_option_id") or "").strip()
    entity_default_law = approved.get(default_option_id, "")
    override_id = str(governing_law_override or "").strip()
    overridden = bool(override_id) and override_id != default_option_id
    effective_option_id = override_id or default_option_id

    entity = entity_party_from_bundle(
        bundle, playbook, governing_law_option_id=override_id, address_id=address_id
    )
    result = generate_nda(entity, intake, playbook=playbook, clause_adapter=clause_adapter)
    # Record governing-law provenance on the manifest (the effective law value +
    # forum are already on the manifest; these let gen-verify validate against the
    # INTENDED option rather than the entity default).
    result.manifest.governing_law_option_id = effective_option_id
    result.manifest.governing_law_overridden = overridden
    result.manifest.entity_default_governing_law_value = entity_default_law
    return result


def generate_and_save_nda(
    entity_id: str,
    intake: CounterpartyIntake,
    matter_id: str,
    *,
    playbook: Mapping[str, Any] | None = None,
    playbook_bundle: ActivePlaybookBundle | None = None,
    repository: Any | None = None,
    based_on_artifact_id: str = "",
    owner_user_id: str = "",
    clause_adapter: ClauseAdapter | None = None,
    use_ai: bool = True,
    governing_law_override: str = "",
    address_id: str = "",
) -> tuple[GenerationResult, Any]:
    """End-to-end: resolve the entity, generate the NDA, save it as an artifact.

    Wires the live ``entity_registry`` (the single source of entity truth) and
    ``artifact_service`` so a caller only needs an ``entity_id`` + intake +
    ``matter_id``. Returns ``(result, artifact)``. ``governing_law_override`` is a
    Playbook governing-law option id (see :func:`generate_nda_for_entity`).
    """

    if playbook is None and playbook_bundle is not None:
        playbook = playbook_bundle.playbook
    if playbook is None:
        from .checker import load_playbook  # noqa: PLC0415

        playbook = load_playbook()

    result = generate_nda_for_entity(
        entity_id,
        intake,
        playbook=playbook,
        playbook_bundle=playbook_bundle,
        clause_adapter=clause_adapter,
        use_ai=use_ai,
        governing_law_override=governing_law_override,
        address_id=address_id,
    )
    # Hard pre-save gate: never persist an off-position draft. The clause adapter
    # is guarded and the intake is sanitised, but this is the last, independent
    # backstop on the SHIP path — a generated NDA that fails its own Playbook or
    # carries a prohibited position must not become a saved artifact.
    assert_generated_nda_is_on_position(result, playbook)
    artifact = save_generated_nda(
        result,
        matter_id,
        repository=repository,
        based_on_artifact_id=based_on_artifact_id,
        owner_user_id=owner_user_id,
    )
    return result, artifact


# --------------------------------------------------------------------------- #
# Pre-save ship gate — the last backstop before a draft becomes an artifact
# --------------------------------------------------------------------------- #
# The prohibited-position families are the SHARED canonical set
# (prohibited_positions.PROHIBITED_POSITION_PATTERNS) so this gate, the AI-adapter
# guard, and gen-verify's independent gate all enforce the same bar. A hit anywhere
# in the generated document (outside the permitted narrow survival carve-out) means
# the draft asserts a position the Playbook bans -> refuse to save.

# A clear, stable message the endpoint surfaces when the ship gate trips.
SAFETY_GATE_MESSAGE = "Generated NDA failed the safety gate and was not saved"


def assert_generated_nda_is_on_position(result: "GenerationResult", playbook: Mapping[str, Any] | None = None) -> None:
    """Refuse to ship a generated NDA that fails its own Playbook or carries a
    prohibited position. Raises :class:`NdaGenerationError` (which the endpoint
    surfaces as a 4xx with a "safety gate" message) so an off-position draft never
    reaches storage.

    PUBLIC so every persistence path can call it — both the convenience
    ``generate_and_save_nda`` AND the HTTP route (which builds the matter +
    artifact itself) MUST run this BEFORE any save, or the AI-first safety
    contract has a hole on the actual ship path.

    This is the load-bearing AI-first backstop on the SHIP path: the per-clause
    adapter guard can fall back to deterministic wording, but a drifted live-AI
    clause that slips through (or a prohibited position outside the six scored
    clauses) must still be caught here before save. Two independent screens:
    1. ``self_check_generated_nda`` — zero native fails AND non_circumvention does
       not fail (the deterministic + AI-first oracles). self_check ALONE is not
       enough: it only scores the six hard clauses, so a non-compete appended to
       mutuality leaves mutuality passing — hence screen 2.
    2. a meaning-based prohibited-position scan (the shared canonical set) over the
       whole document, so a position smuggled OUTSIDE the scored clauses is caught.
    The narrow trade-secret / legal / data-protection survival carve-out the
    Playbook permits ("for as long as ... required") is exempt from the perpetual
    screen so the legitimate survival sentence is not mistaken for a position.
    """

    from .docx_text import extract_docx_text  # noqa: PLC0415
    from .prohibited_positions import PROHIBITED_POSITION_PATTERNS  # noqa: PLC0415

    check = self_check_generated_nda(result.docx_bytes, playbook=playbook)
    if not check.passed:
        problems = check.native_failures + check.dynamic_failures
        raise NdaGenerationError(
            f"{SAFETY_GATE_MESSAGE}: failed its own Playbook ("
            + ", ".join(problems or ["unknown failure"])
            + ")."
        )

    text = extract_docx_text(result.docx_bytes)
    # Scan PER SENTENCE so the permitted-survival exemption applies only to the
    # specific sentence that carries the carve-out scope -- not the whole document
    # (the legitimate trade-secret/data-protection survival sentence must not
    # license a perpetual claim made elsewhere in the draft).
    for sentence in _scan_sentences(text):
        for label, pattern in PROHIBITED_POSITION_PATTERNS:
            if not pattern.search(sentence):
                continue
            if label == "perpetual_confidentiality" and _is_permitted_long_survival(sentence):
                continue
            raise NdaGenerationError(
                f"{SAFETY_GATE_MESSAGE}: carries a prohibited position ({label})."
            )


def _scan_sentences(text: str) -> list[str]:
    """Split the document into sentence/line units for the position scan."""
    return [s.strip() for s in re.split(r"(?<=[.;])\s+|\n+", text) if s.strip()]


def _is_permitted_long_survival(sentence: str) -> bool:
    """The Playbook permits a NARROW trade-secret / legal / data-protection
    survival obligation to last as long as the protected status or law requires.
    Perpetual-survival language scoped to that carve-out is on-position; the same
    language applied to ALL confidential information is drift. Mirrors gen-verify's
    permitted-long-survival logic so the ship gate and the verifier agree -- judged
    on the SAME sentence as the perpetual language, not the whole document."""

    lowered = re.sub(r"\s+", " ", sentence.lower())
    carve_outs = (
        "trade secret", "as long as", "so long as", "applicable law", "law requires",
        "legal obligation", "data protection", "data-protection",
    )
    if not any(token in lowered for token in carve_outs):
        return False
    # If the perpetual language in THIS sentence reaches ALL/ANY confidential
    # information it is no longer the narrow carve-out and is not permitted.
    blanket = ("all confidential information", "any and all information", "all information")
    return not any(token in lowered for token in blanket)


# --------------------------------------------------------------------------- #
# Self-check — the same oracle gen-verify uses
# --------------------------------------------------------------------------- #

# The Playbook's hard clauses, split by review engine. NATIVE clauses are scored
# deterministically (checker.review_nda); the DYNAMIC clause is only emitted on
# the AI-first path. These mirror gen-verify's split so both sides measure the
# generated NDA with the SAME oracle.
NATIVE_CLAUSE_IDS = (
    "mutuality",
    "confidential_information",
    "governing_law",
    "term_and_survival",
    "signatures",
)
DYNAMIC_CLAUSE_IDS = ("non_circumvention",)


@dataclass
class SelfCheckResult:
    """Outcome of the generation self-check."""

    passed: bool
    native_failures: list[str] = field(default_factory=list)
    native_reviews: list[str] = field(default_factory=list)
    dynamic_failures: list[str] = field(default_factory=list)
    overall_status: str = ""

    def __bool__(self) -> bool:
        return self.passed


def self_check_generated_nda(
    docx_bytes: bytes,
    *,
    playbook: Mapping[str, Any] | None = None,
) -> SelfCheckResult:
    """Run the generated NDA through the SAME oracle gen-verify uses.

    The native clauses are scored by the deterministic engine with ``verify=False``
    — the trustworthy, network-free oracle. The key-free *stub* AI reviewer is NOT
    used as the native oracle: it rubber-stamps native clauses and downgrades
    ungrounded passes to "review", which would make a stub-based self-check show a
    false green while a real native defect (e.g. missing survival carve-out or a
    broken execution block) slipped through.

    The dynamic ``non_circumvention`` clause is only emitted on the AI-first path,
    so it is checked separately via the key-free AI-first stub: the stub fails that
    clause iff a prohibited restriction is present, so a fail means the draft
    smuggled one in.

    A generated NDA passes iff it has zero native failures AND non_circumvention
    does not fail. (Native "review" verdicts are surfaced but are not failures —
    the generated NDA is expected to score 0 fails / 0 review, but the contract the
    bar enforces is 0 *fails*.)
    """

    from .checker import load_playbook, review_nda  # noqa: PLC0415
    from .docx_text import extract_docx_text  # noqa: PLC0415

    resolved_playbook = playbook if playbook is not None else load_playbook()
    text = extract_docx_text(docx_bytes)

    # Network-free oracle, as the docstring promises: verify=False drops the AI
    # verifier, ai_enabled=False drops the legacy per-clause AI overlay. Without
    # the latter this made ~5 live model calls (~20s) that never changed the
    # native verdict — and made a SAFETY GATE depend on a non-deterministic model.
    native = review_nda(text, playbook=resolved_playbook, verify=False, ai_enabled=False)
    native_by_id = {str(clause.get("id")): clause for clause in native.get("clauses", [])}
    native_failures: list[str] = []
    native_reviews: list[str] = []
    for clause_id in NATIVE_CLAUSE_IDS:
        clause = native_by_id.get(clause_id)
        if clause is None:
            native_failures.append(f"{clause_id} (not emitted by deterministic engine)")
            continue
        decision = str(clause.get("decision"))
        if decision == "fail":
            native_failures.append(clause_id)
        elif decision == "review":
            native_reviews.append(clause_id)

    dynamic_failures = _self_check_non_circumvention(text, resolved_playbook)

    passed = not native_failures and not dynamic_failures
    return SelfCheckResult(
        passed=passed,
        native_failures=native_failures,
        native_reviews=native_reviews,
        dynamic_failures=dynamic_failures,
        overall_status=str(native.get("overall_status") or ""),
    )


def _self_check_non_circumvention(text: str, playbook: Mapping[str, Any]) -> list[str]:
    """Key-free AI-first pass to surface the dynamic non_circumvention clause."""

    from .ai_assessor import (  # noqa: PLC0415
        _validate_ai_assessment_response,
        build_ai_assessment_packet,
        stub_ai_assessment_response,
    )
    from .ai_first_review import build_ai_first_review_result  # noqa: PLC0415

    packet = build_ai_assessment_packet(text, playbook=playbook)
    raw = stub_ai_assessment_response(packet)
    assessments = _validate_ai_assessment_response(raw, playbook=playbook, packet=packet)
    result = build_ai_first_review_result(text, assessments, playbook=playbook)
    by_id = {str(clause.get("id")): clause for clause in result.get("clauses", [])}

    failures: list[str] = []
    for clause_id in DYNAMIC_CLAUSE_IDS:
        clause = by_id.get(clause_id)
        if clause is None:
            failures.append(f"{clause_id} (not emitted by AI-first engine)")
        elif str(clause.get("decision")) == "fail":
            failures.append(clause_id)
    return failures


# --------------------------------------------------------------------------- #
# Variable-slot filling
# --------------------------------------------------------------------------- #


def _fill_variable_slots(
    document: DocxDocument,
    entity: EntityParty,
    intake: CounterpartyIntake,
    agreement_date: _dt.date,
    manifest: GenerationManifest,
) -> None:
    """Fill the 11 template placeholders, party-aware.

    ``[JURISDICTION OF INCORPORATION]`` and ``[REGISTERED OFFICE ADDRESS]`` each
    appear once per party, so they are filled positionally (first occurrence =
    Company / FIRST party, second = Aspora / SECOND party) rather than globally.
    The per-party fills are applied to the two specific party paragraphs; the
    remaining single-occurrence slots are filled globally.
    """

    day, month, year = _split_date(agreement_date)

    # Per-party paragraphs: the Company block names the FIRST party, the Aspora
    # block the SECOND. They are the two paragraphs that introduce each party.
    first_party_done = False
    second_party_done = False
    for paragraph in document.paragraphs:
        text = paragraph.text
        if not first_party_done and "“Company”" in text and "[COMPANY NAME]" in text:
            _set_paragraph_text(
                paragraph,
                _apply(
                    text,
                    {
                        "[COMPANY NAME]": intake.company_name,
                        "[JURISDICTION OF INCORPORATION]": intake.jurisdiction_of_incorporation,
                        "[REGISTERED OFFICE ADDRESS]": intake.registered_office,
                    },
                ),
            )
            first_party_done = True
        elif not second_party_done and "“Aspora”" in text and "[ASPORA ENTITY LEGAL NAME]" in text:
            _set_paragraph_text(
                paragraph,
                _apply(
                    text,
                    {
                        "[ASPORA ENTITY LEGAL NAME]": entity.legal_name,
                        "[JURISDICTION OF INCORPORATION]": entity.jurisdiction_of_incorporation,
                        "[REGISTERED OFFICE ADDRESS]": entity.registered_office,
                    },
                ),
            )
            second_party_done = True

    if not first_party_done:
        raise NdaGenerationError("Template is missing the Company (FIRST party) paragraph.")
    if not second_party_done:
        raise NdaGenerationError("Template is missing the Aspora (SECOND party) paragraph.")

    # Single-occurrence slots, filled across body paragraphs.
    global_fills = {
        "[•] day of [•], [YEAR]": f"{day} day of {month}, {year}",
        "[BUSINESS DESCRIPTION]": intake.business_description,
        "[GOVERNING LAW]": entity.governing_law_value,
    }
    for paragraph in document.paragraphs:
        text = paragraph.text
        if any(token in text for token in global_fills):
            _set_paragraph_text(paragraph, _apply(text, global_fills))

    # Purpose: the recital names the deal purpose. The template phrases the
    # purpose inline ("certain commercial propositions"); we make it concrete.
    for paragraph in document.paragraphs:
        if "certain commercial propositions" in paragraph.text and "“Purpose”" in paragraph.text:
            _set_paragraph_text(
                paragraph,
                paragraph.text.replace("certain commercial propositions", intake.purpose),
            )
            break

    # Signature block (the table): fill the party names + Aspora signatory.
    _fill_signature_table(document, entity, intake)

    manifest.slot_fills.update(
        {
            "[COMPANY NAME]": intake.company_name,
            "[ASPORA ENTITY LEGAL NAME]": entity.legal_name,
            "[JURISDICTION OF INCORPORATION] (first party)": intake.jurisdiction_of_incorporation,
            "[JURISDICTION OF INCORPORATION] (second party)": entity.jurisdiction_of_incorporation,
            "[REGISTERED OFFICE ADDRESS] (first party)": intake.registered_office,
            "[REGISTERED OFFICE ADDRESS] (second party)": entity.registered_office,
            "[BUSINESS DESCRIPTION]": intake.business_description,
            "[GOVERNING LAW]": entity.governing_law_value,
            # Empty -> the signature block renders a blank fill-line for signing.
            "[AUTHORISED SIGNATORY]": entity.signatory_name or "(blank fill-line)",
            "[DESIGNATION]": entity.signatory_title or "(blank fill-line)",
            "purpose": intake.purpose,
            "agreement_date": agreement_date.isoformat(),
        }
    )


def _fill_signature_table(document: DocxDocument, entity: EntityParty, intake: CounterpartyIntake) -> None:
    """Fill + normalise the signature block to the Playbook execution structure.

    The template's block uses "Designation:" and has no "By:"/"Date:" lines, so
    it fails the Playbook's signatures position (which requires both parties, a
    title, and a date — surfaced via ``By:``/``Title:``/``Date:`` markers). We
    rewrite each party cell to the Playbook ``redline_template`` shape while
    keeping the party-specific fills, so the executed block is Playbook-complete.
    """

    if not document.tables:
        raise NdaGenerationError("Template is missing the signature table.")

    table = document.tables[0]
    cells = table.rows[0].cells
    if len(cells) < 2:
        raise NdaGenerationError("Signature table does not have two party blocks.")

    # The two party signature boxes sit SIDE BY SIDE (counterparty left, Aspora
    # right). The template packs them as two adjacent columns with no gap between;
    # we insert a narrow EMPTY spacer column between them so the boxes are visibly
    # separated (an obvious gap, not cramped-adjacent). The counterparty box stays
    # the LEFT column and Aspora the RIGHT column.
    counterparty_cell, aspora_cell = _space_signature_cells_apart(table)

    # The LEFT cell is the counterparty (FIRST party), the RIGHT cell is Aspora
    # (SECOND party). Each gets its party's DocuSign anchor token so the envelope
    # can drop that signer's signHere/dateSigned tab on the right signature line.
    _write_signature_cell(
        counterparty_cell,
        party_name=intake.company_name,
        signatory_name="",
        signatory_title="",
        anchor=SIGNATURE_ANCHOR_COUNTERPARTY,
    )
    _write_signature_cell(
        aspora_cell,
        party_name=entity.legal_name,
        signatory_name=entity.signatory_name,
        signatory_title=entity.signatory_title,
        anchor=SIGNATURE_ANCHOR_ASPORA,
    )


# Width (in dxa / twentieths of a point; 1440 dxa = 1 inch) of the empty spacer
# column inserted between the two side-by-side signature boxes. The table spans
# the full text width, so a wide (~1.9in) middle spacer pushes the counterparty
# box flush to the LEFT margin and the Aspora box flush to the RIGHT margin with
# an obvious centred gap between them.
_SIGNATURE_SPACER_WIDTH_DXA = 2800


def _space_signature_cells_apart(table: Any) -> tuple[Any, Any]:
    """Insert a narrow empty spacer column between the two side-by-side boxes.

    The template's signature block is one row of two adjacent columns with no gap.
    We add a third grid column between them and give it an empty, narrow cell so
    the LEFT (counterparty) and RIGHT (Aspora) boxes are visibly separated. The two
    party boxes are re-narrowed to make room so the table still fits its original
    width. Returns ``(left_cell, right_cell)`` for the caller to fill.

    Works at the XML grid level (python-docx has no insert-column API). Falls back
    to the bare two-column layout if the grid is not the expected shape, so a
    template change can never make generation fail here.
    """

    from docx.oxml.ns import qn  # noqa: PLC0415

    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    grid_cols = grid.findall(qn("w:gridCol")) if grid is not None else []
    rows = list(tbl.findall(qn("w:tr")))
    # Expected template shape: exactly two grid columns and one row with two cells.
    first_row_cells = rows[0].findall(qn("w:tc")) if rows else []
    if grid is None or len(grid_cols) != 2 or len(rows) != 1 or len(first_row_cells) != 2:
        # Unexpected shape — leave the layout untouched, just hand back the two cells.
        cells = table.rows[0].cells
        return cells[0], cells[-1]

    # Re-budget the column widths: keep the total, carve the spacer out of the pair.
    total = sum(int(gc.get(qn("w:w")) or 0) for gc in grid_cols)
    spacer_w = min(_SIGNATURE_SPACER_WIDTH_DXA, max(0, total - 2))
    box_w = max(1, (total - spacer_w) // 2)
    grid_cols[0].set(qn("w:w"), str(box_w))
    grid_cols[1].set(qn("w:w"), str(box_w))

    # Insert a spacer grid column between the two existing columns.
    spacer_col = grid.makeelement(qn("w:gridCol"), {qn("w:w"): str(spacer_w)})
    grid_cols[0].addnext(spacer_col)

    # Insert a matching empty spacer cell in the row, between the two party cells.
    left_tc, right_tc = first_row_cells[0], first_row_cells[1]
    spacer_tc = _make_spacer_cell(left_tc, spacer_w)
    left_tc.addnext(spacer_tc)

    # Keep each party cell's own width marker in step with the re-budgeted grid
    # column (1 dxa = 1 twip; python-docx widths are EMU, 1 twip = 635 EMU).
    box_w_emu = box_w * 635
    _set_tc_width(left_tc, box_w)
    _set_tc_width(right_tc, box_w)

    # Re-fetch python-docx cell wrappers now the row has three cells.
    row_cells = table.rows[0].cells
    left_cell, right_cell = row_cells[0], row_cells[-1]
    left_cell.width = right_cell.width = box_w_emu
    return left_cell, right_cell


def _set_tc_width(tc: Any, width_dxa: int) -> None:
    """Set the ``<w:tcW>`` width (dxa) on a raw ``<w:tc>`` element."""

    from docx.oxml.ns import qn  # noqa: PLC0415

    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tc_pr)
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = tc_pr.makeelement(qn("w:tcW"), {})
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _make_spacer_cell(template_tc: Any, width_dxa: int) -> Any:
    """Build an empty ``<w:tc>`` spacer cell modelled on ``template_tc``.

    Carries a minimal ``<w:tcPr>`` with the spacer width and a single empty
    paragraph (a ``<w:tc>`` must contain at least one block-level element to be
    valid OOXML).
    """

    from docx.oxml.ns import qn  # noqa: PLC0415

    tc = template_tc.makeelement(qn("w:tc"), {})
    tc_pr = tc.makeelement(qn("w:tcPr"), {})
    tc_w = tc.makeelement(qn("w:tcW"), {qn("w:w"): str(width_dxa), qn("w:type"): "dxa"})
    tc_pr.append(tc_w)
    tc.append(tc_pr)
    tc.append(tc.makeelement(qn("w:p"), {}))
    return tc


def _write_signature_cell(
    cell: Any,
    *,
    party_name: str,
    signatory_name: str,
    signatory_title: str,
    anchor: str = "",
) -> None:
    """Rewrite a signature cell to: header, party name, By/Name/Title/Date lines.

    Mirrors the Playbook ``signatures.redline_template`` so the block carries the
    ``By:`` / ``Title:`` / ``Date:`` markers the signatures checker requires. The
    cell's existing paragraphs are reused in order and any surplus is cleared.

    ``anchor`` is this party's DocuSign anchor token. It is appended to the ``By:``
    signature line so DocuSign drops the signHere tab right on that line (the
    workflow offsets the tab off the marker). The marker is rendered hidden
    (white, 1pt) so it does not show in the executed document, but stays as
    searchable text so the anchor can always be found. Empty ``anchor`` leaves the
    block unchanged (e.g. a non-signing-flow render).
    """

    party_line = f"For {party_name}".strip() if party_name else "For _______________________"
    by_line = "By: _______________________________"
    lines = [
        party_line,
        by_line,
        f"Name: {signatory_name}".rstrip() if signatory_name else "Name: _______________________",
        f"Title: {signatory_title}".rstrip() if signatory_title else "Title: _______________________",
        "Date: _______________________",
    ]
    paragraphs = cell.paragraphs
    for index, line in enumerate(lines):
        if index < len(paragraphs):
            _set_paragraph_text(paragraphs[index], line)
        else:
            cell.add_paragraph(line)
    # Clear any leftover template paragraphs beyond the lines we wrote.
    for paragraph in paragraphs[len(lines):]:
        _set_paragraph_text(paragraph, "")

    # Plant the per-party anchor at the START of the By: line (index 1) as a hidden
    # run, so the marker is searchable text for DocuSign but invisible in the
    # rendered doc.
    #
    # Why the START and not the end: DocuSign places a tab with the LOWER-LEFT
    # corner of the tab at the LOWER-RIGHT corner of the anchor's bounding box, and
    # the (~1in-wide) signHere tab then grows RIGHTWARD across the page. The Aspora
    # signature box is flush against the page's RIGHT margin (see
    # ``_space_signature_cells_apart``), so an anchor at the END of that box's By:
    # line sits at the right margin and a rightward-growing tab runs off the right
    # page edge -> DocuSign rejects the whole envelope with INVALID_USER_OFFSET
    # (HTTP 400). Anchoring at the START of the line puts the tab origin at the
    # cell's LEFT edge, so the tab grows into the blank underscores with room to
    # spare in BOTH the left (counterparty) and right (Aspora) cells.
    if anchor:
        by_paragraph = cell.paragraphs[1]
        _prepend_hidden_anchor_run(by_paragraph, anchor)


def _prepend_hidden_anchor_run(paragraph: Any, anchor: str) -> None:
    """Prepend ``anchor`` to ``paragraph`` as a hidden (white, 1pt) leading run.

    The token must be present as searchable text (so DocuSign can locate it) but
    must not be visible in the executed document. It is planted as the FIRST run on
    the line so DocuSign places the signHere/dateSigned tab at the cell's LEFT edge
    (the tab grows rightward into the blank underscores, never off the page's right
    margin — see ``_write_signature_cell``). A trailing space keeps the hidden
    marker from butting against the visible "By:" text. Colour + tiny size make it
    effectively invisible without relying on the ``vanish`` property (which some
    renderers hide from text extraction, which would defeat anchoring).
    """

    from docx.shared import Pt, RGBColor  # noqa: PLC0415

    run = paragraph.add_run(anchor + " ")
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(1)
    # python-docx appends; move the new run's XML to the FRONT of the paragraph so
    # the marker is the line's first run (the tab anchors at the left edge).
    run_element = run._r
    paragraph._p.remove(run_element)
    _insert_run_first(paragraph, run_element)


def _insert_run_first(paragraph: Any, run_element: Any) -> None:
    """Insert ``run_element`` as the paragraph's first run-level child.

    A ``<w:p>`` may lead with non-run children (``<w:pPr>`` properties). The run
    must go AFTER any ``<w:pPr>`` but BEFORE the first existing ``<w:r>``, so the
    hidden anchor reads as the first text on the line.
    """

    from docx.oxml.ns import qn  # noqa: PLC0415

    p_element = paragraph._p
    first_run = p_element.find(qn("w:r"))
    if first_run is not None:
        first_run.addprevious(run_element)
    else:
        p_element.append(run_element)


# --------------------------------------------------------------------------- #
# Clause realignment to the Playbook
# --------------------------------------------------------------------------- #


def _mutuality_base_statement(playbook: Mapping[str, Any]) -> str:
    """The Playbook's deterministic mutuality statement (pre-adaptation)."""

    clause = _playbook_clause(playbook, "mutuality")
    return str(clause.get("redline_template") or "").strip() or (
        "Each party acts as both a Disclosing Party and a Receiving Party with respect to "
        "Confidential Information it discloses or receives, and the confidentiality obligations "
        "under this Agreement bind each party reciprocally."
    )


def _prefetch_clause_adaptations(
    playbook: Mapping[str, Any],
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
) -> dict[str, str]:
    """Adapt the three realigned clauses concurrently; ``{}`` with no adapter.

    Computes each clause's deterministic base text the same way the per-clause
    alignment does, then hands all three to the AI module's parallel runner so the
    network waits overlap. The returned ``{clause_id: adapted_text}`` is consumed by
    the alignment functions (which fall back to recomputing the base text for any
    clause absent from the map), so behaviour is identical to the old serial path —
    only the wall-clock changes. With no adapter we skip the work entirely.
    """

    if clause_adapter is None:
        return {}
    from .nda_generation_ai import adapt_clauses_in_parallel  # noqa: PLC0415

    context = _adapter_context(intake)
    jobs = [
        ("mutuality", _mutuality_base_statement(playbook), context),
        (CLAUSE_CONFIDENTIAL, _independent_development_base(playbook), context),
        (CLAUSE_TERM, _survival_base(playbook), context),
    ]
    return adapt_clauses_in_parallel(clause_adapter, jobs)


def _align_mutuality(
    document: DocxDocument,
    playbook: Mapping[str, Any],
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
    manifest: GenerationManifest,
    *,
    adapted_text: str | None = None,
) -> None:
    """Make the reciprocal-obligation language explicit per the Playbook.

    The template defines Disclosing/Receiving roles abstractly (recital), but
    leaves the reciprocity implicit, which the Playbook's mutuality position
    flags for review. We insert the Playbook's mutuality statement ("each party
    acts as both a Disclosing Party and a Receiving Party ...") right after the
    role-definition recital so both parties are bound symmetrically.

    ``adapted_text`` carries the pre-resolved (parallel-prefetched) adaptation when
    the engine adapted the clauses up front; absent, the clause is adapted inline
    (or left deterministic when there is no adapter), so this function is correct
    whether or not the prefetch ran.
    """

    statement = _mutuality_base_statement(playbook)
    if adapted_text is not None:
        statement = adapted_text.strip() or statement
    elif clause_adapter is not None:
        adapted = clause_adapter.adapt("mutuality", statement, _adapter_context(intake)).strip()
        statement = adapted or statement

    # Anchor: the recital that defines the Disclosing/Receiving roles.
    for paragraph in document.paragraphs:
        if "Disclosing Party" in paragraph.text and "Receiving Party" in paragraph.text and (
            "shall be referred to as" in paragraph.text
        ):
            _insert_paragraph_after(paragraph, statement)
            manifest.clause_alignments.append(
                "mutuality: inserted reciprocal-obligation statement (Playbook position)"
            )
            return
    raise NdaGenerationError("Template is missing the Disclosing/Receiving role-definition recital.")


def _align_confidential_information(
    document: DocxDocument,
    playbook: Mapping[str, Any],
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
    manifest: GenerationManifest,
    *,
    adapted_text: str | None = None,
) -> None:
    """Add the Playbook's missing exclusion so CI carve-outs match the standard.

    The template's "EXCEPTIONS TO CONFIDENTIAL INFORMATION" lists three carve-outs
    (public / lawful third-party / prior-possession) and omits the Playbook's
    "independently developed without use of Confidential Information" exclusion.
    We append it as a new sub-item to the exception list, in the template's own
    list style, so the definition stays template-shaped but Playbook-complete.

    ``adapted_text`` carries the pre-resolved (parallel-prefetched) adaptation; when
    absent the carve-out is adapted inline, so the function is correct with or
    without the prefetch.
    """

    if adapted_text is not None:
        independent = adapted_text.strip() or _independent_development_base(playbook)
    else:
        independent = _independent_development_sentence(playbook, clause_adapter, intake)

    # Find the last sub-item of the exceptions list (the prior-possession item).
    anchor_index = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().startswith("was previously in the possession of the Receiving Party"):
            anchor_index = index
            break
    if anchor_index is None:
        raise NdaGenerationError("Template is missing the CI exceptions list (prior-possession item).")

    anchor = document.paragraphs[anchor_index]
    # The prior-possession item ends "...written records." We turn its trailing
    # full stop into "; or" and add the independent-development item after it, in
    # the same list paragraph style.
    anchor_text = anchor.text.rstrip()
    if anchor_text.endswith("."):
        anchor_text = anchor_text[:-1] + "; or"
    _set_paragraph_text(anchor, anchor_text)

    _insert_paragraph_after(anchor, independent)

    manifest.clause_alignments.append(
        "confidential_information: added independent-development exclusion (Playbook standard carve-out)"
    )


def _align_term_and_survival(
    document: DocxDocument,
    playbook: Mapping[str, Any],
    term_years: int,
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
    manifest: GenerationManifest,
    *,
    adapted_text: str | None = None,
) -> None:
    """Rewrite the TERM clause to the Playbook term cap + survival carve-out.

    The template's TERM clause caps at "two (2) years" with no survival carve-out.
    We rewrite the clause body to a fixed term of ``term_years`` (already capped
    at the Playbook ``max_term_years``) and append the Playbook's trade-secret /
    legal / data-protection survival carve-out so it passes ``term_and_survival``.

    ``adapted_text`` carries the pre-resolved (parallel-prefetched) survival
    adaptation; when absent it is adapted inline, so the function is correct with
    or without the prefetch.
    """

    if adapted_text is not None:
        survival = adapted_text.strip() or _survival_base(playbook)
    else:
        survival = _survival_sentence(playbook, clause_adapter, intake)

    body = (
        "This Agreement shall become effective on the date of signing of this Agreement and shall "
        f"remain in force, and the confidentiality obligations shall survive, for a fixed period of "
        f"{_year_count_label(term_years, parenthetical=True)} from the date of this Agreement or until the completion of the "
        f"Purpose, whichever is later. {survival}"
    )

    for paragraph in document.paragraphs:
        if paragraph.text.startswith("TERM OF THE AGREEMENT:"):
            # Preserve the bold title run; replace only the body run(s).
            _set_clause_body(paragraph, "TERM OF THE AGREEMENT: ", body)
            manifest.clause_alignments.append(
                f"term_and_survival: term fixed at {term_years}y (<= Playbook max) + survival carve-out injected"
            )
            return
    raise NdaGenerationError("Template is missing the TERM OF THE AGREEMENT clause.")


# Blank-template fallbacks: used ONLY when the Playbook clause carries no template
# text. Normal operation sources the prose LIVE from the Playbook (so editing the
# Playbook changes the generated wording), exactly as ``_align_mutuality`` reads
# the mutuality ``redline_template`` and ``_resolve_term_years`` reads the term cap.
_INDEPENDENT_DEVELOPMENT_FALLBACK = (
    "is independently developed by the receiving Party without use of or reference "
    "to the Confidential Information."
)
_SURVIVAL_FALLBACK = (
    "Notwithstanding the foregoing, trade secrets, information whose confidentiality is "
    "required by law or regulation, and personal data protected by data-protection law shall "
    "remain confidential for as long as the protected status or applicable law requires."
)


def _independent_development_base(playbook: Mapping[str, Any]) -> str:
    """The CI independent-development carve-out, deterministic (pre-adaptation)."""

    clause = _playbook_clause(playbook, CLAUSE_CONFIDENTIAL)
    template = str(clause.get("standard_exclusions_template") or "").strip()
    return _independent_development_item_from_template(template) or _INDEPENDENT_DEVELOPMENT_FALLBACK


def _survival_base(playbook: Mapping[str, Any]) -> str:
    """The survival carve-out statement, deterministic (pre-adaptation)."""

    clause = _playbook_clause(playbook, CLAUSE_TERM)
    template = str(clause.get("redline_template") or "").strip()
    return _survival_statement_from_template(template) or _SURVIVAL_FALLBACK


def _independent_development_sentence(
    playbook: Mapping[str, Any],
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
) -> str:
    """The CI independent-development exclusion list item, sourced from the Playbook.

    Reads ``confidential_information.standard_exclusions_template`` and lifts its
    independent-development carve-out into the template's list-item grammar
    ("is ..."), so a Playbook edit to the standard exclusions flows through to the
    generated clause. Falls back to the literal only when the template is blank.
    """

    base = _independent_development_base(playbook)
    if clause_adapter is None:
        return base
    adapted = clause_adapter.adapt(CLAUSE_CONFIDENTIAL, base, _adapter_context(intake))
    return adapted.strip() or base


def _survival_sentence(
    playbook: Mapping[str, Any],
    clause_adapter: ClauseAdapter | None,
    intake: CounterpartyIntake,
) -> str:
    """The survival carve-out statement, sourced from the Playbook term template.

    Reads ``term_and_survival.redline_template`` and lifts its survival carve-out
    (the "except that ..." tail) into a standalone "Notwithstanding the foregoing,
    ..." statement appended after the fixed-term sentence, so it reads as its own
    statement rather than a sub-clause of the term sentence. A Playbook edit to the
    survival carve-out flows through; falls back to the literal only when blank.
    """

    base = _survival_base(playbook)
    if clause_adapter is None:
        return base
    adapted = clause_adapter.adapt(CLAUSE_TERM, base, _adapter_context(intake))
    return adapted.strip() or base


def _independent_development_item_from_template(template: str) -> str:
    """Lift the independent-development carve-out from the CI exclusions template.

    The ``standard_exclusions_template`` lists carve-outs inline
    ("...or independently developed without use of or reference to Confidential
    Information."). We take the final "or"-joined item and recast it into the
    template's list-item grammar ("is ..."). Returns "" when the template carries
    no independent-development carve-out, so the caller uses the literal fallback.
    """

    if not template:
        return ""
    body = template.rstrip().rstrip(".")
    # The carve-outs are joined by commas with a final "or"; take the last item.
    _, _, last = body.rpartition(", or ")
    fragment = (last or body).strip()
    if "independently develop" not in fragment.lower():
        return ""
    # Recast "independently developed ..." -> "is independently developed ..." so it
    # completes the list stem ("Confidential Information does not include ... that ...").
    if not fragment.lower().startswith("is "):
        fragment = "is " + fragment
    return fragment.rstrip(".") + "."


def _survival_statement_from_template(template: str) -> str:
    """Lift the survival carve-out from the term template into a standalone sentence.

    The ``redline_template`` reads "...survive for a fixed period of up to
    {label}, except that <carve-out>." We take the "except that ..." tail and
    front it with "Notwithstanding the foregoing, " so it stands alone after the
    generator's own fixed-term sentence. Returns "" when there is no carve-out tail.
    """

    if not template:
        return ""
    marker = "except that "
    lowered = template.lower()
    position = lowered.find(marker)
    if position < 0:
        return ""
    carve_out = template[position + len(marker):].strip().rstrip(".").strip()
    if not carve_out:
        return ""
    return "Notwithstanding the foregoing, " + carve_out + "."


def _adapter_context(intake: CounterpartyIntake) -> dict[str, Any]:
    return {
        "counterparty": intake.company_name,
        "purpose": intake.purpose,
        "nda_type": intake.nda_type,
    }


# --------------------------------------------------------------------------- #
# Playbook helpers
# --------------------------------------------------------------------------- #


def _playbook_clause(playbook: Mapping[str, Any], clause_id: str) -> Mapping[str, Any]:
    for clause in playbook.get("clauses", []):
        if clause.get("id") == clause_id:
            return clause
    raise NdaGenerationError(f"Playbook is missing the {clause_id!r} clause.")


def _approved_governing_law_options(playbook: Mapping[str, Any]) -> dict[str, str]:
    """Map approved governing-law option id -> the law value to write."""

    clause = _playbook_clause(playbook, CLAUSE_GOVERNING_LAW)
    options = (clause.get("rules") or {}).get("approved_options") or []
    resolved: dict[str, str] = {}
    for option in options:
        option_id = str(option.get("id") or "").strip()
        value = str(option.get("value") or option.get("label") or "").strip()
        if option_id and value:
            resolved[option_id] = value
    if not resolved:
        raise NdaGenerationError("Playbook governing_law clause has no approved_options.")
    return resolved


def _resolve_term_years(requested: int, playbook: Mapping[str, Any]) -> int:
    """Clamp the requested term to [1, Playbook max_term_years]."""

    clause = _playbook_clause(playbook, CLAUSE_TERM)
    max_years = int(clause.get("max_term_years") or 5)
    try:
        years = int(requested)
    except (TypeError, ValueError):
        years = 2
    if years < 1:
        years = 1
    if years > max_years:
        years = max_years
    return years


# --------------------------------------------------------------------------- #
# docx low-level helpers
# --------------------------------------------------------------------------- #


def _load_template(template_path: Path | str) -> DocxDocument:
    path = Path(template_path)
    if not path.exists():
        raise NdaGenerationError(f"Template not found at {path}.")
    return Document(str(path))


def _document_bytes(document: DocxDocument) -> bytes:
    with BytesIO() as output:
        document.save(output)
        return output.getvalue()


def _apply(text: str, fills: Mapping[str, str]) -> str:
    for token, value in fills.items():
        text = text.replace(token, value)
    return text


def _set_paragraph_text(paragraph: Any, text: str) -> None:
    """Replace a paragraph's text, collapsing it to a single run.

    Placeholders can span runs, so we rewrite the whole paragraph text into the
    first run and clear the rest. Run-level formatting beyond the first run is
    not preserved — acceptable here because the affected paragraphs are plain
    body text or single-format party lines.
    """

    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _set_clause_body(paragraph: Any, title_prefix: str, body: str) -> None:
    """Replace a titled clause's body while keeping the bold title run intact.

    Body clauses are ``<bold title run><body run>``. We keep run0 (the title)
    and rewrite the remainder to ``title_prefix``-stripped ``body``.
    """

    runs = paragraph.runs
    if runs and runs[0].text.strip().rstrip(":").upper() == title_prefix.strip().rstrip(": ").upper():
        runs[0].text = title_prefix
        if len(runs) > 1:
            runs[1].text = body
            for run in runs[2:]:
                run.text = ""
        else:
            paragraph.add_run(body)
    else:
        _set_paragraph_text(paragraph, title_prefix + body)


def _insert_paragraph_after(paragraph: Any, text: str) -> Any:
    """Insert a new paragraph immediately after ``paragraph`` (same parent).

    The new ``<w:p>`` is a run-stripped copy of ``paragraph`` so it inherits the
    source's paragraph properties (list numbering + style + indentation) — the
    copied ``pPr`` already carries them, so nothing is re-applied.
    """

    from docx.text.paragraph import Paragraph as _P

    new_p = _copy_blank_p(paragraph)
    paragraph._p.addnext(new_p)
    new_paragraph = _P(new_p, paragraph._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def _copy_blank_p(paragraph: Any):
    """A fresh empty ``<w:p>`` element to host an inserted paragraph."""

    from docx.oxml.ns import qn
    import copy

    new_p = copy.deepcopy(paragraph._p)
    # Strip all runs from the copy; keep paragraph properties (numbering/style).
    for run in new_p.findall(qn("w:r")):
        new_p.remove(run)
    return new_p


# The template's own variable slots — every one of these must be filled.
_TEMPLATE_SLOTS = frozenset(
    {
        "[COMPANY NAME]",
        "[ASPORA ENTITY LEGAL NAME]",
        "[JURISDICTION OF INCORPORATION]",
        "[REGISTERED OFFICE ADDRESS]",
        "[BUSINESS DESCRIPTION]",
        "[GOVERNING LAW]",
        "[AUTHORISED SIGNATORY]",
        "[DESIGNATION]",
        "[YEAR]",
        "[•]",
    }
)


def _assert_no_unfilled_placeholders(document: DocxDocument, allowed: set[str]) -> None:
    """Fail closed if any *template* slot survived generation.

    Only the template's own ``_TEMPLATE_SLOTS`` count as "unfilled". An entity
    may legitimately carry a bracketed signatory value (the registry ships
    ``"[Authorised Signatory]"`` / ``"[Title]"`` when no signatory is assigned);
    those were filled *from input* and are passed in ``allowed``, so they don't
    trip the guard.
    """

    import re

    leftover: set[str] = set()
    pattern = re.compile(r"\[[^\]]*\]")
    texts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    for text in texts:
        for token in pattern.findall(text):
            if token in _TEMPLATE_SLOTS and token not in allowed:
                leftover.add(token)
    if leftover:
        raise NdaGenerationError(
            "Generated NDA still contains unfilled placeholders: " + ", ".join(sorted(leftover))
        )


# --------------------------------------------------------------------------- #
# Small text utilities
# --------------------------------------------------------------------------- #

def _split_date(date: _dt.date) -> tuple[str, str, str]:
    day = _ordinal(date.day)
    month = date.strftime("%B")
    return day, month, str(date.year)


def _ordinal(day: int) -> str:
    if 11 <= (day % 100) <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{day}{suffix}"


def _default_address_entry(bundle: Mapping[str, Any]) -> Mapping[str, Any] | None:
    """The default-flagged address entry (first address if none is flagged)."""

    addresses = bundle.get("addresses") or []
    for address in addresses:
        if address.get("default"):
            return address
    if addresses:
        return addresses[0]
    return None


def _format_address(entry: Mapping[str, Any]) -> str:
    lines = entry.get("lines") or []
    joined = ", ".join(line for line in lines if line)
    return joined or str(entry.get("label") or "").strip()


def select_bundle_address(bundle: Mapping[str, Any], address_id: str = "") -> tuple[str, str]:
    """Pick the entity address to write, honouring the user's chosen ``address_id``.

    Returns ``(formatted_address, chosen_address_id)``. When ``address_id`` names a
    real address on the bundle, THAT address is used (the product lets the user pick
    a non-default office, e.g. Real Transfer's Belfast registered office instead of
    the default London corporate office). An UNKNOWN/invalid id is rejected — mirrors
    the governing-law override guard, so a stale or tampered client value fails loudly
    rather than silently substituting the default. A blank id falls back to the
    default-flagged address. The returned id is recorded on the manifest for the same
    provenance parity the governing-law override has.
    """

    addresses = bundle.get("addresses") or []
    wanted = str(address_id or "").strip()
    if wanted:
        for address in addresses:
            if str(address.get("id") or "").strip() == wanted:
                return _format_address(address), wanted
        known = sorted(str(a.get("id") or "").strip() for a in addresses if a.get("id"))
        raise NdaGenerationError(
            f"Address id {wanted!r} is not an address on entity bundle "
            f"{str(bundle.get('id') or '')!r} (known: {known})."
        )
    chosen = _default_address_entry(bundle)
    if chosen is None:
        raise NdaGenerationError("Entity bundle has no addresses.")
    return _format_address(chosen), str(chosen.get("id") or "").strip()
