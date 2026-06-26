from io import BytesIO
import json
import posixpath
from pathlib import Path
import unittest
from unittest.mock import patch
import warnings
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from nda_automation.checker import review_nda
from nda_automation.docx_export import (
    A4_PAGE_HEIGHT_TWIPS,
    A4_PAGE_WIDTH_TWIPS,
    DocxExportError,
    SupplementalRedlineUnavailableError,
    accept_all_revisions,
    build_review_report_docx,
    build_source_redline_docx,
    validate_docx_open_health,
)
from nda_automation.redline_xml import (
    _apply_tracked_run_format,
    _needs_inline_space,
    _strip_paragraph_property_revisions,
    _tracked_replace_paragraph,
)
from nda_automation import docx_export, docx_health, export_service, source_redline_docx
from nda_automation.docx_health import verify_export_content_coverage
from nda_automation.inline_diff import diff_text_operations
from nda_automation import docx_text
from nda_automation.docx_text import extract_docx_paragraphs
from nda_automation.review_document import align_document_paragraphs
from nda_automation.redline_actions import (
    REDLINE_DELETE_PARAGRAPH,
    REDLINE_FORMAT_PARAGRAPH,
    REDLINE_INSERT_AFTER_PARAGRAPH,
    REDLINE_REPLACE_PARAGRAPH,
)
from tests.docx_redline_contract import (
    assert_docx_redline_contract,
    inspect_docx_redline_contract,
    paragraph_property_revisions,
)

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W14_NS = {"w14": "http://schemas.microsoft.com/office/word/2010/wordml"}
W15_NS = {"w15": "http://schemas.microsoft.com/office/word/2012/wordml"}
REL_NS = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
STYLE_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"
SETTINGS_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings"
COMMENTS_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
OFFICE_DOCUMENT_RELATIONSHIP_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
DOCUMENT_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
RELATIONSHIPS_CONTENT_TYPE = "application/vnd.openxmlformats-package.relationships+xml"
SETTINGS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_EXTENDED_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"
COMMENTS_EXTENDED_RELATIONSHIP_TYPE = "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
STYLES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"
INLINE_DIFF_VECTORS_PATH = Path(__file__).parent / "fixtures" / "inline_diff_vectors.json"
# Generated from inline_diff_vectors.source.json; use generate_inline_diff_vectors.mjs.
SOURCE_EXPORT_REPORT_LEAKAGE_PHRASES = [
    "NDA Redline",
    "Review Notes",
    "Clause Findings",
    "Proposed Redline",
    "Overall status:",
    "Requirements passed:",
    "Requirements needing review:",
    "Requirements failed:",
    "Checked at:",
    "The Redlined NDA section contains native Word tracked changes.",
    "source paragraph",
]


def mark_ai_executed(result):
    """Stamp the reliable "AI engine ran" marker on a (deterministically built)
    review result so the report renderer surfaces verdicts/findings/comments.

    ``build_review_report_docx`` gates those sections on
    ``review_state.review_was_ai_executed`` (P0: never ship deterministic verdicts
    to a counterparty). Tests that exercise the verdict/finding/comment rendering
    must therefore present an AI-executed result; ``review_nda`` alone is the
    deterministic engine and is correctly treated as "AI has not run".
    """
    result["active_review_engine"] = {"executed_engine": "ai_first"}
    return result


def docx_xml_roots(docx_bytes):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        assert archive.testzip() is None
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return ET.fromstring(settings_xml), ET.fromstring(document_xml), document_xml


def assert_source_export_has_no_report_leakage(testcase, docx_bytes, extra_forbidden=()):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        testcase.assertIsNone(archive.testzip())
        document_xml = archive.read("word/document.xml").decode("utf-8")
    for phrase in [*SOURCE_EXPORT_REPORT_LEAKAGE_PHRASES, *extra_forbidden]:
        testcase.assertNotIn(phrase, document_xml)


def docx_document_relationship_targets(docx_bytes):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        relationships_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    relationships_root = ET.fromstring(relationships_xml)
    return {
        relationship.attrib["Type"]: relationship.attrib["Target"]
        for relationship in relationships_root.findall(".//rel:Relationship", REL_NS)
    }


def docx_content_type_overrides(docx_bytes):
    content_type_ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    with ZipFile(BytesIO(docx_bytes)) as archive:
        content_types_xml = archive.read("[Content_Types].xml").decode("utf-8")
    content_types_root = ET.fromstring(content_types_xml)
    return {
        override.attrib["PartName"]: override.attrib["ContentType"]
        for override in content_types_root.findall(".//ct:Override", content_type_ns)
    }


def docx_content_types(docx_bytes):
    content_type_ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    with ZipFile(BytesIO(docx_bytes)) as archive:
        content_types_xml = archive.read("[Content_Types].xml").decode("utf-8")
    content_types_root = ET.fromstring(content_types_xml)
    defaults = {
        default.attrib["Extension"]: default.attrib["ContentType"]
        for default in content_types_root.findall(".//ct:Default", content_type_ns)
    }
    overrides = {
        override.attrib["PartName"]: override.attrib["ContentType"]
        for override in content_types_root.findall(".//ct:Override", content_type_ns)
    }
    return defaults, overrides


def docx_comments(docx_bytes):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        comments_xml = archive.read("word/comments.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    return ET.fromstring(comments_xml), ET.fromstring(document_xml), comments_xml, document_xml


def docx_comments_extended(docx_bytes):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        comments_extended_xml = archive.read("word/commentsExtended.xml").decode("utf-8")
    return ET.fromstring(comments_extended_xml), comments_extended_xml


def assert_every_xml_part_parses(testcase, docx_bytes):
    with ZipFile(BytesIO(docx_bytes)) as archive:
        testcase.assertIsNone(archive.testzip())
        for name in archive.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                ET.fromstring(archive.read(name))


def relationship_targets(archive, relationship_part):
    relationships_root = ET.fromstring(archive.read(relationship_part).decode("utf-8"))
    return [
        relationship.attrib
        for relationship in relationships_root.findall(".//rel:Relationship", REL_NS)
    ]


def resolve_relationship_target(relationship_part, target):
    if target.startswith("/"):
        return target.removeprefix("/")
    if relationship_part == "_rels/.rels":
        base_dir = ""
    else:
        rels_dir = posixpath.dirname(relationship_part)
        base_dir = posixpath.dirname(rels_dir)
    return posixpath.normpath(posixpath.join(base_dir, target))


def _delete_and_insert_review_result(paragraphs):
    """Build a review result with one delete + one insert redline edit.

    Decouples the source-export delete+insert mechanism test from the deterministic
    engine: it deletes the paragraph mentioning circumvention and inserts a
    governing-law clause after the first paragraph.
    """
    paragraph_records = [
        {
            "id": f"p{index}",
            "index": index,
            "source_index": paragraph.get("source_index", index),
            "text": str(paragraph["text"]),
        }
        for index, paragraph in enumerate(paragraphs, start=1)
    ]
    delete_paragraph = next(p for p in paragraph_records if "circumvent" in p["text"])
    anchor_paragraph = paragraph_records[0]
    insert_text = "This Agreement shall be governed by the laws of England and Wales."
    return {
        "paragraphs": paragraph_records,
        "clauses": [],
        "redline_edits": [
            {
                "id": "r1",
                "clause_id": "non_circumvention",
                "clause_name": "Non-Circumvention",
                "paragraph_id": delete_paragraph["id"],
                "paragraph_index": delete_paragraph["index"],
                "source_index": delete_paragraph["source_index"],
                "action": REDLINE_DELETE_PARAGRAPH,
                "original_text": delete_paragraph["text"],
                "replacement_text": "",
            },
            {
                "id": "r2",
                "clause_id": "governing_law",
                "clause_name": "Governing Law",
                "paragraph_id": anchor_paragraph["id"],
                "paragraph_index": anchor_paragraph["index"],
                "source_index": anchor_paragraph["source_index"],
                "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                "original_text": "",
                "replacement_text": insert_text,
                "anchor_text": anchor_paragraph["text"],
                "insert_text": insert_text,
            },
        ],
    }


def assert_docx_package_healthy(testcase, docx_bytes, require_styles=False):
    testcase.assertEqual(validate_docx_open_health(docx_bytes, require_styles=require_styles), [])
    required_parts = {
        "[Content_Types].xml",
        "_rels/.rels",
        "word/document.xml",
        "word/_rels/document.xml.rels",
        "word/settings.xml",
    }
    if require_styles:
        required_parts.add("word/styles.xml")

    with ZipFile(BytesIO(docx_bytes)) as archive:
        testcase.assertIsNone(archive.testzip())
        names = set(archive.namelist())
        testcase.assertTrue(required_parts <= names, f"Missing DOCX parts: {sorted(required_parts - names)}")

        defaults, overrides = docx_content_types(docx_bytes)
        testcase.assertEqual(defaults.get("rels"), RELATIONSHIPS_CONTENT_TYPE)
        testcase.assertEqual(defaults.get("xml"), "application/xml")
        testcase.assertEqual(overrides.get("/word/document.xml"), DOCUMENT_CONTENT_TYPE)
        testcase.assertEqual(overrides.get("/word/settings.xml"), SETTINGS_CONTENT_TYPE)
        if "word/styles.xml" in names:
            testcase.assertEqual(overrides.get("/word/styles.xml"), STYLES_CONTENT_TYPE)

        package_relationships = relationship_targets(archive, "_rels/.rels")
        office_document_targets = [
            resolve_relationship_target("_rels/.rels", relationship["Target"])
            for relationship in package_relationships
            if relationship.get("Type") == OFFICE_DOCUMENT_RELATIONSHIP_TYPE
        ]
        testcase.assertEqual(office_document_targets, ["word/document.xml"])

        document_relationships = relationship_targets(archive, "word/_rels/document.xml.rels")
        relationship_targets_by_type = {
            relationship["Type"]: resolve_relationship_target("word/_rels/document.xml.rels", relationship["Target"])
            for relationship in document_relationships
            if relationship.get("TargetMode") != "External"
        }
        for target in relationship_targets_by_type.values():
            testcase.assertIn(target, names)
        testcase.assertEqual(relationship_targets_by_type.get(SETTINGS_RELATIONSHIP_TYPE), "word/settings.xml")
        if require_styles:
            testcase.assertEqual(relationship_targets_by_type.get(STYLE_RELATIONSHIP_TYPE), "word/styles.xml")
        elif STYLE_RELATIONSHIP_TYPE in relationship_targets_by_type:
            testcase.assertEqual(relationship_targets_by_type[STYLE_RELATIONSHIP_TYPE], "word/styles.xml")


def tracked_deleted_text(document_root):
    return [
        "".join(node.text or "" for node in deletion.findall(".//w:delText", W_NS))
        for deletion in document_root.findall(".//w:del", W_NS)
    ]


def tracked_inserted_text(document_root):
    return [
        "".join(node.text or "" for node in insertion.findall(".//w:t", W_NS))
        for insertion in document_root.findall(".//w:ins", W_NS)
    ]


def paragraph_mark_revisions(document_root, kind):
    return document_root.findall(f".//w:pPr/w:rPr/w:{kind}", W_NS)


def revision_text_for_state(node, accepted):
    tag = node.tag.rsplit("}", 1)[-1]
    if tag == "del":
        return "" if accepted else "".join(revision_text_for_state(child, accepted=False) for child in list(node))
    if tag == "ins":
        return "".join(revision_text_for_state(child, accepted=True) for child in list(node)) if accepted else ""
    if tag in {"t", "delText"}:
        return node.text or ""
    if tag == "br":
        return "\n"
    return "".join(revision_text_for_state(child, accepted) for child in list(node))


def make_source_docx(paragraphs, include_package_rels=True):
    body = "".join(
        f"<w:p><w:r><w:t>{escape_xml(paragraph)}</w:t></w:r></w:p>"
        for paragraph in paragraphs
    )
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{body}</w:body>
</w:document>"""
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    package_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    document_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""
    with BytesIO() as output:
        with ZipFile(output, "w", ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types_xml)
            if include_package_rels:
                archive.writestr("_rels/.rels", package_rels_xml)
            archive.writestr("word/document.xml", document_xml)
            archive.writestr("word/_rels/document.xml.rels", document_rels_xml)
            archive.writestr("customXml/item1.xml", "<custom>preserved</custom>")
        return output.getvalue()


def make_source_docx_with_footer(body_paragraphs, footer_paragraphs):
    """A native-DOCX source whose body holds ``body_paragraphs`` and whose
    ``word/footer1.xml`` part holds ``footer_paragraphs`` (e.g. a company-letterhead
    footer). ``extract_docx_paragraphs`` tags the footer paragraphs
    ``source_kind="supplemental"`` and appends them after the body, exactly like the
    real Moorwand letterhead NDA whose footer tripped the content-coverage gate."""
    footer_body = "".join(
        f"<w:p><w:r><w:t>{escape_xml(paragraph)}</w:t></w:r></w:p>"
        for paragraph in footer_paragraphs
    )
    footer_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"{footer_body}</w:ftr>"
    )
    return replace_docx_parts(
        make_source_docx(body_paragraphs),
        {"word/footer1.xml": footer_xml},
    )


def make_source_docx_with_internal_blank_line_paragraph(prefix_paragraphs, internal_blocks):
    """A source DOCX whose final paragraph is ONE physical <w:p> holding several
    logical blocks separated by a hard blank line (two <w:br/>), as the real
    extractor produces. The extractor reads it as one paragraph whose text has an
    internal "\\n\\n"; align_document_paragraphs re-splits it into one review
    paragraph per block, all sharing that physical paragraph's source_index."""
    runs = []
    for index, block in enumerate(internal_blocks):
        if index:
            runs.append("<w:br/><w:br/>")
        runs.append(f"<w:t>{escape_xml(block)}</w:t>")
    blank_line_paragraph = f"<w:p><w:r>{''.join(runs)}</w:r></w:p>"
    body = "".join(
        f"<w:p><w:r><w:t>{escape_xml(paragraph)}</w:t></w:r></w:p>"
        for paragraph in prefix_paragraphs
    ) + blank_line_paragraph
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{body}</w:body>
</w:document>"""
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    package_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    document_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""
    with BytesIO() as output:
        with ZipFile(output, "w", ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", content_types_xml)
            archive.writestr("_rels/.rels", package_rels_xml)
            archive.writestr("word/document.xml", document_xml)
            archive.writestr("word/_rels/document.xml.rels", document_rels_xml)
        return output.getvalue()


def make_source_docx_with_table(before_paragraphs, table_cells, after_paragraphs):
    before = "".join(
        f"<w:p><w:r><w:t>{escape_xml(paragraph)}</w:t></w:r></w:p>"
        for paragraph in before_paragraphs
    )
    rows = "".join(
        "<w:tr>"
        + "".join(
            f"<w:tc><w:p><w:r><w:t>{escape_xml(cell)}</w:t></w:r></w:p></w:tc>"
            for cell in row
        )
        + "</w:tr>"
        for row in table_cells
    )
    table = f"<w:tbl>{rows}</w:tbl>"
    after = "".join(
        f"<w:p><w:r><w:t>{escape_xml(paragraph)}</w:t></w:r></w:p>"
        for paragraph in after_paragraphs
    )
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>{before}{table}{after}</w:body>
</w:document>"""
    return replace_docx_parts(
        make_source_docx([*before_paragraphs, *[cell for row in table_cells for cell in row], *after_paragraphs]),
        {"word/document.xml": document_xml},
    )


def replace_docx_parts(docx_bytes, replacements):
    with ZipFile(BytesIO(docx_bytes), "r") as source_archive:
        with BytesIO() as output:
            with ZipFile(output, "w", ZIP_DEFLATED) as patched_archive:
                written = set()
                for item in source_archive.infolist():
                    data = replacements.get(item.filename, source_archive.read(item.filename))
                    if isinstance(data, str):
                        data = data.encode("utf-8")
                    patched_archive.writestr(item, data)
                    written.add(item.filename)
                for name, data in replacements.items():
                    if name in written:
                        continue
                    if isinstance(data, str):
                        data = data.encode("utf-8")
                    patched_archive.writestr(name, data)
            return output.getvalue()


def duplicate_docx_part(docx_bytes, part_name):
    with ZipFile(BytesIO(docx_bytes), "r") as source_archive:
        with BytesIO() as output:
            with ZipFile(output, "w", ZIP_DEFLATED) as patched_archive:
                for item in source_archive.infolist():
                    data = source_archive.read(item.filename)
                    patched_archive.writestr(item, data)
                    if item.filename == part_name:
                        with warnings.catch_warnings():
                            warnings.simplefilter("ignore", UserWarning)
                            patched_archive.writestr(item, data)
            return output.getvalue()


def unsafe_document_xml():
    return """<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE w:document [
  <!ENTITY a "aaaaaaaaaa">
  <!ENTITY b "&a;&a;&a;&a;&a;&a;&a;&a;&a;&a;">
]>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body><w:p><w:r><w:t>&b;</w:t></w:r></w:p></w:body>
</w:document>"""


def escape_xml(value):
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def track_changes_contract_review_result():
    paragraphs = [
        {
            "id": "p1",
            "index": 1,
            "source_index": 1,
            "text": "This Agreement shall be governed by the laws of California.",
        },
        {
            "id": "p2",
            "index": 2,
            "source_index": 2,
            "text": "Insert after this paragraph.",
        },
        {
            "id": "p3",
            "index": 3,
            "source_index": 3,
            "text": "The Recipient must not circumvent the Company.",
        },
    ]
    redline_edits = [
        {
            "id": "r1",
            "action": REDLINE_REPLACE_PARAGRAPH,
            "paragraph_id": "p1",
            "source_index": 1,
            "original_text": paragraphs[0]["text"],
            "replacement_text": "This Agreement shall be governed by the laws of England and Wales.",
        },
        {
            "id": "r2",
            "action": REDLINE_INSERT_AFTER_PARAGRAPH,
            "paragraph_id": "p2",
            "source_index": 2,
            "insert_text": "New required clause.",
        },
        {
            "id": "r3",
            "action": REDLINE_DELETE_PARAGRAPH,
            "paragraph_id": "p3",
            "source_index": 3,
            "original_text": paragraphs[2]["text"],
        },
    ]
    return {
        "overall_status": "does_not_meet_requirements",
        "requirements_passed": 0,
        "requirements_failed": 3,
        "checked_at": "2026-05-31T00:00:00+00:00",
        "paragraphs": paragraphs,
        "clauses": [],
        "redline_edits": redline_edits,
    }


def assert_track_changes_contract(testcase, docx_bytes, redline_edits):
    return assert_docx_redline_contract(testcase, docx_bytes, redline_edits)


def inline_diff_vectors():
    with INLINE_DIFF_VECTORS_PATH.open(encoding="utf-8") as fixture:
        return json.load(fixture)


def inline_diff_operation_vectors():
    return [
        vector
        for vector in inline_diff_vectors()
        if "original" in vector and "replacement" in vector and "operations" in vector
    ]


def inline_diff_spacing_pairs():
    return [
        pair
        for vector in inline_diff_vectors()
        for pair in vector.get("spacing_pairs", [])
    ]


def expand_inline_diff_vector(vector):
    original = vector["original"]
    replacement = vector["replacement"]
    operations = [(operation["type"], operation["token"]) for operation in vector["operations"]]
    return original, replacement, operations


class DocxExportTests(unittest.TestCase):
    def test_inline_diff_operations_match_shared_vectors(self):
        for vector in inline_diff_operation_vectors():
            with self.subTest(vector["name"]):
                original, replacement, expected_operations = expand_inline_diff_vector(vector)
                self.assertEqual(diff_text_operations(original, replacement), expected_operations)

    def test_inline_spacing_matches_shared_vectors(self):
        for pair in inline_diff_spacing_pairs():
            with self.subTest(f"{pair['previous_token']} + {pair['token']}"):
                self.assertEqual(
                    _needs_inline_space(pair["previous_token"], pair["token"]),
                    pair["needs_space"],
                )

    def test_tracked_replace_paragraph_reconstructs_shared_vectors(self):
        for vector in inline_diff_operation_vectors():
            with self.subTest(vector["name"]):
                paragraph_xml, _next_revision_id = _tracked_replace_paragraph(
                    vector["original"],
                    vector["replacement"],
                    7,
                )

                root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
                paragraph = root.find(".//w:p", W_NS)
                self.assertEqual(revision_text_for_state(paragraph, accepted=False), vector["original"])
                self.assertEqual(revision_text_for_state(paragraph, accepted=True), vector["replacement"])

    def test_tracked_replace_paragraph_preserves_punctuation_spacing(self):
        original = "This Agreement (California) applies."
        replacement = "This Agreement (England and Wales) applies."

        paragraph_xml, next_revision_id = _tracked_replace_paragraph(original, replacement, 7)

        root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
        paragraph = root.find(".//w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        self.assertEqual(next_revision_id, 9)

    def test_tracked_replace_paragraph_preserves_grouped_numbers_and_non_ascii_words(self):
        original = "Payment cap is 1,000 for café records."
        replacement = "Payment cap is 1,000 for café documents."

        paragraph_xml, next_revision_id = _tracked_replace_paragraph(original, replacement, 7)

        root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
        paragraph = root.find(".//w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        self.assertEqual(next_revision_id, 9)

    def test_tracked_replace_paragraph_preserves_currency_amount_spacing(self):
        original = "Payment cap is $100 for records."
        replacement = "Payment cap is $100 for documents."

        paragraph_xml, next_revision_id = _tracked_replace_paragraph(original, replacement, 7)

        root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
        paragraph = root.find(".//w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        self.assertEqual(next_revision_id, 9)

    def test_tracked_replace_paragraph_preserves_spaced_numeric_list_items(self):
        original = "Payment caps are 1, 2, 3, 400 for different classes."
        replacement = "Payment caps are 1, 2, 3, 400 for different categories."

        paragraph_xml, next_revision_id = _tracked_replace_paragraph(original, replacement, 7)

        root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
        paragraph = root.find(".//w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        self.assertEqual(next_revision_id, 9)

    def test_tracked_replace_paragraph_preserves_newlines(self):
        original = "Line one\nLine two"
        replacement = "Line one\nLine three"

        paragraph_xml, next_revision_id = _tracked_replace_paragraph(original, replacement, 7)

        root = ET.fromstring(f'<root xmlns:w="{W_NS["w"]}">{paragraph_xml}</root>')
        paragraph = root.find(".//w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        self.assertEqual(next_revision_id, 9)

    def test_review_report_docx_opens_with_track_changes_enabled(self):
        result = mark_ai_executed(review_nda("This Agreement shall be governed by the laws of California."))

        docx_bytes = build_review_report_docx(result, title="California NDA")

        assert_docx_package_healthy(self, docx_bytes, require_styles=True)
        settings_root, document_root, document_xml = docx_xml_roots(docx_bytes)
        relationship_targets = docx_document_relationship_targets(docx_bytes)

        self.assertEqual(relationship_targets[STYLE_RELATIONSHIP_TYPE], "styles.xml")
        self.assertEqual(relationship_targets[SETTINGS_RELATIONSHIP_TYPE], "settings.xml")
        self.assertIsNotNone(settings_root.find(".//w:trackRevisions", W_NS))
        self.assertIsNotNone(settings_root.find(".//w:revisionView", W_NS))
        self.assertIn("NDA Redline", document_xml)
        self.assertIn("The Redlined NDA section contains native Word tracked changes.", document_xml)
        self.assertIn("Governing Law - CHECK", document_xml)
        self.assertIn("Template options", document_xml)
        self.assertIn("This Agreement shall be governed by the laws of the DIFC.", document_xml)
        deletions = document_root.findall(".//w:del", W_NS)
        insertions = document_root.findall(".//w:ins", W_NS)
        deleted_text = tracked_deleted_text(document_root)
        inserted_text = tracked_inserted_text(document_root)
        self.assertGreaterEqual(len(deletions), 1)
        self.assertGreaterEqual(len(insertions), 1)
        self.assertTrue(any("California" in text for text in deleted_text))
        self.assertFalse(any("This Agreement shall be governed by the laws of California." in text for text in deleted_text))
        self.assertTrue(any("England and Wales" in text for text in inserted_text))

    def test_review_report_docx_omits_verdicts_when_ai_review_has_not_run(self):
        # P0 gate: this report can ship to the counterparty. A deterministic
        # review (review_nda alone -> review_was_ai_executed is False) must NOT
        # present verdicts/counts/per-clause findings as authoritative review
        # notes. The Redlined NDA content / tracked edits stay intact.
        result = review_nda("This Agreement shall be governed by the laws of California.")
        result["review_comments"] = [
            {
                "author": "Reviewer",
                "clause_id": "governing_law",
                "clause_name": "Governing Law",
                "paragraph_id": "p1",
                "text": "Deterministic note that must not ship.",
            }
        ]

        docx_bytes = build_review_report_docx(result, title="California NDA")

        assert_docx_package_healthy(self, docx_bytes, require_styles=True)
        settings_root, document_root, document_xml = docx_xml_roots(docx_bytes)
        with ZipFile(BytesIO(docx_bytes)) as archive:
            package_names = set(archive.namelist())

        # The "AI has not run" notice replaces the verdict body.
        self.assertIn("AI review has not been run on this document.", document_xml)
        # No verdict / count / per-clause-finding lines.
        self.assertNotIn("Overall status:", document_xml)
        self.assertNotIn("Requirements passed:", document_xml)
        self.assertNotIn("Requirements needing review:", document_xml)
        self.assertNotIn("Requirements failed:", document_xml)
        self.assertNotIn("Clause Findings", document_xml)
        self.assertNotIn("Governing Law - CHECK", document_xml)
        self.assertNotIn("Template options", document_xml)
        # Deterministic reviewer comments are not anchored either: no comments
        # part is emitted and the body carries no comment references.
        self.assertNotIn("word/comments.xml", package_names)
        self.assertEqual(document_root.findall(".//w:commentReference", W_NS), [])
        self.assertNotIn("Deterministic note that must not ship.", document_xml)
        # The actual NDA content / Redlined NDA section is still present.
        self.assertIsNotNone(settings_root.find(".//w:trackRevisions", W_NS))
        self.assertIn("NDA Redline", document_xml)
        self.assertIn("Redlined NDA", document_xml)
        self.assertTrue(any("California" in text for text in tracked_deleted_text(document_root)))

    def test_review_report_docx_strips_invalid_xml_characters(self):
        result = review_nda("This Agreement shall be governed by the laws of California.\x08\ud800\ufdd0\U0001fffe")

        docx_bytes = build_review_report_docx(result, title="California\x08\udfff\ufdef\U0010ffffNDA")

        assert_docx_package_healthy(self, docx_bytes, require_styles=True)
        with ZipFile(BytesIO(docx_bytes)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            core_xml = archive.read("docProps/core.xml").decode("utf-8")
        ET.fromstring(document_xml)
        ET.fromstring(core_xml)
        self.assertNotIn("\x08", document_xml)
        self.assertNotIn("\x08", core_xml)
        self.assertNotIn("\ud800", document_xml)
        self.assertNotIn("\udfff", core_xml)
        self.assertNotIn("\ufdd0", document_xml)
        self.assertNotIn("\ufdef", core_xml)
        self.assertNotIn("\U0001fffe", document_xml)
        self.assertNotIn("\U0010ffff", core_xml)

    def test_source_docx_export_strips_invalid_xml_characters_from_redlines(self):
        source_docx = make_source_docx(["Intro paragraph.\ufdd0", "This Agreement shall be governed by the laws of California."])
        review_result = {
            "paragraphs": [
                {
                    "id": "p1",
                    "index": 1,
                    "source_index": 1,
                    "text": "Intro paragraph.\ufdd0",
                },
                {
                    "id": "p2",
                    "index": 2,
                    "source_index": 2,
                    "text": "This Agreement shall be governed by the laws of California.",
                }
            ],
            "redline_edits": [
                {
                    "id": "r1",
                    "paragraph_id": "p2",
                    "paragraph_index": 2,
                    "source_index": 2,
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "original_text": "This Agreement shall be governed by the laws of California.\ud800\U0002fffe",
                    "replacement_text": "This Agreement shall be governed by the laws of England and Wales.\udfff\ufdef",
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        ET.fromstring(document_xml)
        self.assertIn("England and Wales", document_xml)
        self.assertNotIn("\ud800", document_xml)
        self.assertNotIn("\udfff", document_xml)
        self.assertNotIn("\ufdd0", document_xml)
        self.assertNotIn("\ufdef", document_xml)
        self.assertNotIn("\U0002fffe", document_xml)

    def test_review_report_docx_marks_replace_paragraph_redlines_inline(self):
        result = review_nda("The confidentiality obligations survive for seven years.")

        _settings_root, document_root, _document_xml = docx_xml_roots(build_review_report_docx(result))

        deleted_text = tracked_deleted_text(document_root)
        inserted_text = tracked_inserted_text(document_root)
        self.assertTrue(any("seven" in text for text in deleted_text))
        self.assertFalse(any("The confidentiality obligations survive for seven years." in text for text in deleted_text))
        self.assertTrue(any("a fixed period of up to five" in text for text in inserted_text))

    def test_source_docx_export_redlines_original_document_inline(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
            "The confidentiality obligations survive for three years.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        assert_source_export_has_no_report_leakage(self, redlined_docx)
        settings_root, document_root, document_xml = docx_xml_roots(redlined_docx)
        relationship_targets = docx_document_relationship_targets(redlined_docx)
        content_type_overrides = docx_content_type_overrides(redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertEqual(archive.read("customXml/item1.xml").decode("utf-8"), "<custom>preserved</custom>")

        deleted_text = tracked_deleted_text(document_root)
        inserted_text = tracked_inserted_text(document_root)
        self.assertIsNotNone(settings_root.find(".//w:trackRevisions", W_NS))
        self.assertEqual(relationship_targets[SETTINGS_RELATIONSHIP_TYPE], "settings.xml")
        self.assertEqual(
            content_type_overrides["/word/settings.xml"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml",
        )
        self.assertIn("Intro paragraph.", document_xml)
        self.assertTrue(any("California" in text for text in deleted_text))
        self.assertFalse(any("This Agreement shall be governed by the laws of California." in text for text in deleted_text))
        self.assertTrue(any("England and Wales" in text for text in inserted_text))

    def test_source_docx_facade_drops_manual_insert_before_rendering(self):
        source_docx = make_source_docx(["Anchor paragraph."])
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": "Anchor paragraph."}],
            "redline_edits": [
                {
                    "id": "manual-insert",
                    "clause_id": "manual_viewer_edit",
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "insert_text": "Unsafe manual insertion.",
                }
            ],
        }

        redlined_docx = source_redline_docx.build_source_redline_docx(source_docx, review_result)

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        self.assertEqual(tracked_inserted_text(document_root), [])

    def test_source_docx_facade_rejects_unhealthy_rendered_docx(self):
        source_docx = make_source_docx(["Anchor paragraph."])

        with patch.object(source_redline_docx, "validate_docx_open_health", return_value=["broken package"]):
            with self.assertRaisesRegex(DocxExportError, "redline failed validation"):
                source_redline_docx.build_source_redline_docx(
                    source_docx,
                    {"paragraphs": [], "redline_edits": []},
                )

    def test_source_docx_export_collapses_duplicate_document_xml_entries(self):
        source_docx = duplicate_docx_part(
            make_source_docx(["This Agreement shall be governed by the laws of California."]),
            "word/document.xml",
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertEqual(archive.namelist().count("word/document.xml"), 1)
            document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("England and Wales", document_xml)
        self.assertIn("w:del", document_xml)

    def test_source_redline_preserves_run_formatting_on_replaced_paragraph(self):
        # A bold + colored source paragraph must keep that run formatting after a
        # replace redline; otherwise the redline silently strips bold/italic/color.
        bold_paragraph = (
            '<w:p><w:r><w:rPr><w:b/><w:color w:val="FF0000"/></w:rPr>'
            "<w:t>This Agreement shall be governed by the laws of California.</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{bold_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(run):
            return "".join(node.text or "" for node in run.findall("w:t", W_NS) + run.findall("w:delText", W_NS))

        # Only the runs derived from the replaced bold paragraph must keep the
        # source formatting; brand-new inserted clauses have no source run to
        # inherit from and are correctly left unformatted.
        replaced_runs = [
            run
            for run in document_root.findall(".//w:ins/w:r", W_NS)
            if "England and Wales" in run_text(run)
        ] + [
            run
            for run in document_root.findall(".//w:del/w:r", W_NS)
            if "California" in run_text(run)
        ]
        self.assertEqual(len(replaced_runs), 2, "expected the inserted and deleted runs of the replaced paragraph")
        for run in replaced_runs:
            self.assertIsNotNone(run.find("w:rPr/w:b", W_NS), "replaced paragraph lost source bold formatting")
            self.assertIsNotNone(run.find("w:rPr/w:color", W_NS), "replaced paragraph lost source color formatting")

    def test_source_redline_maps_deleted_text_to_matching_source_run_formatting(self):
        # Mixed-run edited paragraphs must not flatten every tracked run to the
        # first source run's rPr. The deleted "California" text is in the second
        # source run, so it must keep that run's bold/red formatting.
        mixed_paragraph = (
            "<w:p>"
            "<w:r><w:t>This Agreement shall be governed by the laws of </w:t></w:r>"
            '<w:r><w:rPr><w:b/><w:color w:val="FF0000"/></w:rPr><w:t>California</w:t></w:r>'
            "<w:r><w:t>.</w:t></w:r>"
            "</w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{mixed_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        california_runs = [
            run
            for run in document_root.findall(".//w:del/w:r", W_NS)
            if "".join(node.text or "" for node in run.findall("w:delText", W_NS)) == "California"
        ]
        self.assertEqual(len(california_runs), 1, "expected exactly one tracked-deleted California run")
        rpr = california_runs[0].find("w:rPr", W_NS)
        self.assertIsNotNone(rpr, "deleted source run lost its rPr")
        self.assertIsNotNone(rpr.find("w:b", W_NS), "deleted source run lost bold")
        color = rpr.find("w:color", W_NS)
        self.assertIsNotNone(color, "deleted source run lost color")
        self.assertEqual(color.get(f"{{{W_NS['w']}}}val"), "FF0000")

    def test_source_redline_refuses_replace_that_would_drop_hyperlink(self):
        linked_text = "Click here to review the NDA."
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            '<w:p><w:hyperlink w:anchor="bookmark"><w:r><w:t>'
            f"{linked_text}"
            "</w:t></w:r></w:hyperlink></w:p>"
            "</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder"]),
            {"word/document.xml": document_xml},
        )
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": linked_text}],
            "redline_edits": [
                {
                    "id": "edit-1",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": linked_text,
                    "replacement_text": "Click this link to review the NDA.",
                }
            ],
        }

        with self.assertRaisesRegex(DocxExportError, "inline objects, hyperlinks, fields, or note references"):
            build_source_redline_docx(source_docx, review_result)

    def test_source_redline_refuses_delete_that_would_drop_inline_drawing(self):
        drawing_text = "Logo clause."
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>Logo clause.</w:t><w:drawing><w:inline/></w:drawing></w:r></w:p>"
            "</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder"]),
            {"word/document.xml": document_xml},
        )
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": drawing_text}],
            "redline_edits": [
                {
                    "id": "edit-1",
                    "status": "proposed",
                    "action": REDLINE_DELETE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": drawing_text,
                    "replacement_text": "",
                }
            ],
        }

        with self.assertRaisesRegex(DocxExportError, "inline objects, hyperlinks, fields, or note references"):
            build_source_redline_docx(source_docx, review_result)

    def test_source_redline_format_paragraph_emits_tracked_pprchange(self):
        # A format_paragraph redline (alignment left->center AND font ->Arial)
        # must emit the new pPr (jc=center, run-default rFonts=Arial) plus a
        # pPrChange recording the from-state, with the paragraph TEXT untouched.
        formatted_paragraph = (
            '<w:p><w:pPr><w:jc w:val="left"/><w:spacing w:after="120"/></w:pPr>'
            '<w:r><w:rPr><w:rFonts w:ascii="Aptos"/></w:rPr>'
            "<w:t>This Agreement shall be governed by the laws of California.</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_text = "This Agreement shall be governed by the laws of California."
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {"scope": "paragraph", "property": "alignment", "from": "left", "to": "center"},
                        {"scope": "paragraph", "property": "font", "from": "Aptos", "to": "Arial"},
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        # Package opens and every XML part parses.
        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )
        ppr = formatted.find("w:pPr", W_NS)
        self.assertIsNotNone(ppr, "formatted paragraph lost its pPr")

        # New (current) state: alignment centered, run-default font Arial.
        self.assertEqual(ppr.find("w:jc", W_NS).get(f"{{{W_NS['w']}}}val"), "center")
        rfonts = ppr.find("w:rPr/w:rFonts", W_NS)
        self.assertIsNotNone(rfonts, "formatted paragraph lost its run-default rFonts")
        self.assertEqual(rfonts.get(f"{{{W_NS['w']}}}ascii"), "Arial")
        # Other existing pPr children are preserved.
        self.assertIsNotNone(ppr.find("w:spacing", W_NS), "format redline dropped an unrelated pPr child")

        # The tracked change: a pPrChange whose nested original pPr reflects the
        # SOURCE alignment (left) and carries NO nested pPrChange of its own.
        pprchange = ppr.find("w:pPrChange", W_NS)
        self.assertIsNotNone(pprchange, "format redline did not emit a pPrChange")
        self.assertTrue(str(pprchange.get(f"{{{W_NS['w']}}}id") or "").strip(), "pPrChange missing a revision id")
        original_ppr = pprchange.find("w:pPr", W_NS)
        self.assertIsNotNone(original_ppr, "pPrChange missing the original pPr")
        self.assertEqual(original_ppr.find("w:jc", W_NS).get(f"{{{W_NS['w']}}}val"), "left")
        self.assertIsNone(original_ppr.find("w:pPrChange", W_NS), "nested original pPr must not carry a stale pPrChange")

        # The paragraph TEXT is unchanged: no tracked insert/delete on it.
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])
        self.assertEqual(run_text(formatted), formatted_text)

        # Content-coverage gate passes (a format redline only restyles; the source
        # text is fully and faithfully covered).
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_source_redline_format_paragraph_emits_tracked_run_bold(self):
        # A format_paragraph redline carrying a single run op (bold a slice of the
        # paragraph) must split the run so the covered span carries <w:b/> plus an
        # <w:rPrChange> recording the prior (un-bold) rPr, with text byte-identical.
        formatted_text = "This Agreement shall be governed by the laws of California."
        formatted_paragraph = (
            "<w:p>"
            '<w:r><w:rPr><w:rFonts w:ascii="Aptos"/></w:rPr>'
            f"<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        # Bold the word "governed" (offsets into the concatenated paragraph text).
        bold_start = formatted_text.index("governed")
        bold_end = bold_start + len("governed")
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-run-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "bold",
                            "start": bold_start,
                            "end": bold_end,
                            "from": False,
                            "to": True,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        # Package opens and every XML part parses.
        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )

        # Paragraph text is byte-identical, with no tracked insert/delete.
        self.assertEqual(run_text(formatted), formatted_text)
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])

        # The runs split into before / covered / after, with text preserved exactly.
        runs = formatted.findall("w:r", W_NS)
        run_texts = ["".join(n.text or "" for n in r.findall("w:t", W_NS)) for r in runs]
        self.assertEqual("".join(run_texts), formatted_text)
        self.assertIn("governed", run_texts)
        before = formatted_text[:bold_start]
        after = formatted_text[bold_end:]
        self.assertIn(before, run_texts)
        self.assertIn(after, run_texts)

        covered = next(r for r, t in zip(runs, run_texts) if t == "governed")
        covered_rpr = covered.find("w:rPr", W_NS)
        self.assertIsNotNone(covered_rpr, "covered run lost its rPr")
        # New (current) state: <w:b/> present.
        self.assertIsNotNone(covered_rpr.find("w:b", W_NS), "covered run did not gain <w:b/>")
        # The from-state record: an rPrChange whose nested original rPr has NO <w:b/>
        # and carries no nested rPrChange of its own.
        rprchange = covered_rpr.find("w:rPrChange", W_NS)
        self.assertIsNotNone(rprchange, "covered run did not emit an rPrChange")
        self.assertTrue(
            str(rprchange.get(f"{{{W_NS['w']}}}id") or "").strip(),
            "rPrChange missing a revision id",
        )
        original_rpr = rprchange.find("w:rPr", W_NS)
        self.assertIsNotNone(original_rpr, "rPrChange missing the original rPr")
        self.assertIsNone(original_rpr.find("w:b", W_NS), "from-state rPr must not carry <w:b/>")
        self.assertIsNone(
            original_rpr.find("w:rPrChange", W_NS),
            "nested original rPr must not carry a stale rPrChange",
        )
        # The before/after runs keep their original (un-bold) formatting.
        before_run = next(r for r, t in zip(runs, run_texts) if t == before)
        self.assertIsNone(before_run.find("w:rPr/w:b", W_NS), "before run must not be bold")
        self.assertIsNone(before_run.find(".//w:rPrChange", W_NS), "before run must not be tracked")

        # Content-coverage gate passes (a run-format redline only restyles).
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_source_redline_format_paragraph_emits_tracked_run_size(self):
        # A run "size" op (set a slice to 16pt) splits the run so the covered span
        # carries <w:sz w:val="32"/> + <w:szCs w:val="32"/> (half-points) plus an
        # <w:rPrChange> recording the prior (un-sized) rPr, with text byte-identical.
        formatted_text = "This Agreement shall be governed by the laws of California."
        formatted_paragraph = (
            "<w:p>"
            '<w:r><w:rPr><w:rFonts w:ascii="Aptos"/></w:rPr>'
            f"<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        size_start = formatted_text.index("governed")
        size_end = size_start + len("governed")
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-size-run-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "size",
                            "start": size_start,
                            "end": size_end,
                            "from": 0,
                            "to": 16,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )
        # Text byte-identical, no tracked insert/delete.
        self.assertEqual(run_text(formatted), formatted_text)
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])

        runs = formatted.findall("w:r", W_NS)
        run_texts = ["".join(n.text or "" for n in r.findall("w:t", W_NS)) for r in runs]
        self.assertEqual("".join(run_texts), formatted_text)
        covered = next(r for r, t in zip(runs, run_texts) if t == "governed")
        covered_rpr = covered.find("w:rPr", W_NS)
        self.assertIsNotNone(covered_rpr, "covered run lost its rPr")
        # New (current) state: <w:sz>/<w:szCs> in half-points (16pt -> 32).
        sz = covered_rpr.find("w:sz", W_NS)
        self.assertIsNotNone(sz, "covered run did not gain <w:sz/>")
        self.assertEqual(sz.get(f"{{{W_NS['w']}}}val"), "32")
        szcs = covered_rpr.find("w:szCs", W_NS)
        self.assertIsNotNone(szcs, "covered run did not gain <w:szCs/>")
        self.assertEqual(szcs.get(f"{{{W_NS['w']}}}val"), "32")
        # The from-state record: an rPrChange whose nested original rPr has NO size.
        rprchange = covered_rpr.find("w:rPrChange", W_NS)
        self.assertIsNotNone(rprchange, "covered run did not emit an rPrChange")
        self.assertTrue(
            str(rprchange.get(f"{{{W_NS['w']}}}id") or "").strip(),
            "rPrChange missing a revision id",
        )
        original_rpr = rprchange.find("w:rPr", W_NS)
        self.assertIsNotNone(original_rpr, "rPrChange missing the original rPr")
        self.assertIsNone(original_rpr.find("w:sz", W_NS), "from-state rPr must not carry <w:sz/>")
        # The before/after runs keep their original (un-sized) formatting.
        before = formatted_text[:size_start]
        before_run = next(r for r, t in zip(runs, run_texts) if t == before)
        self.assertIsNone(before_run.find("w:rPr/w:sz", W_NS), "before run must not be resized")
        self.assertIsNone(before_run.find(".//w:rPrChange", W_NS), "before run must not be tracked")

        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_source_redline_format_paragraph_emits_tracked_paragraph_size(self):
        # A paragraph "size" op (whole-paragraph 14pt) sets the paragraph-mark run
        # default <w:pPr><w:rPr><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr> plus a
        # pPrChange recording the from-state, with the paragraph TEXT untouched.
        formatted_text = "This Agreement shall be governed by the laws of California."
        formatted_paragraph = (
            '<w:p><w:pPr><w:spacing w:after="120"/></w:pPr>'
            '<w:r><w:rPr><w:rFonts w:ascii="Aptos"/></w:rPr>'
            f"<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-size-para-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {"scope": "paragraph", "property": "size", "from": 0, "to": 14},
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )
        ppr = formatted.find("w:pPr", W_NS)
        self.assertIsNotNone(ppr, "formatted paragraph lost its pPr")
        # New (current) state: run-default sz/szCs in half-points (14pt -> 28).
        sz = ppr.find("w:rPr/w:sz", W_NS)
        self.assertIsNotNone(sz, "paragraph lost its run-default <w:sz/>")
        self.assertEqual(sz.get(f"{{{W_NS['w']}}}val"), "28")
        szcs = ppr.find("w:rPr/w:szCs", W_NS)
        self.assertIsNotNone(szcs, "paragraph lost its run-default <w:szCs/>")
        self.assertEqual(szcs.get(f"{{{W_NS['w']}}}val"), "28")
        # An unrelated pPr child is preserved.
        self.assertIsNotNone(ppr.find("w:spacing", W_NS), "size redline dropped an unrelated pPr child")
        # The tracked change: a pPrChange whose nested original pPr carries no size.
        pprchange = ppr.find("w:pPrChange", W_NS)
        self.assertIsNotNone(pprchange, "size redline did not emit a pPrChange")
        original_ppr = pprchange.find("w:pPr", W_NS)
        self.assertIsNotNone(original_ppr, "pPrChange missing the original pPr")
        self.assertIsNone(original_ppr.find("w:rPr/w:sz", W_NS), "from-state pPr must not carry a size")
        # Text unchanged.
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])
        self.assertEqual(run_text(formatted), formatted_text)

        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_source_redline_run_size_inserts_sz_in_schema_order(self):
        # When the source run's rPr already carries a CT_RPr child that must FOLLOW
        # <w:sz> (here <w:u>), a run size op must insert <w:sz>/<w:szCs> BEFORE it so
        # the rPr stays in schema sequence (sz < szCs < u), not blindly appended.
        formatted_text = "This Agreement shall be governed by the laws of California."
        formatted_paragraph = (
            "<w:p>"
            '<w:r><w:rPr><w:rFonts w:ascii="Aptos"/><w:u w:val="single"/></w:rPr>'
            f"<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        size_start = formatted_text.index("governed")
        size_end = size_start + len("governed")
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-size-order-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {"scope": "run", "property": "size", "start": size_start, "end": size_end, "from": 0, "to": 12},
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            p for p in document_root.findall(".//w:p", W_NS) if run_text(p) == formatted_text
        )
        runs = formatted.findall("w:r", W_NS)
        run_texts = ["".join(n.text or "" for n in r.findall("w:t", W_NS)) for r in runs]
        covered = next(r for r, t in zip(runs, run_texts) if t == "governed")
        children = [child.tag for child in list(covered.find("w:rPr", W_NS))]
        sz_tag = f"{{{W_NS['w']}}}sz"
        szcs_tag = f"{{{W_NS['w']}}}szCs"
        u_tag = f"{{{W_NS['w']}}}u"
        self.assertIn(sz_tag, children)
        self.assertIn(szcs_tag, children)
        self.assertIn(u_tag, children)
        # Schema order: sz before szCs, both before the pre-existing <w:u>.
        self.assertLess(children.index(sz_tag), children.index(szcs_tag))
        self.assertLess(children.index(szcs_tag), children.index(u_tag))

    def test_accept_all_revisions_flattens_to_clean_doc(self):
        # A tracked manual replace -> build_source_redline_docx (tracked) ->
        # accept_all_revisions -> a CLEAN doc: the edit is applied with no <w:ins>,
        # <w:del>, <w:rPrChange>, <w:pPrChange> or trackRevisions setting left.
        original_text = "Air India Limited is hereby named the FIRST PARTY."
        edited_text = "Air India Private Limited is hereby named the FIRST PARTY."
        body_paragraph = f"<w:p><w:r><w:t>{original_text}</w:t></w:r></w:p>"
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{body_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(p["text"]) for p in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        pid = next(str(p["id"]) for p in paragraphs if str(p["text"]) == original_text)
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "action_label": "Your edit",
                    "is_manual": True,
                    "whole_paragraph": False,
                    "paragraph_id": pid,
                    "original_text": original_text,
                    "replacement_text": edited_text,
                }
            ],
        }
        tracked = build_source_redline_docx(source_docx, review_result)
        _s, tracked_root, _t = docx_xml_roots(tracked)
        self.assertTrue(tracked_root.findall(".//w:ins", W_NS), "tracked doc should carry insertions")

        clean = accept_all_revisions(tracked)
        # A clean doc deliberately has Track Changes OFF, so the redline-oriented
        # health check doesn't apply; assert package + XML integrity directly.
        with ZipFile(BytesIO(clean)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        settings_root, clean_root, _c = docx_xml_roots(clean)
        # No revision markup remains anywhere.
        for tag in ("w:ins", "w:del", "w:delText", "w:rPrChange", "w:pPrChange"):
            self.assertEqual(clean_root.findall(f".//{tag}", W_NS), [], f"clean doc still has {tag}")
        if settings_root is not None:
            self.assertIsNone(settings_root.find("w:trackRevisions", W_NS), "trackRevisions should be cleared")
        # The edit is applied: the new text is present and the replaced original gone.
        full = "".join(n.text or "" for n in clean_root.findall(".//w:t", W_NS))
        self.assertIn("Air India Private Limited is hereby named the FIRST PARTY.", full)
        self.assertNotIn("Air India Limited is hereby named", full)

    def test_accept_all_revisions_keeps_replacement_run_formatting(self):
        # A replace redline carrying replacement_runs (the edited paragraph's run model,
        # here one rich formatted run) -> build_source_redline_docx (tracked) ->
        # accept_all_revisions -> a CLEAN doc whose new text keeps the run formatting:
        # the formatted word lands on a run with its rPr, and no <w:ins>/<w:del> remains.
        original_text = "These obligations survive termination of this Agreement."
        edited_text = "These obligations survive expiry of this Agreement."
        body_paragraph = f"<w:p><w:r><w:t>{original_text}</w:t></w:r></w:p>"
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{body_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(p["text"]) for p in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        pid = next(str(p["id"]) for p in paragraphs if str(p["text"]) == original_text)
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "action_label": "Your edit",
                    "is_manual": True,
                    "whole_paragraph": False,
                    "paragraph_id": pid,
                    "original_text": original_text,
                    "replacement_text": edited_text,
                    # Run model: a plain head, a formatted "expiry", then a plain tail.
                    # Its joined text equals replacement_text.
                    "replacement_runs": [
                        {"text": "These obligations survive "},
                        {
                            "text": "expiry",
                            "bold": True,
                            "underline": True,
                            "strike": True,
                            "color": "00AAFF",
                            "highlight": "yellow",
                            "size": 12,
                        },
                        {"text": " of this Agreement."},
                    ],
                }
            ],
        }
        tracked = build_source_redline_docx(source_docx, review_result)
        _s, tracked_root, _t = docx_xml_roots(tracked)
        self.assertTrue(tracked_root.findall(".//w:ins", W_NS), "tracked doc should carry insertions")

        clean = accept_all_revisions(tracked)
        # Package + XML integrity (clean doc has Track Changes OFF).
        with ZipFile(BytesIO(clean)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, clean_root, _c = docx_xml_roots(clean)
        # No revision markup remains anywhere.
        for tag in ("w:ins", "w:del", "w:delText", "w:rPrChange", "w:pPrChange"):
            self.assertEqual(clean_root.findall(f".//{tag}", W_NS), [], f"clean doc still has {tag}")
        # The edited text is applied and the original is gone.
        full = "".join(n.text or "" for n in clean_root.findall(".//w:t", W_NS))
        self.assertEqual(full.count("These obligations survive expiry of this Agreement."), 1)
        self.assertNotIn("survive termination", full)
        # The run model is honoured: the "expiry" run carries the widened formatting,
        # and no other run of this paragraph carries the bold flag.
        bold_run_texts = []
        expiry_rpr = None
        for run in clean_root.findall(".//w:r", W_NS):
            run_text_value = "".join(t.text or "" for t in run.findall("w:t", W_NS))
            if run_text_value == "expiry":
                expiry_rpr = run.find("w:rPr", W_NS)
            if run.find("w:rPr/w:b", W_NS) is None:
                continue
            bold_run_texts.append(run_text_value)
        self.assertIn("expiry", bold_run_texts)
        self.assertNotIn("These obligations survive ", bold_run_texts)
        self.assertNotIn(" of this Agreement.", bold_run_texts)
        self.assertIsNotNone(expiry_rpr, "formatted replacement run missing rPr")
        self.assertIsNotNone(expiry_rpr.find("w:strike", W_NS))
        self.assertEqual(expiry_rpr.find("w:color", W_NS).get(f"{{{W_NS['w']}}}val"), "00AAFF")
        self.assertEqual(expiry_rpr.find("w:highlight", W_NS).get(f"{{{W_NS['w']}}}val"), "yellow")
        self.assertEqual(expiry_rpr.find("w:u", W_NS).get(f"{{{W_NS['w']}}}val"), "single")
        self.assertEqual(expiry_rpr.find("w:sz", W_NS).get(f"{{{W_NS['w']}}}val"), "24")
        self.assertEqual(expiry_rpr.find("w:szCs", W_NS).get(f"{{{W_NS['w']}}}val"), "24")
        expiry_child_tags = [child.tag.rsplit("}", 1)[-1] for child in list(expiry_rpr)]
        self.assertEqual(
            expiry_child_tags,
            ["b", "strike", "color", "sz", "szCs", "highlight", "u"],
        )

    def test_source_redline_format_paragraph_mixes_paragraph_and_run_ops(self):
        # One format_paragraph redline carrying BOTH a paragraph op (alignment) and a
        # run op (italic a slice): the paragraph gains a pPrChange AND the covered run
        # gains <w:i/> + an rPrChange, with the paragraph text untouched.
        formatted_text = "The confidentiality obligations survive for three years."
        formatted_paragraph = (
            '<w:p><w:pPr><w:jc w:val="left"/></w:pPr>'
            f"<w:r><w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        italic_start = formatted_text.index("survive")
        italic_end = italic_start + len("survive")
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-mixed-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {"scope": "paragraph", "property": "alignment", "from": "left", "to": "center"},
                        {
                            "scope": "run",
                            "property": "italic",
                            "start": italic_start,
                            "end": italic_end,
                            "from": False,
                            "to": True,
                        },
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )

        # Paragraph op landed: centered alignment + a pPrChange recording left.
        ppr = formatted.find("w:pPr", W_NS)
        self.assertIsNotNone(ppr, "formatted paragraph lost its pPr")
        self.assertEqual(ppr.find("w:jc", W_NS).get(f"{{{W_NS['w']}}}val"), "center")
        pprchange = ppr.find("w:pPrChange", W_NS)
        self.assertIsNotNone(pprchange, "mixed redline did not emit a pPrChange")
        self.assertEqual(
            pprchange.find("w:pPr/w:jc", W_NS).get(f"{{{W_NS['w']}}}val"),
            "left",
        )

        # Run op landed: the covered span carries <w:i/> + an rPrChange (no <w:i/> in
        # the from-state), and the text is byte-identical with no ins/del.
        self.assertEqual(run_text(formatted), formatted_text)
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])
        runs = formatted.findall("w:r", W_NS)
        run_texts = ["".join(n.text or "" for n in r.findall("w:t", W_NS)) for r in runs]
        self.assertEqual("".join(run_texts), formatted_text)
        covered = next(r for r, t in zip(runs, run_texts) if t == "survive")
        covered_rpr = covered.find("w:rPr", W_NS)
        self.assertIsNotNone(covered_rpr.find("w:i", W_NS), "covered run did not gain <w:i/>")
        rprchange = covered_rpr.find("w:rPrChange", W_NS)
        self.assertIsNotNone(rprchange, "covered run did not emit an rPrChange")
        self.assertIsNone(
            rprchange.find("w:rPr/w:i", W_NS),
            "from-state rPr must not carry <w:i/>",
        )

        # Content-coverage gate passes.
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_source_redline_run_format_offsets_account_for_leading_whitespace(self):
        # BUG 1A residual (offset-seam, strip): the frontend indexes into the STRIPPED
        # paragraph text (docx_text._paragraph_text .strip()s it), but the non-split
        # export path measured run-op offsets from the RAW <w:p>. A paragraph with
        # leading whitespace shifted the <w:rPrChange> left (bolding "world" hit
        # "llo w"). _apply_tracked_run_format now shifts run-op offsets by the
        # leading-whitespace count (raw_text minus raw_text.lstrip()), so the bold
        # lands on exactly the selected word while the raw run structure is kept.
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            '<w:p><w:r><w:t xml:space="preserve">    Hello world</w:t></w:r></w:p>'
            "</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_text = "Hello world"  # leading whitespace stripped in the FE space
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        self.assertEqual(formatted_text.index("world"), 6)
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-ws-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "bold",
                            "start": 6,  # index of "world" in "Hello world"
                            "end": 11,
                            "from": False,
                            "to": True,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:body/w:p", W_NS)
            if "".join(
                node.text or ""
                for run in paragraph.findall("w:r", W_NS)
                for node in run.findall("w:t", W_NS)
            ).strip()
            == formatted_text
        )
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])
        bold_runs = [
            "".join(node.text or "" for node in run.findall("w:t", W_NS))
            for run in formatted.findall("w:r", W_NS)
            if run.find("w:rPr/w:b", W_NS) is not None
        ]
        # The bold + <w:rPrChange> cover EXACTLY "world" -- not a whitespace-shifted "llo w".
        self.assertEqual(bold_runs, ["world"])

    def test_source_redline_run_format_offsets_account_for_tab(self):
        # BUG 1A (offset-seam, tab): the frontend's offset space renders <w:tab> as
        # "\t" (docx_text._run_text), so a run-format op's start/end are measured over
        # tab-aware text. The backend splitter must measure runs the same way -- before
        # the fix it accumulated only the <w:t> bytes, so a tab BEFORE the selection
        # shifted the <w:rPrChange> onto the wrong characters (and could run past end).
        # Here the selected word lives in its own pure-<w:t> run after a tab run; the
        # bold must land EXACTLY on it.
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            "<w:p>"
            "<w:r><w:t>AB</w:t></w:r>"
            "<w:r><w:tab/></w:r>"
            "<w:r><w:t>CD</w:t></w:r>"
            "</w:p></w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        # The tab-bearing paragraph's text in the FE offset space is "AB\tCD" (len 5).
        formatted_text = "AB\tCD"
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        self.assertEqual(formatted_text.index("CD"), 3)  # tab counts as one char
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-tab-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "bold",
                            "start": 3,  # index of "CD" in "AB\tCD"
                            "end": 5,
                            "from": False,
                            "to": True,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def all_text(paragraph):
            parts = []
            for run in paragraph.findall("w:r", W_NS):
                if run.find("w:tab", W_NS) is not None:
                    parts.append("\t")
                parts.append("".join(n.text or "" for n in run.findall("w:t", W_NS)))
            return "".join(parts)

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:body/w:p", W_NS)
            if all_text(paragraph) == formatted_text
        )
        # Text is byte-identical (tab preserved), no ins/del.
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])
        self.assertIsNotNone(formatted.find("w:r/w:tab", W_NS), "tab must survive")

        bold_runs = [
            "".join(n.text or "" for n in run.findall("w:t", W_NS))
            for run in formatted.findall("w:r", W_NS)
            if run.find("w:rPr/w:b", W_NS) is not None
        ]
        # The bold + <w:rPrChange> cover EXACTLY "CD" -- not a tab-shifted "D".
        self.assertEqual(bold_runs, ["CD"])
        covered = next(
            run
            for run in formatted.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == "CD"
        )
        self.assertIsNotNone(covered.find("w:rPr/w:rPrChange", W_NS), "covered run must be tracked")
        # The text before the selection ("AB") keeps its un-bold, untracked state.
        before = next(
            run
            for run in formatted.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == "AB"
        )
        self.assertIsNone(before.find("w:rPr/w:b", W_NS), "AB must not be bold")
        self.assertIsNone(before.find(".//w:rPrChange", W_NS), "AB must not be tracked")

    def test_run_format_offsets_account_for_break(self):
        # BUG 1A (offset-seam, break): the FE renders <w:br> as "\n" (docx_text._run_text)
        # in the offset space the run-format start/end index into. Unlike a <w:tab>, the
        # aligner treats a soft line break as a review-paragraph boundary, so a <w:br>
        # never survives inside ONE review paragraph end-to-end -- this exercises the
        # offset math directly at _apply_tracked_run_format (where the fix lives). A
        # <w:br> in its own run BEFORE the selection must advance the running offset by
        # one ("\n"); the bold then lands on exactly the selected <w:t> segment, not a
        # shifted one. Before the fix the break contributed 0 and the op slipped left.
        paragraph_xml = (
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:r><w:t>Line</w:t></w:r>"
            "<w:r><w:br/></w:r>"
            "<w:r><w:t>Next word</w:t></w:r>"
            "</w:p>"
        )
        source_p = ET.fromstring(paragraph_xml)
        # Offset space (mirror of docx_text._run_text): "Line" + "\n" + "Next word".
        offset_text = "Line\nNext word"
        bold_start = offset_text.index("word")  # 10: Line(4)+\n(1)+"Next "(5)
        self.assertEqual(bold_start, 10)
        bold_end = bold_start + len("word")
        rebuilt, next_rev = _apply_tracked_run_format(
            source_p,
            [
                {
                    "scope": "run",
                    "property": "bold",
                    "start": bold_start,
                    "end": bold_end,
                    "from": False,
                    "to": True,
                }
            ],
            7,
        )
        self.assertEqual(next_rev, 8)  # exactly one revision consumed
        # Break survives, text byte-identical, no ins/del.
        self.assertIsNotNone(rebuilt.find("w:r/w:br", W_NS), "break must survive")
        self.assertEqual(rebuilt.findall(".//w:ins", W_NS), [])
        self.assertEqual(rebuilt.findall(".//w:del", W_NS), [])
        bold_runs = [
            "".join(n.text or "" for n in run.findall("w:t", W_NS))
            for run in rebuilt.findall("w:r", W_NS)
            if run.find("w:rPr/w:b", W_NS) is not None
        ]
        # The bold + <w:rPrChange> cover EXACTLY "word" -- the break did not shift it.
        self.assertEqual(bold_runs, ["word"])
        covered = next(
            run
            for run in rebuilt.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == "word"
        )
        self.assertIsNotNone(covered.find("w:rPr/w:rPrChange", W_NS), "covered run must be tracked")
        # The preceding "Next " stays un-bold/untracked (op started at the right place).
        prefix = next(
            run
            for run in rebuilt.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == "Next "
        )
        self.assertIsNone(prefix.find("w:rPr/w:b", W_NS), '"Next " must not be bold')
        self.assertIsNone(prefix.find(".//w:rPrChange", W_NS), '"Next " must not be tracked')

    def test_run_format_drops_op_past_end_of_offset_space(self):
        # BUG 1A belt-and-braces: an op whose range exceeds the paragraph's offset-space
        # length is clipped/dropped (fail safe) rather than silently mis-placed. Here
        # the run text is "Hello" (len 5) but the op runs to 99 -- it must clip to the
        # available "lo" tail, never index past the end.
        paragraph_xml = (
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:r><w:t>Hello</w:t></w:r></w:p>"
        )
        source_p = ET.fromstring(paragraph_xml)
        rebuilt, next_rev = _apply_tracked_run_format(
            source_p,
            [{"scope": "run", "property": "bold", "start": 3, "end": 99, "from": False, "to": True}],
            1,
        )
        # Text intact; the bold clips to the in-range tail ("lo"), never past the end.
        self.assertEqual(
            "".join(n.text or "" for n in rebuilt.findall(".//w:t", W_NS)), "Hello"
        )
        bold_runs = [
            "".join(n.text or "" for n in run.findall("w:t", W_NS))
            for run in rebuilt.findall("w:r", W_NS)
            if run.find("w:rPr/w:b", W_NS) is not None
        ]
        self.assertEqual(bold_runs, ["lo"])
        # A fully out-of-range op is dropped entirely (no revision, no change).
        rebuilt2, next_rev2 = _apply_tracked_run_format(
            ET.fromstring(paragraph_xml),
            [{"scope": "run", "property": "bold", "start": 50, "end": 99, "from": False, "to": True}],
            1,
        )
        self.assertEqual(next_rev2, 1)
        self.assertEqual(rebuilt2.findall(".//w:b", W_NS), [])
        self.assertEqual(rebuilt2.findall(".//w:rPrChange", W_NS), [])

    def test_source_redline_run_format_lands_on_correct_split_block(self):
        # BUG 1B (split-block re-basing): one physical <w:p> holds two logical blocks
        # split on a blank line. The frontend's run-format op offsets are relative to
        # the SINGLE block's text. Before the fix the block-aware path applied them to
        # the WHOLE physical <w:p>, so bolding "Second" (block-local {0,6}) landed the
        # <w:rPrChange> on "First " in the FIRST block.
        source_docx = make_source_docx_with_internal_blank_line_paragraph(
            ["Alpha clause."], ["First block.", "Second one."]
        )
        extracted = extract_docx_paragraphs(source_docx)
        self.assertEqual(extracted[-1]["text"], "First block.\n\nSecond one.")
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        aligned = align_document_paragraphs(extracted, source_text)
        aligned_by_text = {str(paragraph["text"]): paragraph for paragraph in aligned}
        first_block = aligned_by_text["First block."]
        second_block = aligned_by_text["Second one."]
        self.assertEqual(first_block["source_index"], second_block["source_index"])

        # Bold "Second" -- block-local offsets into "Second one." (0..6).
        review_result = {
            "paragraphs": aligned,
            "redline_edits": [
                {
                    "id": "fmt-block-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": second_block["id"],
                    "paragraph_index": second_block["index"],
                    "source_index": second_block["source_index"],
                    "original_text": "Second one.",
                    "replacement_text": "Second one.",
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "bold",
                            "start": 0,
                            "end": 6,  # "Second"
                            "from": False,
                            "to": True,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        def bold_runs(paragraph):
            return [
                "".join(n.text or "" for n in run.findall("w:t", W_NS))
                for run in paragraph.findall("w:r", W_NS)
                if run.find("w:rPr/w:b", W_NS) is not None
            ]

        first_paragraph = next(
            p for p in document_root.findall(".//w:body/w:p", W_NS) if run_text(p) == "First block."
        )
        second_paragraph = next(
            p for p in document_root.findall(".//w:body/w:p", W_NS) if run_text(p) == "Second one."
        )
        # The bold (and its rPrChange) land on the SECOND block's "Second"...
        self.assertEqual(bold_runs(second_paragraph), ["Second"])
        covered = next(
            run
            for run in second_paragraph.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == "Second"
        )
        self.assertIsNotNone(covered.find("w:rPr/w:rPrChange", W_NS))
        # ...and the FIRST block is untouched -- no bold, no tracked rPrChange leaked.
        self.assertEqual(bold_runs(first_paragraph), [])
        self.assertEqual(first_paragraph.findall(".//w:rPrChange", W_NS), [])

    def test_source_redline_run_format_enables_bold_on_explicit_off_run(self):
        # BUG 2 (be-rprchange): a source run whose rPr already carries an explicit-off
        # <w:b w:val="false"/> shows as un-bold in the FE (docx_text._toggle_property
        # treats false/0/off/none as OFF), so enabling bold emits to:true. Before the
        # fix _set_run_toggle only inserted a toggle when ABSENT, so the explicit-off
        # val survived and the run stayed un-bold (a phantom no-op revision). The new
        # <w:b/> must have NO falsy val, while the nested rPrChange original keeps the
        # val="false" so Word can roll back.
        formatted_text = "Heading text"
        formatted_paragraph = (
            "<w:p>"
            '<w:r><w:rPr><w:b w:val="false"/></w:rPr>'
            f"<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-falsy-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {
                            "scope": "run",
                            "property": "bold",
                            "start": 0,
                            "end": len(formatted_text),
                            "from": False,
                            "to": True,
                        }
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(paragraph):
            return "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS))

        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:body/w:p", W_NS)
            if run_text(paragraph) == formatted_text
        )
        covered = next(
            run
            for run in formatted.findall("w:r", W_NS)
            if "".join(n.text or "" for n in run.findall("w:t", W_NS)) == formatted_text
        )
        covered_rpr = covered.find("w:rPr", W_NS)
        new_bold = covered_rpr.find("w:b", W_NS)
        self.assertIsNotNone(new_bold, "covered run did not gain <w:b/>")
        # The new (current) <w:b/> is genuinely ON -- no falsy val survived.
        self.assertNotIn(
            (new_bold.get(f"{{{W_NS['w']}}}val") or new_bold.get("val") or "").strip().lower(),
            {"false", "0", "off", "none"},
        )
        # And docx_text agrees the run now reads as bold.
        self.assertTrue(docx_text._toggle_property(covered_rpr, "b"))
        # The from-state record keeps the ACTUAL original (explicit-off) so Word rolls back.
        rprchange = covered_rpr.find("w:rPrChange", W_NS)
        self.assertIsNotNone(rprchange, "covered run did not emit an rPrChange")
        original_rpr = rprchange.find("w:rPr", W_NS)
        self.assertIsNotNone(original_rpr, "rPrChange missing the original rPr")
        original_bold = original_rpr.find("w:b", W_NS)
        self.assertIsNotNone(original_bold, "from-state must keep the original <w:b/>")
        self.assertEqual(
            (original_bold.get(f"{{{W_NS['w']}}}val") or original_bold.get("val") or "").strip().lower(),
            "false",
            "from-state <w:b/> must keep its explicit-off val for rollback",
        )
        # Text byte-identical, no ins/del.
        self.assertEqual(run_text(formatted), formatted_text)
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])

    def test_manual_export_redline_sanitizes_widened_format_ops(self):
        cleaned = export_service.clean_manual_export_redline(
            {
                "id": "fmt-rich",
                "action": REDLINE_FORMAT_PARAGRAPH,
                "paragraph_id": "p1",
                "original_text": "Alpha beta gamma delta.",
                "replacement_text": "Alpha beta gamma delta.",
                "format_ops": [
                    {"scope": "run", "property": "underline", "start": 0, "end": 5, "to": True},
                    {"scope": "run", "property": "strike", "start": 6, "end": 10, "to": True},
                    {"scope": "run", "property": "color", "start": 11, "end": 16, "to": "#00aaff"},
                    {"scope": "run", "property": "highlight", "start": 17, "end": 22, "to": "yellow"},
                    {"scope": "run", "property": "color", "start": 0, "end": 5, "to": "not-a-color"},
                    {"scope": "run", "property": "highlight", "start": 0, "end": 5, "to": "sparkle"},
                    {"scope": "paragraph", "property": "underline", "to": True},
                    {"scope": "paragraph", "property": "color", "to": "#00aaff"},
                ],
            }
        )

        self.assertIsNotNone(cleaned)
        self.assertEqual(
            cleaned["format_ops"],
            [
                {"scope": "run", "property": "underline", "to": True, "from": False, "start": 0, "end": 5},
                {"scope": "run", "property": "strike", "to": True, "from": False, "start": 6, "end": 10},
                {"scope": "run", "property": "color", "to": "00AAFF", "from": "", "start": 11, "end": 16},
                {"scope": "run", "property": "highlight", "to": "yellow", "from": "", "start": 17, "end": 22},
            ],
        )

    def test_manual_export_redline_sanitizes_widened_replacement_runs(self):
        cleaned = export_service.clean_manual_export_redline(
            {
                "id": "replace-rich",
                "action": REDLINE_REPLACE_PARAGRAPH,
                "paragraph_id": "p1",
                "original_text": "These obligations survive termination.",
                "replacement_text": "These obligations survive expiry.",
                "replacement_runs": [
                    {"text": "These obligations survive "},
                    {
                        "text": "expiry",
                        "bold": True,
                        "underline": True,
                        "strike": True,
                        "color": "#00aaff",
                        "highlight": "yellow",
                        "size": 12,
                    },
                    {"text": "."},
                ],
            }
        )

        self.assertIsNotNone(cleaned)
        self.assertEqual(
            cleaned["replacement_runs"],
            [
                {"text": "These obligations survive "},
                {
                    "text": "expiry",
                    "bold": True,
                    "underline": True,
                    "strike": True,
                    "size": 12,
                    "color": "00AAFF",
                    "highlight": "yellow",
                },
                {"text": "."},
            ],
        )

    def test_source_redline_format_paragraph_emits_widened_tracked_run_properties(self):
        formatted_text = "Alpha beta gamma delta."
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"<w:p><w:r><w:t>{formatted_text}</w:t></w:r></w:p>"
            "</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-rich-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": [
                        {"scope": "run", "property": "underline", "start": 0, "end": 5, "from": False, "to": True},
                        {"scope": "run", "property": "strike", "start": 6, "end": 10, "from": False, "to": True},
                        {"scope": "run", "property": "color", "start": 11, "end": 16, "from": "", "to": "00AAFF"},
                        {"scope": "run", "property": "highlight", "start": 17, "end": 22, "from": "", "to": "yellow"},
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        formatted = next(
            paragraph
            for paragraph in document_root.findall(".//w:body/w:p", W_NS)
            if "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS)) == formatted_text
        )
        self.assertEqual(formatted.findall(".//w:ins", W_NS), [])
        self.assertEqual(formatted.findall(".//w:del", W_NS), [])

        runs_by_text = {
            "".join(node.text or "" for node in run.findall("w:t", W_NS)): run
            for run in formatted.findall("w:r", W_NS)
        }
        alpha_rpr = runs_by_text["Alpha"].find("w:rPr", W_NS)
        beta_rpr = runs_by_text["beta"].find("w:rPr", W_NS)
        gamma_rpr = runs_by_text["gamma"].find("w:rPr", W_NS)
        delta_rpr = runs_by_text["delta"].find("w:rPr", W_NS)
        self.assertEqual(alpha_rpr.find("w:u", W_NS).get(f"{{{W_NS['w']}}}val"), "single")
        self.assertIsNotNone(beta_rpr.find("w:strike", W_NS))
        self.assertEqual(gamma_rpr.find("w:color", W_NS).get(f"{{{W_NS['w']}}}val"), "00AAFF")
        self.assertEqual(delta_rpr.find("w:highlight", W_NS).get(f"{{{W_NS['w']}}}val"), "yellow")
        for rpr in (alpha_rpr, beta_rpr, gamma_rpr, delta_rpr):
            self.assertIsNotNone(rpr.find("w:rPrChange", W_NS), "tracked run format missing rollback state")
            child_tags = [
                child.tag.rsplit("}", 1)[-1]
                for child in list(rpr)
            ]
            self.assertEqual(child_tags[-1], "rPrChange", "rPrChange must remain the final rPr child")
            self.assertEqual(child_tags, sorted(child_tags, key=("strike", "color", "highlight", "u", "rPrChange").index))

    def test_export_content_coverage_passes_for_full_source_redline(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
            "The confidentiality obligations survive for three years.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in paragraphs)
        result = review_nda(source_text, paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        # A real redline only adds tracked text, so it always covers the source.
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=result["redline_edits"],
            ),
            [],
        )

    def test_export_content_coverage_passes_when_source_table_survives(self):
        source_docx = make_source_docx_with_table(
            ["Intro paragraph."],
            [["Table A", "Table B"]],
            ["The confidentiality obligations survive for three years."],
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in paragraphs)
        result = review_nda(source_text, paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=result["redline_edits"],
                source_docx=source_docx,
            ),
            [],
        )

    def test_export_content_coverage_flags_structural_table_drop(self):
        source_docx = make_source_docx_with_table(
            ["Intro paragraph."],
            [["Table A", "Table B"]],
            ["Outro paragraph."],
        )
        dropped_table_docx = make_source_docx(["Intro paragraph.", "Table A", "Table B", "Outro paragraph."])

        errors = verify_export_content_coverage(
            dropped_table_docx,
            "Intro paragraph.\n\nTable A\n\nTable B\n\nOutro paragraph.",
            source_docx=source_docx,
        )

        self.assertEqual(len(errors), 1)
        self.assertIn("structural counts", errors[0])
        self.assertIn("tables source=1 export=0", errors[0])
        self.assertNotIn("Intro paragraph", errors[0])

    def test_export_content_coverage_flags_empty_body(self):
        empty_docx = make_source_docx([])
        errors = verify_export_content_coverage(
            empty_docx,
            "This Agreement shall be governed by the laws of California.",
        )
        self.assertTrue(errors)

    def test_export_content_coverage_flags_truncated_export(self):
        tiny_docx = make_source_docx(["x"])
        big_source = "This Agreement shall be governed by the laws of California. " * 20
        errors = verify_export_content_coverage(tiny_docx, big_source)
        self.assertTrue(errors)

    def test_export_content_coverage_flags_reordered_source_body(self):
        source_text = "First source paragraph.\n\nSecond source paragraph."
        reordered_docx = make_source_docx(["Second source paragraph.", "First source paragraph."])

        errors = verify_export_content_coverage(reordered_docx, source_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("misplaced, duplicated, or dropped source content", errors[0])

    def test_export_content_coverage_flags_duplicated_source_body(self):
        source_text = "Only source paragraph."
        duplicated_docx = make_source_docx(["Only source paragraph.", "Only source paragraph."])

        errors = verify_export_content_coverage(duplicated_docx, source_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("expected 1", errors[0])

    def test_export_content_coverage_flags_misplaced_expected_insert(self):
        source_docx = make_source_docx(["Anchor one.", "Anchor two."])
        source_text = "Anchor one.\n\nAnchor two."
        misplaced_result = {
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": "Anchor one."},
                {"id": "p2", "index": 2, "source_index": 2, "text": "Anchor two."},
            ],
            "redline_edits": [
                {
                    "id": "r1",
                    "paragraph_id": "p2",
                    "paragraph_index": 2,
                    "source_index": 2,
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "insert_text": "Inserted clause.",
                }
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, misplaced_result)
        expected_redlines = [{
            **misplaced_result["redline_edits"][0],
            "paragraph_id": "p1",
            "paragraph_index": 1,
            "source_index": 1,
        }]

        errors = verify_export_content_coverage(
            redlined_docx,
            source_text,
            expected_redline_edits=expected_redlines,
        )

        self.assertEqual(len(errors), 1)
        self.assertIn("paragraph sequence does not match", errors[0])

    def test_export_content_coverage_flags_duplicated_expected_insert(self):
        source_docx = make_source_docx(["Anchor paragraph."])
        source_text = "Anchor paragraph."
        duplicated_insert_result = {
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": "Anchor paragraph."},
            ],
            "redline_edits": [
                {
                    "id": "r1",
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "insert_text": "Inserted clause.",
                },
                {
                    "id": "r1-copy",
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "insert_text": "Inserted clause.",
                },
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, duplicated_insert_result)

        errors = verify_export_content_coverage(
            redlined_docx,
            source_text,
            expected_redline_edits=duplicated_insert_result["redline_edits"][:1],
        )

        self.assertEqual(len(errors), 1)
        self.assertIn("paragraph sequence does not match", errors[0])

    def test_export_content_coverage_passes_paragraph_with_tab_separator(self):
        # Regression: a native-DOCX paragraph whose words are separated by a
        # <w:tab/> (pervasive in legal NDAs: section-number/title separators,
        # indentation, signature blocks) must NOT trip the accepted-paragraph
        # sequence gate. The exported-side extractor previously dropped the tab
        # entirely ("Section 1.Confidential...") while the expected side mapped
        # it to whitespace ("Section 1. Confidential..."), so the equal-count
        # sequences mismatched and a correct redline raised content_coverage.
        tabbed_document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>Section 1.</w:t></w:r><w:r><w:tab/></w:r>"
            "<w:r><w:t>Confidential Information means all data.</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>The obligations survive for three years.</w:t></w:r></w:p>"
            "</w:body></w:document>"
        )
        docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": tabbed_document_xml},
        )
        # Source text carries the tab as real whitespace (extractors collapse any
        # whitespace run to a single space), so a faithful export matches.
        source_text = (
            "Section 1.\tConfidential Information means all data."
            "\n\nThe obligations survive for three years."
        )

        self.assertEqual(verify_export_content_coverage(docx, source_text), [])

    def test_export_content_coverage_passes_paragraph_with_carriage_return(self):
        # Regression sibling: a soft line break expressed as <w:cr/> inside a
        # paragraph was also dropped by the exported-side extractor while the
        # expected side mapped it to "\n", causing a spurious mismatch.
        cr_document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>First line</w:t><w:cr/><w:t>second line.</w:t></w:r></w:p>"
            "</w:body></w:document>"
        )
        docx = replace_docx_parts(
            make_source_docx(["placeholder one"]),
            {"word/document.xml": cr_document_xml},
        )
        source_text = "First line\nsecond line."

        self.assertEqual(verify_export_content_coverage(docx, source_text), [])

    def test_export_content_coverage_still_flags_drop_in_tabbed_paragraph(self):
        # The tab/cr alignment must NOT weaken the gate: genuinely dropped
        # content in a tab-bearing document is still caught. Here the export is
        # missing the entire second source paragraph.
        tabbed_document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>Section 1.</w:t></w:r><w:r><w:tab/></w:r>"
            "<w:r><w:t>Confidential Information means all data.</w:t></w:r></w:p>"
            "</w:body></w:document>"
        )
        docx = replace_docx_parts(
            make_source_docx(["placeholder one"]),
            {"word/document.xml": tabbed_document_xml},
        )
        source_text = (
            "Section 1.\tConfidential Information means all data."
            "\n\nThe obligations survive for three years."
        )

        errors = verify_export_content_coverage(docx, source_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("paragraph sequence does not match", errors[0])
        self.assertNotIn("Confidential Information", errors[0])

    def test_export_content_coverage_passes_for_letterhead_footer_native_docx(self):
        # Regression for the Moorwand letterhead bug: a native-DOCX matter whose source
        # has a non-empty footer. ``extracted_text`` joins BODY + footer (supplemental),
        # but the reviewed-DOCX export reconstructs only the body (footer copied through
        # verbatim), so the body-only export can never reproduce the footer text. Feeding
        # the full supplemental-inclusive ``extracted_text`` to the body-only gate
        # false-rejected a faithful export with HTTP 500. The fix passes the BODY-ONLY
        # text instead, which the resolver derives from the source bytes.
        from nda_automation import redline_export_service

        source_docx = make_source_docx_with_footer(
            [
                "Confidential Information means all data disclosed under this Agreement.",
                "The confidentiality obligations survive for three years.",
            ],
            [
                "Moorwand Ltd | Registered office address: Fora, 3 Lloyds Avenue | "
                "London | EC3N 3DS | Company No. 08491211",
            ],
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        full_extracted_text = "\n\n".join(str(paragraph["text"]) for paragraph in paragraphs)
        # The footer text is in the supplemental-inclusive extracted_text...
        self.assertIn("Company No. 08491211", full_extracted_text)
        # A FAITHFUL reviewed export reconstructs the body and copies the footer through
        # verbatim; here we model the faithful body-only export bytes (no body content
        # dropped/reordered). The body text is unchanged from the source body.
        faithful_export = make_source_docx_with_footer(
            [
                "Confidential Information means all data disclosed under this Agreement.",
                "The confidentiality obligations survive for three years.",
            ],
            [
                "Moorwand Ltd | Registered office address: Fora, 3 Lloyds Avenue | "
                "London | EC3N 3DS | Company No. 08491211",
            ],
        )

        # BUG (origin/main): the full extracted_text requires the footer text on the
        # body-only export side. The export side reads word/document.xml only (the body),
        # so the footer paragraph in the EXPECTED sequence has no counterpart -> the
        # accepted-paragraph sequence diverges and the gate false-rejects.
        full_text_errors = verify_export_content_coverage(faithful_export, full_extracted_text)
        self.assertTrue(
            full_text_errors,
            "expected the full supplemental-inclusive text to trip the body-only gate "
            "(this is the bug condition that must hold before the fix)",
        )

        # FIX: the body-only expected text the resolver derives from the source bytes
        # excludes the footer, so the faithful export now PASSES.
        body_only_text = redline_export_service._body_extracted_text_from_docx_bytes(
            source_docx, "moorwand-nda.docx"
        )
        self.assertNotIn("Company No. 08491211", body_only_text)
        self.assertIn("Confidential Information means all data", body_only_text)
        self.assertEqual(
            verify_export_content_coverage(faithful_export, body_only_text),
            [],
        )

    def test_body_extracted_text_from_docx_bytes_strips_footer(self):
        from nda_automation import redline_export_service

        source_docx = make_source_docx_with_footer(
            ["Body clause one.", "Body clause two."],
            ["Footer: Company No. 99999999"],
        )
        body_only = redline_export_service._body_extracted_text_from_docx_bytes(
            source_docx, "letterhead.docx"
        )
        self.assertEqual(body_only, "Body clause one.\n\nBody clause two.")

    def test_body_extracted_text_from_docx_bytes_safe_on_non_docx_or_missing(self):
        from nda_automation import redline_export_service

        # Non-DOCX filename and missing bytes both return "" (caller falls back), never raise.
        self.assertEqual(
            redline_export_service._body_extracted_text_from_docx_bytes(b"%PDF-1.4", "scan.pdf"),
            "",
        )
        self.assertEqual(
            redline_export_service._body_extracted_text_from_docx_bytes(None, "x.docx"),
            "",
        )
        # Corrupt DOCX bytes degrade to "" rather than propagating an extraction error.
        self.assertEqual(
            redline_export_service._body_extracted_text_from_docx_bytes(b"not a zip", "x.docx"),
            "",
        )

    def test_body_expected_source_text_precedence(self):
        from nda_automation import redline_export_service

        field = redline_export_service.BODY_EXTRACTED_TEXT_FIELD
        # 1. The explicit body-only field wins.
        self.assertEqual(
            redline_export_service._body_expected_source_text(
                {field: "Body only.", "extracted_text": "Body only.\n\nFooter."}
            ),
            "Body only.",
        )
        # 2. With no field but marked paragraphs, supplemental paragraphs are filtered.
        self.assertEqual(
            redline_export_service._body_expected_source_text(
                {
                    "paragraphs": [
                        {"text": "Body A", "source_kind": "paragraph"},
                        {"text": "Footer", "source_kind": "supplemental"},
                    ],
                    "extracted_text": "Body A\n\nFooter",
                }
            ),
            "Body A",
        )
        # 3. With unmarked paragraphs (the deferred path), fall back to extracted_text
        #    verbatim -- the resolver does not silently drop body text it cannot classify.
        self.assertEqual(
            redline_export_service._body_expected_source_text(
                {
                    "paragraphs": [{"text": "Body A"}, {"text": "Footer"}],
                    "extracted_text": "Body A\n\nFooter",
                }
            ),
            "Body A\n\nFooter",
        )

    def test_export_content_coverage_still_flags_body_drop_with_footer_source(self):
        # PROTECTION INTACT: with the footer excluded from the expected text, a genuine
        # BODY paragraph drop is still caught. Three substantial clauses, one dropped --
        # the surviving two keep the export above the 0.5 length ratio so the stronger
        # accepted-paragraph SEQUENCE check (not just the ratio precheck) is what fires.
        from nda_automation import redline_export_service

        clause_one = "Confidential Information means all data disclosed under this Agreement."
        clause_two = "The confidentiality obligations survive for a period of three years."
        clause_three = "This Agreement shall be governed by the laws of England and Wales."
        source_docx = make_source_docx_with_footer(
            [clause_one, clause_two, clause_three],
            ["Footer: Company No. 12345678"],
        )
        body_only_text = redline_export_service._body_extracted_text_from_docx_bytes(
            source_docx, "letterhead.docx"
        )
        # A faithful body-only EXPORT missing the MIDDLE clause (clause_two dropped).
        dropped_export = make_source_docx([clause_one, clause_three])

        errors = verify_export_content_coverage(dropped_export, body_only_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("misplaced, duplicated, or dropped source content", errors[0])
        self.assertNotIn("confidentiality obligations", errors[0])

    def test_export_content_coverage_still_flags_body_reorder_with_footer_source(self):
        from nda_automation import redline_export_service

        source_docx = make_source_docx_with_footer(
            ["First body clause.", "Second body clause."],
            ["Footer line."],
        )
        body_only_text = redline_export_service._body_extracted_text_from_docx_bytes(
            source_docx, "letterhead.docx"
        )
        reordered_export = make_source_docx(["Second body clause.", "First body clause."])

        errors = verify_export_content_coverage(reordered_export, body_only_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("misplaced, duplicated, or dropped source content", errors[0])

    def test_export_content_coverage_still_flags_body_duplication_with_footer_source(self):
        from nda_automation import redline_export_service

        source_docx = make_source_docx_with_footer(
            ["Only body clause."],
            ["Footer line."],
        )
        body_only_text = redline_export_service._body_extracted_text_from_docx_bytes(
            source_docx, "letterhead.docx"
        )
        duplicated_export = make_source_docx(["Only body clause.", "Only body clause."])

        errors = verify_export_content_coverage(duplicated_export, body_only_text)

        self.assertEqual(len(errors), 1)
        self.assertIn("expected 1", errors[0])

    def test_export_content_coverage_logs_truncated_diff_on_sequence_failure(self):
        # PERMANENT DIAGNOSABILITY: a coverage rejection emits ONE bounded WARNING
        # carrying the failed check, the diverging paragraph index, and a truncated
        # whitespace-collapsed window of each side -- never the full document. A REORDER
        # keeps the total length identical (so the ratio precheck passes) and trips the
        # stronger accepted-paragraph SEQUENCE check.
        long_first = ("Alpha " * 200).strip() + "."  # ~1200 chars, far beyond 80.
        long_second = ("Bravo " * 200).strip() + "."
        source_text = f"{long_first}\n\n{long_second}"
        reordered_export = make_source_docx([long_second, long_first])

        with self.assertLogs("nda_automation.docx_health", level="WARNING") as captured:
            errors = verify_export_content_coverage(reordered_export, source_text)

        self.assertTrue(errors)
        self.assertEqual(len(captured.records), 1)
        message = captured.records[0].getMessage()
        self.assertIn("content-coverage gate rejected export", message)
        self.assertIn("check=sequence", message)
        self.assertIn("diverging_paragraph_index=0", message)
        # The whole 1200-char source paragraph must NOT appear; only a bounded window.
        self.assertNotIn(long_first, message)
        self.assertNotIn(long_second, message)
        # The window is single-line (whitespace collapsed) and short.
        self.assertNotIn("\n", message)
        self.assertLessEqual(message.count("Alpha"), 20)
        self.assertLessEqual(message.count("Bravo"), 20)
        # The client-facing error string stays count-only (no document text leaked).
        self.assertNotIn("Alpha", errors[0])
        self.assertNotIn("Bravo", errors[0])

    def test_export_content_coverage_logs_truncated_diff_on_ratio_failure(self):
        tiny_docx = make_source_docx(["x"])
        big_source = "Confidential disclosure obligations. " * 30

        with self.assertLogs("nda_automation.docx_health", level="WARNING") as captured:
            errors = verify_export_content_coverage(tiny_docx, big_source)

        self.assertTrue(errors)
        self.assertEqual(len(captured.records), 1)
        message = captured.records[0].getMessage()
        self.assertIn("check=ratio", message)
        self.assertNotIn("\n", message)
        # Bounded window: the repeated phrase appears at most a couple of times, never 30.
        self.assertLessEqual(message.count("Confidential"), 3)

    def test_export_content_coverage_ignores_missing_source_text(self):
        docx = make_source_docx(["Some body text."])
        self.assertEqual(verify_export_content_coverage(docx, ""), [])

    def test_export_content_coverage_rejects_excessive_uncompressed_docx_size(self):
        oversized_docx = make_source_docx(["A" * 4096])

        with patch.object(docx_text, "MAX_DOCX_UNCOMPRESSED_BYTES", 1024):
            errors = verify_export_content_coverage(oversized_docx, "Expected source text.")

        self.assertEqual(errors, ["Exported document body contains no text."])

    def test_docx_open_health_flags_duplicate_zip_entries(self):
        source_docx = duplicate_docx_part(make_source_docx(["Some body text."]), "word/document.xml")

        errors = validate_docx_open_health(source_docx)

        self.assertIn("DOCX package contains duplicate entries: word/document.xml.", errors)

    def test_source_docx_export_writes_native_word_comments(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)
        result["review_comments"] = [
            {
                "author": "Reviewer",
                "clause_id": "governing_law",
                "clause_name": "Governing Law",
                "paragraph_id": "p2",
                "text": "Confirm England and Wales is acceptable.",
            }
        ]

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        relationship_targets = docx_document_relationship_targets(redlined_docx)
        content_type_overrides = docx_content_type_overrides(redlined_docx)
        comments_root, document_root, comments_xml, document_xml = docx_comments(redlined_docx)
        self.assertEqual(relationship_targets[COMMENTS_RELATIONSHIP_TYPE], "comments.xml")
        self.assertEqual(content_type_overrides["/word/comments.xml"], COMMENTS_CONTENT_TYPE)
        self.assertIn("Confirm England and Wales is acceptable.", comments_xml)
        self.assertEqual(len(comments_root.findall(".//w:comment", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentRangeStart", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentRangeEnd", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentReference", W_NS)), 1)
        self.assertIn('w:commentRangeStart w:id="0"', document_xml)

    def test_source_docx_export_does_not_count_nested_textbox_paragraphs_for_comment_anchors(self):
        target_text = "This Agreement shall be governed by the laws of California."
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body>"
            "<w:p><w:r><w:t>Intro paragraph.</w:t>"
            "<w:drawing><w:txbxContent><w:p><w:r><w:t>Textbox note.</w:t></w:r></w:p></w:txbxContent></w:drawing>"
            "</w:r></w:p>"
            f"<w:p><w:r><w:t>{target_text}</w:t></w:r></w:p>"
            "</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        review_result = {
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": paragraphs[0]["text"]},
                {"id": "p2", "index": 2, "source_index": 2, "text": target_text},
            ],
            "review_comments": [
                {
                    "author": "Reviewer",
                    "clause_id": "governing_law",
                    "clause_name": "Governing Law",
                    "paragraph_id": "p2",
                    "text": "Comment should target the governing law paragraph.",
                }
            ],
            "redline_edits": [],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        _comments_root, _document_root, comments_xml, exported_xml = docx_comments(redlined_docx)
        self.assertIn("Comment should target the governing law paragraph.", comments_xml)
        target_index = exported_xml.index(target_text)
        comment_start_index = exported_xml.index("commentRangeStart")
        comment_end_index = exported_xml.index("commentRangeEnd")
        textbox_index = exported_xml.index("Textbox note.")
        self.assertGreater(comment_start_index, textbox_index)
        self.assertLess(comment_start_index, target_index)
        self.assertGreater(comment_end_index, target_index)

    def test_source_docx_export_anchors_selected_text_comments(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement includes ordinary operational terms.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)
        result["review_comments"] = [
            {
                "author": "Reviewer",
                "paragraph_id": "p1",
                "scope": "selection",
                "selected_text": "Intro",
                "selection_start": 0,
                "selection_end": 5,
                "text": "Selected text comment.",
            }
        ]

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        _comments_root, document_root, comments_xml, document_xml = docx_comments(redlined_docx)
        self.assertIn("Selected text comment.", comments_xml)
        self.assertEqual(len(document_root.findall(".//w:commentRangeStart", W_NS)), 1)
        start_index = document_xml.index("commentRangeStart")
        selected_index = document_xml.index("Intro")
        end_index = document_xml.index("commentRangeEnd")
        self.assertLess(start_index, selected_index)
        self.assertLess(selected_index, end_index)

    def test_source_docx_export_keeps_comment_anchor_on_original_paragraph_after_insert_redline(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement includes ordinary operational terms.",
            "Insert the new carve-out after this paragraph.",
            "Comment target should remain here.",
        ])
        review_result = {
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": "Intro paragraph."},
                {
                    "id": "p2",
                    "index": 2,
                    "source_index": 2,
                    "text": "This Agreement includes ordinary operational terms.",
                },
                {
                    "id": "p3",
                    "index": 3,
                    "source_index": 3,
                    "text": "Insert the new carve-out after this paragraph.",
                },
                {
                    "id": "p4",
                    "index": 4,
                    "source_index": 4,
                    "text": "Comment target should remain here.",
                },
            ],
            "review_comments": [
                {
                    "author": "Reviewer",
                    "paragraph_id": "p4",
                    "text": "Comment belongs to the original fourth paragraph.",
                }
            ],
            "redline_edits": [
                {
                    "id": "r1",
                    "paragraph_id": "p3",
                    "paragraph_index": 3,
                    "source_index": 3,
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "insert_text": "Inserted redline paragraph must not steal the comment.",
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _comments_root, document_root, comments_xml, _document_xml = docx_comments(redlined_docx)
        self.assertIn("Comment belongs to the original fourth paragraph.", comments_xml)

        anchored_paragraph = self._paragraph_with_comment_start(document_root, "0")
        self.assertIsNotNone(anchored_paragraph)
        anchored_text = revision_text_for_state(anchored_paragraph, accepted=True)
        self.assertIn("Comment target should remain here.", anchored_text)
        self.assertNotIn("Inserted redline paragraph must not steal the comment.", anchored_text)

    def test_review_report_docx_writes_native_word_comments(self):
        result = mark_ai_executed(review_nda(
            "This Agreement shall be governed by the laws of California.\n\n"
            "The confidentiality obligations survive for three years."
        ))
        result["review_comments"] = [
            {
                "author": "Reviewer",
                "clause_id": "term_and_survival",
                "clause_name": "Term and Survival",
                "paragraph_id": "p2",
                "text": "Check whether a three-year term is required commercially.",
            }
        ]

        report_docx = build_review_report_docx(result)

        assert_docx_package_healthy(self, report_docx, require_styles=True)
        relationship_targets = docx_document_relationship_targets(report_docx)
        content_type_overrides = docx_content_type_overrides(report_docx)
        comments_root, document_root, comments_xml, _document_xml = docx_comments(report_docx)
        self.assertEqual(relationship_targets[COMMENTS_RELATIONSHIP_TYPE], "comments.xml")
        self.assertEqual(content_type_overrides["/word/comments.xml"], COMMENTS_CONTENT_TYPE)
        self.assertIn("Check whether a three-year term is required commercially.", comments_xml)
        self.assertEqual(len(comments_root.findall(".//w:comment", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentReference", W_NS)), 1)

    def test_source_docx_export_writes_resolved_threaded_comments(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)
        result["review_comments"] = [
            {
                "id": "root-1",
                "author": "Reviewer",
                "clause_id": "governing_law",
                "clause_name": "Governing Law",
                "paragraph_id": "p2",
                "text": "Should this be England and Wales instead?",
                "resolved": True,
            },
            {
                "id": "reply-1",
                "parent_id": "root-1",
                "author": "Counsel",
                "paragraph_id": "p2",
                "text": "Agreed, switching to England and Wales.",
            },
        ]

        redlined_docx = build_source_redline_docx(source_docx, result)

        # The package opens as a zip and every XML part parses.
        assert_docx_package_healthy(self, redlined_docx)
        assert_every_xml_part_parses(self, redlined_docx)

        comments_root, document_root, comments_xml, document_xml = docx_comments(redlined_docx)
        comments_extended_root, comments_extended_xml = docx_comments_extended(redlined_docx)

        # word/comments.xml has 2 <w:comment> entries, each with a <w:p> carrying a w14:paraId.
        comments = comments_root.findall(".//w:comment", W_NS)
        self.assertEqual(len(comments), 2)
        para_ids = []
        for comment in comments:
            comment_para_ids = [
                paragraph.attrib.get(f"{{{W14_NS['w14']}}}paraId")
                for paragraph in comment.findall(".//w:p", W_NS)
            ]
            self.assertTrue(comment_para_ids and all(comment_para_ids))
            for para_id in comment_para_ids:
                self.assertRegex(para_id, r"^[0-9A-F]{8}$")
            para_ids.append(comment_para_ids[0])
        self.assertIn("Should this be England and Wales instead?", comments_xml)
        self.assertIn("Agreed, switching to England and Wales.", comments_xml)

        # Only the ROOT has the in-document range + a CommentReference run.
        self.assertEqual(len(document_root.findall(".//w:commentRangeStart", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentRangeEnd", W_NS)), 1)
        self.assertEqual(len(document_root.findall(".//w:commentReference", W_NS)), 1)

        # word/commentsExtended.xml has 2 <w15:commentEx>; the reply points at the
        # root's paraId and both are done="1".
        comment_ex = comments_extended_root.findall(".//w15:commentEx", W15_NS)
        self.assertEqual(len(comment_ex), 2)
        by_para_id = {entry.attrib.get(f"{{{W15_NS['w15']}}}paraId"): entry for entry in comment_ex}
        root_para_id = self._root_comment_para_id(comments_root)
        self.assertIn(root_para_id, by_para_id)
        reply_entry = next(
            entry
            for entry in comment_ex
            if entry.attrib.get(f"{{{W15_NS['w15']}}}paraIdParent")
        )
        self.assertEqual(reply_entry.attrib[f"{{{W15_NS['w15']}}}paraIdParent"], root_para_id)
        for entry in comment_ex:
            self.assertEqual(entry.attrib.get(f"{{{W15_NS['w15']}}}done"), "1")
        # The root entry itself carries no paraIdParent.
        self.assertNotIn(f"{{{W15_NS['w15']}}}paraIdParent", by_para_id[root_para_id].attrib)

        # [Content_Types].xml + word/_rels/document.xml.rels register the new part.
        content_type_overrides = docx_content_type_overrides(redlined_docx)
        relationship_targets = docx_document_relationship_targets(redlined_docx)
        self.assertEqual(
            content_type_overrides["/word/commentsExtended.xml"], COMMENTS_EXTENDED_CONTENT_TYPE
        )
        self.assertEqual(
            relationship_targets[COMMENTS_EXTENDED_RELATIONSHIP_TYPE], "commentsExtended.xml"
        )

    def test_review_report_docx_writes_resolved_threaded_comments(self):
        result = mark_ai_executed(review_nda(
            "This Agreement shall be governed by the laws of California.\n\n"
            "The confidentiality obligations survive for three years."
        ))
        result["review_comments"] = [
            {
                "id": "root-1",
                "author": "Reviewer",
                "clause_id": "term_and_survival",
                "clause_name": "Term and Survival",
                "paragraph_id": "p2",
                "text": "Is a three-year survival period acceptable?",
                "resolved": True,
            },
            {
                "id": "reply-1",
                "parent_id": "root-1",
                "author": "Counsel",
                "paragraph_id": "p2",
                "text": "Yes, three years is fine.",
            },
        ]

        report_docx = build_review_report_docx(result)

        assert_docx_package_healthy(self, report_docx, require_styles=True)
        assert_every_xml_part_parses(self, report_docx)

        comments_root, document_root, _comments_xml, _document_xml = docx_comments(report_docx)
        comments_extended_root, _comments_extended_xml = docx_comments_extended(report_docx)
        self.assertEqual(len(comments_root.findall(".//w:comment", W_NS)), 2)
        # Only the root is anchored in the report body.
        self.assertEqual(len(document_root.findall(".//w:commentReference", W_NS)), 1)
        comment_ex = comments_extended_root.findall(".//w15:commentEx", W15_NS)
        self.assertEqual(len(comment_ex), 2)
        self.assertTrue(
            any(entry.attrib.get(f"{{{W15_NS['w15']}}}paraIdParent") for entry in comment_ex)
        )
        for entry in comment_ex:
            self.assertEqual(entry.attrib.get(f"{{{W15_NS['w15']}}}done"), "1")

    def test_source_docx_export_marks_unresolved_thread_not_done(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)
        result["review_comments"] = [
            {
                "id": "root-1",
                "author": "Reviewer",
                "paragraph_id": "p2",
                "text": "Open question on governing law.",
            },
            {
                # A stray resolved flag on a reply must NOT resolve the thread; only
                # the root's resolved state drives done.
                "id": "reply-1",
                "parent_id": "root-1",
                "author": "Counsel",
                "paragraph_id": "p2",
                "text": "Still thinking about it.",
                "resolved": True,
            },
        ]

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_every_xml_part_parses(self, redlined_docx)
        comments_extended_root, _xml = docx_comments_extended(redlined_docx)
        comment_ex = comments_extended_root.findall(".//w15:commentEx", W15_NS)
        self.assertEqual(len(comment_ex), 2)
        for entry in comment_ex:
            self.assertEqual(entry.attrib.get(f"{{{W15_NS['w15']}}}done"), "0")

    def test_source_docx_export_appends_to_existing_comments_extended_part(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "This Agreement shall be governed by the laws of California.",
        ])
        existing_comments_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml">'
            '<w:comment w:id="0" w:author="Prior" w:date="2024-01-01T00:00:00Z" w:initials="P">'
            '<w:p w14:paraId="0ABCDEF0"><w:r><w:t>Pre-existing comment.</w:t></w:r></w:p>'
            '</w:comment></w:comments>'
        )
        existing_comments_extended_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
            '<w15:commentEx w15:paraId="0ABCDEF0" w15:done="0"/>'
            '</w15:commentsEx>'
        )
        source_docx = replace_docx_parts(
            source_docx,
            {
                "word/comments.xml": existing_comments_xml,
                "word/commentsExtended.xml": existing_comments_extended_xml,
            },
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)
        result["review_comments"] = [
            {
                "id": "root-1",
                "author": "Reviewer",
                "paragraph_id": "p2",
                "text": "A new top-level comment.",
                "resolved": True,
            }
        ]

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_every_xml_part_parses(self, redlined_docx)
        comments_root, _document_root, _comments_xml, _document_xml = docx_comments(redlined_docx)
        comments_extended_root, _xml = docx_comments_extended(redlined_docx)
        # Pre-existing comment is retained and the new one appended (no duplication).
        self.assertEqual(len(comments_root.findall(".//w:comment", W_NS)), 2)
        comment_ex = comments_extended_root.findall(".//w15:commentEx", W15_NS)
        self.assertEqual(len(comment_ex), 2)
        para_ids = [entry.attrib.get(f"{{{W15_NS['w15']}}}paraId") for entry in comment_ex]
        self.assertIn("0ABCDEF0", para_ids)
        # The single commentsExtended override/relationship is not duplicated.
        with ZipFile(BytesIO(redlined_docx)) as archive:
            self.assertEqual(archive.namelist().count("word/commentsExtended.xml"), 1)
        content_type_overrides = docx_content_type_overrides(redlined_docx)
        self.assertEqual(
            content_type_overrides["/word/commentsExtended.xml"], COMMENTS_EXTENDED_CONTENT_TYPE
        )

    @staticmethod
    def _root_comment_para_id(comments_root):
        """The w14:paraId on the comment that the document anchors (the only one with
        a commentRangeStart) is the thread root; here we resolve it via the comment
        whose w:id is 0 -- the first/root comment appended in the build."""
        for comment in comments_root.findall(".//w:comment", W_NS):
            if comment.attrib.get(f"{{{W_NS['w']}}}id") == "0":
                paragraph = comment.find(".//w:p", W_NS)
                return paragraph.attrib.get(f"{{{W14_NS['w14']}}}paraId")
        return None

    @staticmethod
    def _paragraph_with_comment_start(document_root, comment_id):
        for paragraph in document_root.findall(".//w:p", W_NS):
            for marker in paragraph.findall(".//w:commentRangeStart", W_NS):
                if marker.attrib.get(f"{{{W_NS['w']}}}id") == comment_id:
                    return paragraph
        return None

    def test_source_docx_export_preserves_ignorable_namespace_prefixes(self):
        source_text = "This Agreement shall be governed by the laws of California."
        document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
  xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
  xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
  mc:Ignorable="w14 wp14">
  <w:body><w:p><w:r><w:t>{escape_xml(source_text)}</w:t></w:r></w:p></w:body>
</w:document>"""
        source_docx = replace_docx_parts(
            make_source_docx([source_text]),
            {"word/document.xml": document_xml},
        )
        result = review_nda(source_text)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            redlined_xml = archive.read("word/document.xml").decode("utf-8")
        ET.fromstring(redlined_xml)
        self.assertIn('mc:Ignorable="w14 wp14"', redlined_xml)
        self.assertIn('xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"', redlined_xml)
        self.assertIn('xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"', redlined_xml)
        self.assertIn('xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"', redlined_xml)
        self.assertNotIn("ns0:Ignorable", redlined_xml)
        self.assertNotIn("xmlns:ns", redlined_xml)

    def test_source_docx_export_preserves_grouped_numbers_and_non_ascii_words(self):
        original = "Payment cap is 1,000 for café records."
        replacement = "Payment cap is 1,000 for café documents."
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 1,
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraph = document_root.find(".//w:body/w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)

    def test_manual_freeform_replace_redlines_at_character_level(self):
        # A free-form manual edit (clause_id == manual_viewer_edit, no whole_paragraph)
        # must diff at the CHARACTER level like the frontend preview: "color" -> "colour"
        # keeps "colo", inserts only "u", keeps "r" -- never strikes the whole word.
        original = "Favourite color is blue."
        replacement = "Favourite colour is blue."
        source_docx = make_source_docx([original])
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraph = document_root.find(".//w:body/w:p", W_NS)
        # Accepting/rejecting still reconstructs both states exactly.
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        # Char-level: only "u" is inserted, nothing is deleted (pure insertion).
        self.assertEqual(tracked_inserted_text(document_root), ["u"])
        self.assertEqual(tracked_deleted_text(document_root), [])
        # NOT a whole-word ~~color~~ colour replacement: "colour" never appears whole
        # inside an insert run, and "color" is never struck whole inside a delete run.
        self.assertNotIn("colour", tracked_inserted_text(document_root))
        self.assertNotIn("color", tracked_deleted_text(document_root))

    def test_manual_freeform_replace_deletes_only_changed_chars_verbatim(self):
        # Deleting a word inside a quoted phrase: only " India" (with its leading space)
        # is struck, verbatim -- no spurious spacing introduced around the curly quotes.
        original = "The party known as “Air India” is referenced."
        replacement = "The party known as “Air” is referenced."
        source_docx = make_source_docx([original])
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraph = document_root.find(".//w:body/w:p", W_NS)
        # Exactly one deletion, carrying just " India" verbatim (leading space kept).
        self.assertEqual(tracked_deleted_text(document_root), [" India"])
        self.assertEqual(tracked_inserted_text(document_root), [])
        # Both states reconstruct byte-identically -- the curly quotes are untouched.
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)

    def test_manual_whole_paragraph_replace_does_not_use_char_level(self):
        # A manual redline explicitly flagged whole_paragraph (e.g. a clause/governing
        # -law pick surfaced through the manual path) must stay on the token/whole
        # path, NOT the char-level diff. A char-level diff would fragment "California"
        # into single-letter runs (sharing letters with "Wales"); the token path
        # strikes the whole word " California" instead.
        original = "Governed by the laws of California."
        replacement = "Governed by the laws of England and Wales."
        source_docx = make_source_docx([original])
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                    "whole_paragraph": True,
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraph = document_root.find(".//w:body/w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        # Token/whole path: the whole word " California" is struck as ONE delete run,
        # not fragmented into single-character runs the way char-level would.
        self.assertEqual(tracked_deleted_text(document_root), [" California"])
        # Every insertion is a whole word, never a lone letter (the char-level
        # signature). The token path inserts " England", " and", " Wales".
        for inserted in tracked_inserted_text(document_root):
            self.assertGreater(len(inserted.strip()), 1)

    def test_ai_clause_replace_still_uses_token_level_path(self):
        # An AI/clause redline (clause_id != manual_viewer_edit) keeps the token-level
        # path: a word-boundary edit redlines whole words, not individual characters.
        original = "The confidentiality obligations survive for seven years."
        replacement = "The confidentiality obligations survive for ten years."
        source_docx = make_source_docx([original])
        review_result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "c1",
                    "clause_id": "term",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraph = document_root.find(".//w:body/w:p", W_NS)
        self.assertEqual(revision_text_for_state(paragraph, accepted=False), original)
        self.assertEqual(revision_text_for_state(paragraph, accepted=True), replacement)
        # Token-level: the whole word "seven" is struck and the whole word "ten"
        # inserted -- a char-level diff would have kept the shared "en".
        self.assertTrue(any("seven" in text for text in tracked_deleted_text(document_root)))
        self.assertTrue(any("ten" in text for text in tracked_inserted_text(document_root)))

    def test_source_docx_export_preserves_multiple_redlines_on_same_source_paragraph(self):
        original = "This Agreement is governed by California. Confidentiality survives for seven years."
        governing_law_replacement = "This Agreement is governed by the laws of England and Wales."
        term_replacement = "Confidentiality survives for a fixed period of up to five years."
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 2,
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": governing_law_replacement,
                },
                {
                    "id": "r2",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": term_replacement,
                },
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in paragraphs
        ]
        self.assertIn((original, governing_law_replacement), states)
        self.assertIn((original, term_replacement), states)
        self.assertEqual(sum(1 for rejected, _accepted in states if rejected == original), 2)
        assert_track_changes_contract(self, redlined_docx, review_result["redline_edits"])

    def test_source_docx_export_real_split_block_paragraph_preserves_both_blocks(self):
        # The REAL shape (the fabricated [1,2,2] source_index fixture can never
        # occur: the extractor assigns source_index strictly by physical <w:p>
        # ordinal). ONE physical <w:p> holds two logical blocks separated by a hard
        # blank line; align_document_paragraphs re-splits it into two review
        # paragraphs that SHARE that physical paragraph's source_index. Built
        # end-to-end through the real extractor + aligner so the regression matches
        # production. Before the fix this either hard-failed ("could not anchor 1
        # approved redline") or rebuilt the whole <w:p> from one block's text,
        # silently destroying the other block.
        source_docx = make_source_docx_with_internal_blank_line_paragraph(
            ["Alpha clause."], ["Beta one.", "Beta two."]
        )
        extracted = extract_docx_paragraphs(source_docx)
        # The extractor yields ONE physical paragraph for the two blocks (its text
        # carries the internal blank line) -- never two paragraphs with one
        # source_index each.
        self.assertEqual(extracted[-1]["text"], "Beta one.\n\nBeta two.")
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        aligned = align_document_paragraphs(extracted, source_text)
        aligned_by_text = {str(paragraph["text"]): paragraph for paragraph in aligned}
        beta_one = aligned_by_text["Beta one."]
        beta_two = aligned_by_text["Beta two."]
        # Two distinct review paragraphs, shared provenance source_index.
        self.assertNotEqual(beta_one["id"], beta_two["id"])
        self.assertEqual(beta_one["source_index"], beta_two["source_index"])

        review_result = {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 2,
            "checked_at": "2026-06-01T00:00:00+00:00",
            "extracted_text": source_text,
            "paragraphs": aligned,
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": beta_one["id"],
                    "paragraph_index": beta_one["index"],
                    "source_index": beta_one["source_index"],
                    "original_text": "Beta one.",
                    "replacement_text": "First requirement satisfied.",
                },
                {
                    "id": "r2",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": beta_two["id"],
                    "paragraph_index": beta_two["index"],
                    "source_index": beta_two["source_index"],
                    "original_text": "Beta two.",
                    "replacement_text": "Second requirement satisfied.",
                },
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in paragraphs
        ]
        accepted_texts = [accepted for _rejected, accepted in states]
        # Each block is redlined on its own tracked paragraph -- no block clobbered.
        self.assertIn(("Beta one.", "First requirement satisfied."), states)
        self.assertIn(("Beta two.", "Second requirement satisfied."), states)
        # ...and neither original block survives un-redlined (the data-loss symptom).
        self.assertNotIn("Beta one.", accepted_texts)
        self.assertNotIn("Beta two.", accepted_texts)
        # Alpha + the two re-emitted block paragraphs: no block dropped, no spurious
        # duplicate inserted from two redlines piling onto one paragraph.
        self.assertEqual(len(paragraphs), 3)
        # The content-coverage gate (now always run on the direct-upload path)
        # agrees the sequence is intact.
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )
        assert_track_changes_contract(self, redlined_docx, review_result["redline_edits"])

    def test_source_docx_export_split_block_preserves_unedited_sibling_block(self):
        # Only ONE block of a split-block physical <w:p> is redlined. The sibling
        # block must survive verbatim -- the rebuild-from-one-block bug used to drop
        # it entirely.
        source_docx = make_source_docx_with_internal_blank_line_paragraph(
            ["Alpha clause."], ["Beta one.", "Beta two."]
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        aligned = align_document_paragraphs(extracted, source_text)
        beta_two = {str(p["text"]): p for p in aligned}["Beta two."]
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "extracted_text": source_text,
            "paragraphs": aligned,
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r2",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": beta_two["id"],
                    "paragraph_index": beta_two["index"],
                    "source_index": beta_two["source_index"],
                    "original_text": "Beta two.",
                    "replacement_text": "Second requirement satisfied.",
                },
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        accepted_all = "\n".join(
            revision_text_for_state(paragraph, accepted=True) for paragraph in paragraphs
        )
        # The untouched block survives verbatim; the edited block is redlined.
        self.assertIn("Beta one.", accepted_all)
        self.assertIn("Second requirement satisfied.", accepted_all)
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )
        assert_track_changes_contract(self, redlined_docx, review_result["redline_edits"])

    def test_source_docx_export_rejects_suspicious_compression_ratio(self):
        source_docx = make_source_docx(["A" * 4096])

        with patch.object(docx_text, "MAX_DOCX_ENTRY_COMPRESSION_RATIO", 2):
            with self.assertRaises(DocxExportError):
                build_source_redline_docx(source_docx, {"paragraphs": [], "redline_edits": []})

    def test_source_docx_export_rejects_suspicious_compression_ratio_before_zipfile_open(self):
        source_docx = make_source_docx(["A" * 4096])

        with (
            patch.object(docx_text, "MAX_DOCX_ENTRY_COMPRESSION_RATIO", 2),
            patch.object(docx_export, "ZipFile", side_effect=AssertionError("ZipFile should not open")),
        ):
            with self.assertRaises(DocxExportError):
                build_source_redline_docx(source_docx, {"paragraphs": [], "redline_edits": []})

    def test_source_docx_export_rejects_xml_dtd_entity_declarations(self):
        source_docx = replace_docx_parts(
            make_source_docx(["Safe body text."]),
            {"word/document.xml": unsafe_document_xml()},
        )

        with self.assertRaises(DocxExportError):
            build_source_redline_docx(source_docx, {"paragraphs": [], "redline_edits": []})

    def test_source_docx_export_prefers_text_anchors_over_stale_source_index(self):
        source_docx = make_source_docx([
            "Intro paragraph.",
            "Insert after this paragraph.",
            "This Agreement shall be governed by the laws of California.",
            "The Recipient must not circumvent the Company.",
        ])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 3,
            "checked_at": "2026-05-31T00:00:00+00:00",
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": "Intro paragraph."},
                {"id": "p2", "index": 2, "source_index": 2, "text": "Insert after this paragraph."},
                {
                    "id": "p3",
                    "index": 3,
                    "source_index": 3,
                    "text": "This Agreement shall be governed by the laws of California.",
                },
                {"id": "p4", "index": 4, "source_index": 4, "text": "The Recipient must not circumvent the Company."},
            ],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p3",
                    "source_index": 1,
                    "original_text": "This Agreement shall be governed by the laws of California.",
                    "replacement_text": "This Agreement shall be governed by the laws of England and Wales.",
                },
                {
                    "id": "r2",
                    "action": REDLINE_INSERT_AFTER_PARAGRAPH,
                    "paragraph_id": "p2",
                    "source_index": 1,
                    "anchor_text": "Insert after this paragraph.",
                    "insert_text": "New required clause.",
                },
                {
                    "id": "r3",
                    "action": REDLINE_DELETE_PARAGRAPH,
                    "paragraph_id": "p4",
                    "source_index": 1,
                    "original_text": "The Recipient must not circumvent the Company.",
                },
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in paragraphs
        ]
        self.assertIn(("Intro paragraph.", "Intro paragraph."), states)
        self.assertIn(("Insert after this paragraph.", "Insert after this paragraph."), states)
        self.assertIn(("", "New required clause."), states)
        self.assertIn(
            (
                "This Agreement shall be governed by the laws of California.",
                "This Agreement shall be governed by the laws of England and Wales.",
            ),
            states,
        )
        self.assertIn(("The Recipient must not circumvent the Company.", ""), states)

    def test_source_docx_export_rejects_ambiguous_text_anchor_without_source_index(self):
        source_docx = make_source_docx([
            "Duplicate paragraph.",
            "Duplicate paragraph.",
        ])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 1,
            "checked_at": "2026-05-31T00:00:00+00:00",
            "paragraphs": [
                {"id": "p1", "index": 1, "text": "Duplicate paragraph."},
            ],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "original_text": "Duplicate paragraph.",
                    "replacement_text": "Changed paragraph.",
                },
            ],
        }

        with self.assertRaisesRegex(DocxExportError, "could not anchor 1 approved redline"):
            build_source_redline_docx(source_docx, review_result)

    def _supplemental_part_review_result(self):
        return {
            "overall_status": "does_not_meet_requirements",
            "requirements_passed": 0,
            "requirements_failed": 1,
            "checked_at": "2026-05-31T00:00:00+00:00",
            "paragraphs": [
                {
                    "id": "p1",
                    "index": 1,
                    "source_part": "header1",
                    "text": "This Agreement shall be governed by the laws of California.",
                },
            ],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_part": "header1",
                    "original_text": "This Agreement shall be governed by the laws of California.",
                    "replacement_text": "This Agreement shall be governed by the laws of England and Wales.",
                },
            ],
        }

    def test_source_docx_export_fails_closed_on_supplemental_part_redline(self):
        # A header/footer paragraph is extracted and reviewed, so an APPROVED redline
        # can target it -- but this body-only export cannot write header1.xml/footer1.xml.
        # The OLD behaviour silently logged-and-dropped it while reporting success (the
        # header/footer analogue of the PDF silent-drop P0). Strict (the default for
        # send/approve/export) must now FAIL CLOSED instead.
        source_docx = make_source_docx([
            "This Agreement shall be governed by the laws of California.",
        ])
        review_result = self._supplemental_part_review_result()

        with self.assertRaises(SupplementalRedlineUnavailableError) as ctx:
            build_source_redline_docx(source_docx, review_result)
        self.assertEqual(ctx.exception.count, 1)
        self.assertIsInstance(ctx.exception, DocxExportError)

    def test_source_docx_export_reports_supplemental_redline_incomplete_lenient(self):
        # Lenient (preview/draft/diagnostic): the file is still produced, but the
        # unapplied header redline is surfaced as incomplete -- never silently dropped
        # under a clean/successful package.
        source_docx = make_source_docx([
            "This Agreement shall be governed by the laws of California.",
        ])
        review_result = self._supplemental_part_review_result()

        package = source_redline_docx.build_source_redline_package(
            source_docx, review_result, strict=False
        )

        self.assertEqual(len(package.anchor_uncertain_redlines), 1)
        self.assertEqual(package.anchor_uncertain_redlines[0].get("id"), "r1")
        assert_docx_package_healthy(self, package.data)
        _settings_root, document_root, _document_xml = docx_xml_roots(package.data)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        # Body untouched (the header redline did not land in the body), but the package
        # is explicitly flagged incomplete rather than presented as clean success.
        self.assertEqual(
            [
                (
                    revision_text_for_state(paragraph, accepted=False),
                    revision_text_for_state(paragraph, accepted=True),
                )
                for paragraph in paragraphs
            ],
            [
                (
                    "This Agreement shall be governed by the laws of California.",
                    "This Agreement shall be governed by the laws of California.",
                ),
            ],
        )

    def test_source_docx_export_repairs_missing_package_relationships(self):
        source_docx = make_source_docx(
            ["This Agreement shall be governed by the laws of California."],
            include_package_rels=False,
        )
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)

    def test_source_docx_export_repairs_malformed_required_package_metadata(self):
        source_docx = make_source_docx([
            "This Agreement shall be governed by the laws of California.",
        ])
        source_docx = replace_docx_parts(source_docx, {
            "[Content_Types].xml": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/xml"/>
  <Default Extension="xml" ContentType="text/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/xml"/>
  <Override PartName="/word/settings.xml" ContentType="application/xml"/>
</Types>""",
            "_rels/.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/missing-document.xml"/>
</Relationships>""",
            "word/_rels/document.xml.rels": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="missing-settings.xml"/>
</Relationships>""",
            "word/settings.xml": """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:trackRevisions w:val="0"/>
</w:settings>""",
        })
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        defaults, overrides = docx_content_types(redlined_docx)
        self.assertEqual(defaults["rels"], RELATIONSHIPS_CONTENT_TYPE)
        self.assertEqual(defaults["xml"], "application/xml")
        self.assertEqual(overrides["/word/document.xml"], DOCUMENT_CONTENT_TYPE)
        self.assertEqual(overrides["/word/settings.xml"], SETTINGS_CONTENT_TYPE)
        with ZipFile(BytesIO(redlined_docx)) as archive:
            package_relationships = relationship_targets(archive, "_rels/.rels")
        office_document_targets = [
            resolve_relationship_target("_rels/.rels", relationship["Target"])
            for relationship in package_relationships
            if relationship.get("Type") == OFFICE_DOCUMENT_RELATIONSHIP_TYPE
        ]
        self.assertEqual(office_document_targets, ["word/document.xml"])
        self.assertEqual(docx_document_relationship_targets(redlined_docx)[SETTINGS_RELATIONSHIP_TYPE], "settings.xml")

    def test_source_docx_export_adds_missing_section_properties(self):
        source_docx = make_source_docx([
            "This Agreement shall be governed by the laws of California.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        result = review_nda("\n\n".join(str(paragraph["text"]) for paragraph in paragraphs), paragraphs=paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        section = document_root.find(".//w:body/w:sectPr", W_NS)
        self.assertIsNotNone(section)
        page_size = section.find("w:pgSz", W_NS)
        self.assertIsNotNone(page_size)
        self.assertEqual(page_size.get(f"{{{W_NS['w']}}}w"), A4_PAGE_WIDTH_TWIPS)
        self.assertEqual(page_size.get(f"{{{W_NS['w']}}}h"), A4_PAGE_HEIGHT_TWIPS)

    def test_source_docx_export_marks_paragraph_deletes_and_insertions(self):
        source_docx = make_source_docx([
            "The parties will discuss a possible transaction.",
            "The Recipient must not circumvent the Company or deal directly with introduced parties.",
        ])
        paragraphs = extract_docx_paragraphs(source_docx)
        # Exercise the delete + insert source-export mechanism directly via redline_edits
        # (a delete for the prohibited paragraph, an insert for the governing-law gap),
        # independent of which engine produced them.
        result = _delete_and_insert_review_result(paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_source_export_has_no_report_leakage(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        deleted_text = tracked_deleted_text(document_root)
        inserted_text = tracked_inserted_text(document_root)
        self.assertTrue(any("must not circumvent" in text for text in deleted_text))
        self.assertTrue(any("This Agreement shall be governed by the laws of England and Wales." in text for text in inserted_text))
        self.assertEqual(paragraph_property_revisions(document_root, "del"), [])
        self.assertEqual(paragraph_property_revisions(document_root, "ins"), [])

    def test_source_docx_export_strips_source_paragraph_property_revisions(self):
        source_docx = make_source_docx([
            "The parties will discuss a possible transaction.",
            "The Recipient must not circumvent the Company or deal directly with introduced parties.",
        ])
        with ZipFile(BytesIO(source_docx)) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        document_xml = document_xml.replace(
            "<w:p><w:r><w:t>The parties will discuss a possible transaction.</w:t></w:r></w:p>",
            '<w:p><w:pPr><w:rPr><w:ins w:id="97" w:author="source" w:date="2026-06-01T00:00:00Z" />'
            '<w:del w:id="98" w:author="source" w:date="2026-06-01T00:00:00Z" /></w:rPr></w:pPr>'
            "<w:r><w:t>The parties will discuss a possible transaction.</w:t></w:r></w:p>",
        )
        source_docx = replace_docx_parts(source_docx, {"word/document.xml": document_xml})
        paragraphs = extract_docx_paragraphs(source_docx)
        result = _delete_and_insert_review_result(paragraphs)

        redlined_docx = build_source_redline_docx(source_docx, result)

        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        self.assertEqual(paragraph_property_revisions(document_root, "del"), [])
        self.assertEqual(paragraph_property_revisions(document_root, "ins"), [])
        self.assertTrue(any("must not circumvent" in text for text in tracked_deleted_text(document_root)))
        self.assertTrue(any("England and Wales" in text for text in tracked_inserted_text(document_root)))

    def test_strip_paragraph_property_revisions_handles_paragraph_properties_root(self):
        paragraph_properties = ET.fromstring(
            '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:rPr><w:b/><w:ins w:id="97" w:author="source" w:date="2026-06-01T00:00:00Z" />'
            '<w:del w:id="98" w:author="source" w:date="2026-06-01T00:00:00Z" />'
            '<w:rPrChange w:id="99" w:author="source" w:date="2026-06-01T00:00:00Z">'
            '<w:rPr><w:i/></w:rPr></w:rPrChange></w:rPr>'
            '<w:pPrChange w:id="100" w:author="source" w:date="2026-06-01T00:00:00Z">'
            '<w:pPr><w:spacing w:after="120"/></w:pPr></w:pPrChange>'
            "</w:pPr>"
        )

        _strip_paragraph_property_revisions(paragraph_properties)

        self.assertEqual(paragraph_properties.findall(".//w:ins", W_NS), [])
        self.assertEqual(paragraph_properties.findall(".//w:del", W_NS), [])
        self.assertEqual(paragraph_properties.findall(".//w:pPrChange", W_NS), [])
        self.assertEqual(paragraph_properties.findall(".//w:rPrChange", W_NS), [])
        self.assertIsNotNone(paragraph_properties.find(".//w:b", W_NS))

    def test_source_docx_export_matches_redline_actions_at_paragraph_level(self):
        source_docx = make_source_docx([
            "This Agreement shall be governed by the laws of California.",
            "Insert after this paragraph.",
            "The Recipient must not circumvent the Company.",
        ])
        review_result = track_changes_contract_review_result()

        redlined_docx = build_source_redline_docx(source_docx, review_result)

        assert_source_export_has_no_report_leakage(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in paragraphs
        ]
        self.assertIn(
            (
                "This Agreement shall be governed by the laws of California.",
                "This Agreement shall be governed by the laws of England and Wales.",
            ),
            states,
        )
        self.assertIn(("Insert after this paragraph.", "Insert after this paragraph."), states)
        self.assertIn(("", "New required clause."), states)
        self.assertIn(("The Recipient must not circumvent the Company.", ""), states)
        assert_track_changes_contract(self, redlined_docx, review_result["redline_edits"])

    def test_review_report_docx_marks_delete_paragraph_redlines(self):
        # A delete_paragraph redline edit renders as a tracked deletion. (non_circumvention
        # is now a dynamic clause; its delete redline arrives via redline_edits, which is
        # what the report renderer consumes — exercise that mechanism directly.)
        offending = "The Recipient must not circumvent the Company or deal directly with introduced parties."
        result = {
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": offending}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "clause_name": "Non-Circumvention",
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "action": REDLINE_DELETE_PARAGRAPH,
                    "original_text": offending,
                    "replacement_text": "",
                }
            ],
        }

        _settings_root, document_root, _document_xml = docx_xml_roots(build_review_report_docx(result))

        deleted_text = tracked_deleted_text(document_root)
        self.assertTrue(any(offending in text for text in deleted_text))
        self.assertEqual(paragraph_property_revisions(document_root, "del"), [])

    def test_review_report_docx_marks_insert_after_paragraph_redlines(self):
        result = review_nda("The parties will discuss a possible transaction.")

        _settings_root, document_root, _document_xml = docx_xml_roots(build_review_report_docx(result))

        inserted_text = tracked_inserted_text(document_root)
        self.assertTrue(any("This Agreement shall be governed by the laws of England and Wales." in text for text in inserted_text))
        self.assertEqual(paragraph_property_revisions(document_root, "ins"), [])

    def test_docx_open_health_rejects_revision_markup_inside_paragraph_properties(self):
        result = review_nda("The parties will discuss a possible transaction.")
        docx_bytes = build_review_report_docx(result)
        with ZipFile(BytesIO(docx_bytes), "r") as source_archive:
            document_xml = source_archive.read("word/document.xml").decode("utf-8")
            document_xml = document_xml.replace(
                "<w:p><w:ins ",
                '<w:p><w:pPr><w:rPr><w:ins w:id="999" w:author="bad" w:date="2026-05-31T00:00:00Z" /></w:rPr></w:pPr><w:ins ',
                1,
            )
            with BytesIO() as output:
                with ZipFile(output, "w", ZIP_DEFLATED) as patched_archive:
                    for item in source_archive.infolist():
                        data = document_xml.encode("utf-8") if item.filename == "word/document.xml" else source_archive.read(item.filename)
                        patched_archive.writestr(item, data)
                patched_docx = output.getvalue()

        errors = validate_docx_open_health(patched_docx, require_styles=True)

        self.assertIn("document.xml contains insertion revision markup inside paragraph properties.", errors)

    def test_docx_open_health_rejects_xml_dtd_entity_declarations(self):
        docx_bytes = replace_docx_parts(
            build_review_report_docx(review_nda("The parties will discuss a possible transaction.")),
            {"word/document.xml": unsafe_document_xml()},
        )

        errors = validate_docx_open_health(docx_bytes, require_styles=True)

        self.assertTrue(
            any("unsupported XML DTD/entity declarations" in error for error in errors),
            errors,
        )

    def test_docx_open_health_rejects_excessive_uncompressed_docx_size_before_integrity_check(self):
        oversized_docx = make_source_docx(["A" * 4096])

        with (
            patch.object(docx_text, "MAX_DOCX_UNCOMPRESSED_BYTES", 1024),
            patch.object(ZipFile, "testzip", side_effect=AssertionError("integrity check should not run")),
            patch.object(docx_health, "parse_docx_xml", side_effect=AssertionError("XML parse should not run")),
        ):
            errors = validate_docx_open_health(oversized_docx)

        self.assertEqual(errors, ["The Word document is too large after decompression."])

    def test_docx_open_health_rejects_suspicious_compression_ratio_before_zipfile_open(self):
        suspicious_docx = make_source_docx(["A" * 4096])

        with (
            patch.object(docx_text, "MAX_DOCX_ENTRY_COMPRESSION_RATIO", 2),
            patch.object(docx_health, "ZipFile", side_effect=AssertionError("ZipFile should not open")),
        ):
            errors = validate_docx_open_health(suspicious_docx)

        self.assertEqual(errors, ["The Word document uses a suspicious compression ratio."])

    def test_review_report_docx_preserves_track_changes_contract_by_redline_action(self):
        result = track_changes_contract_review_result()

        report_docx = build_review_report_docx(result)

        inspection = assert_track_changes_contract(self, report_docx, result["redline_edits"])
        self.assertEqual(inspection.summary["expected_replace"], 1)
        self.assertEqual(inspection.summary["expected_insert"], 1)
        self.assertEqual(inspection.summary["expected_delete"], 1)

    def test_redline_contract_inspector_rejects_accepted_only_exports(self):
        result = track_changes_contract_review_result()
        accepted_only_result = {**result, "redline_edits": []}

        report_docx = build_review_report_docx(accepted_only_result)
        inspection = inspect_docx_redline_contract(report_docx, result["redline_edits"])

        self.assertTrue(any("replace did not preserve" in error for error in inspection.errors))
        self.assertTrue(any("insert did not preserve" in error for error in inspection.errors))
        self.assertTrue(any("delete did not preserve" in error for error in inspection.errors))

    def test_review_report_docx_splits_multi_paragraph_insertions(self):
        result = review_nda(
            "This Agreement shall be governed by the laws of the DIFC.\n\n"
            "The confidentiality obligations survive for three (3) years."
        )

        _settings_root, document_root, _document_xml = docx_xml_roots(build_review_report_docx(result))

        inserted_text = tracked_inserted_text(document_root)
        self.assertTrue(any("For [Party 1 legal name]" in text for text in inserted_text))
        self.assertTrue(any("For [Party 2 legal name]" in text for text in inserted_text))
        self.assertGreaterEqual(len([text for text in inserted_text if text.startswith("For [Party")]), 2)
        self.assertEqual(paragraph_mark_revisions(document_root, "ins"), [])

    def _build_run_format_docx(self, formatted_text, format_ops, source_rpr=""):
        # Shared scaffold for the underline/strike/color/highlight round-trip tests:
        # builds a one-paragraph source doc, runs the format_ops through the tracked
        # source-redline export, and returns (redlined_bytes, source_text). The
        # ``source_rpr`` lets a test seed a pre-existing rPr to stress schema ordering.
        formatted_paragraph = (
            f"<w:p><w:r>{source_rpr}<w:t>{formatted_text}</w:t></w:r></w:p>"
        )
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{formatted_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(paragraph["text"]) for paragraph in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        formatted_id = next(
            str(paragraph["id"])
            for paragraph in paragraphs
            if str(paragraph["text"]) == formatted_text
        )
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "fmt-roundtrip-1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_FORMAT_PARAGRAPH,
                    "action_label": "Format paragraph",
                    "paragraph_id": formatted_id,
                    "original_text": formatted_text,
                    "replacement_text": formatted_text,
                    "format_ops": format_ops,
                }
            ],
        }
        return build_source_redline_docx(source_docx, review_result), source_text

    def test_run_underline_strike_color_highlight_round_trip(self):
        # The hard gate: a single paragraph carrying underline + strike + text color +
        # named highlight runs must export to a DOCX that (1) opens cleanly under
        # python-docx, (2) keeps its rPr children in schema order, and (3) serialises
        # each property correctly (u/strike present, color is a 6-hex, highlight is a
        # NAMED palette value -- never a raw hex).
        formatted_text = "This Agreement is governed by the laws of California."
        u_start = formatted_text.index("Agreement")
        u_end = u_start + len("Agreement")
        strike_start = formatted_text.index("governed")
        strike_end = strike_start + len("governed")
        color_start = formatted_text.index("laws")
        color_end = color_start + len("laws")
        hl_start = formatted_text.index("California")
        hl_end = hl_start + len("California")
        format_ops = [
            {"scope": "run", "property": "underline", "start": u_start, "end": u_end, "from": False, "to": True},
            {"scope": "run", "property": "strike", "start": strike_start, "end": strike_end, "from": False, "to": True},
            {"scope": "run", "property": "color", "start": color_start, "end": color_end, "from": "", "to": "FF0000"},
            {"scope": "run", "property": "highlight", "start": hl_start, "end": hl_end, "from": "", "to": "yellow"},
        ]
        redlined_docx, source_text = self._build_run_format_docx(formatted_text, format_ops)

        # (1) Package health + python-docx load (Document(path) must succeed).
        assert_docx_package_healthy(self, redlined_docx)
        from docx import Document as _Document

        document = _Document(BytesIO(redlined_docx))
        self.assertTrue(any(formatted_text == p.text for p in document.paragraphs))

        # Content coverage gate: the text is unchanged, only formatting moved.
        self.assertEqual(
            verify_export_content_coverage(redlined_docx, source_text), []
        )

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)

        def run_text(node):
            return "".join(t.text or "" for t in node.findall(".//w:t", W_NS))

        formatted = next(
            p for p in document_root.findall(".//w:p", W_NS) if run_text(p) == formatted_text
        )

        def covered_rpr(word):
            run = next(
                r
                for r in formatted.findall("w:r", W_NS)
                if "".join(t.text or "" for t in r.findall("w:t", W_NS)) == word
            )
            return run.find("w:rPr", W_NS)

        val_attr = f"{{{W_NS['w']}}}val"

        # (3a) underline -> <w:u w:val="single">.
        u_rpr = covered_rpr("Agreement")
        u_el = u_rpr.find("w:u", W_NS)
        self.assertIsNotNone(u_el)
        self.assertEqual(u_el.get(val_attr), "single")

        # (3b) strike -> <w:strike> present.
        strike_rpr = covered_rpr("governed")
        self.assertIsNotNone(strike_rpr.find("w:strike", W_NS))

        # (3c) color -> <w:color w:val="FF0000"> (6-hex, no leading #).
        color_rpr = covered_rpr("laws")
        color_el = color_rpr.find("w:color", W_NS)
        self.assertIsNotNone(color_el)
        self.assertEqual(color_el.get(val_attr), "FF0000")

        # (3d) highlight -> NAMED palette value, never a raw hex.
        highlight_rpr = covered_rpr("California")
        highlight_el = highlight_rpr.find("w:highlight", W_NS)
        self.assertIsNotNone(highlight_el)
        self.assertEqual(highlight_el.get(val_attr), "yellow")

        # (2) rPr child schema order across every covered run: the emitted children
        # must appear in CT_RPr sequence (strike < color < sz < szCs < highlight < u).
        canonical = ["rFonts", "b", "bCs", "i", "iCs", "strike", "dstrike",
                     "color", "sz", "szCs", "highlight", "u", "rPrChange"]
        order_index = {name: i for i, name in enumerate(canonical)}
        for rpr in formatted.findall(".//w:rPr", W_NS):
            local_names = [child.tag.split("}", 1)[-1] for child in list(rpr)]
            ranked = [order_index[name] for name in local_names if name in order_index]
            self.assertEqual(ranked, sorted(ranked), f"rPr out of schema order: {local_names}")

    def test_run_format_inserts_before_preexisting_later_rpr_children(self):
        # Stress the schema-ordering insert: the source run already carries <w:u> and
        # <w:highlight> (both LATER than strike/color in CT_RPr). Adding strike + color
        # must place them BEFORE the existing u/highlight, not append them out of order.
        formatted_text = "Confidential information is defined herein."
        source_rpr = '<w:rPr><w:u w:val="single"/><w:highlight w:val="green"/></w:rPr>'
        format_ops = [
            {"scope": "run", "property": "strike", "start": 0, "end": len(formatted_text), "from": False, "to": True},
            {"scope": "run", "property": "color", "start": 0, "end": len(formatted_text), "from": "", "to": "0000FF"},
        ]
        redlined_docx, _source_text = self._build_run_format_docx(
            formatted_text, format_ops, source_rpr=source_rpr
        )
        assert_docx_package_healthy(self, redlined_docx)
        from docx import Document as _Document

        _Document(BytesIO(redlined_docx))  # must not raise

        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        formatted = next(
            p
            for p in document_root.findall(".//w:p", W_NS)
            if "".join(t.text or "" for t in p.findall(".//w:t", W_NS)) == formatted_text
        )
        canonical = ["rFonts", "b", "bCs", "i", "iCs", "strike", "dstrike",
                     "color", "sz", "szCs", "highlight", "u", "rPrChange"]
        order_index = {name: i for i, name in enumerate(canonical)}
        covered = next(r for r in formatted.findall("w:r", W_NS) if r.find("w:rPr", W_NS) is not None)
        rpr = covered.find("w:rPr", W_NS)
        local_names = [child.tag.split("}", 1)[-1] for child in list(rpr)]
        # All four props present and in canonical order: strike < color < highlight < u.
        for name in ("strike", "color", "highlight", "u"):
            self.assertIn(name, local_names, f"{name} missing from rPr")
        ranked = [order_index[name] for name in local_names if name in order_index]
        self.assertEqual(ranked, sorted(ranked), f"rPr out of schema order: {local_names}")

    def test_replacement_runs_preserve_underline_strike_color_highlight(self):
        # The replace path: a manual replace redline carries replacement_runs with the
        # four new props. The clean (non-tracked) export must re-emit the inserted text
        # WITH each property, opening cleanly under python-docx.
        original_text = "Air India Limited is the FIRST PARTY."
        edited_text = "Air India Private Limited is the FIRST PARTY."
        body_paragraph = f"<w:p><w:r><w:t>{original_text}</w:t></w:r></w:p>"
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:body><w:p><w:r><w:t>Intro paragraph.</w:t></w:r></w:p>"
            f"{body_paragraph}</w:body></w:document>"
        )
        source_docx = replace_docx_parts(
            make_source_docx(["placeholder one", "placeholder two"]),
            {"word/document.xml": document_xml},
        )
        extracted = extract_docx_paragraphs(source_docx)
        source_text = "\n\n".join(str(p["text"]) for p in extracted)
        paragraphs = align_document_paragraphs(extracted, source_text)
        pid = next(str(p["id"]) for p in paragraphs if str(p["text"]) == original_text)
        review_result = {
            "paragraphs": paragraphs,
            "redline_edits": [
                {
                    "id": "m1",
                    "clause_id": "manual_viewer_edit",
                    "status": "proposed",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "action_label": "Your edit",
                    "is_manual": True,
                    "whole_paragraph": False,
                    "paragraph_id": pid,
                    "original_text": original_text,
                    "replacement_text": edited_text,
                    "replacement_runs": [
                        {"text": "Air India "},
                        {"text": "Private", "underline": True, "color": "FF0000"},
                        {"text": " Limited is the "},
                        {"text": "FIRST PARTY", "strike": True, "highlight": "cyan"},
                        {"text": "."},
                    ],
                }
            ],
        }

        redlined_docx = build_source_redline_docx(source_docx, review_result)
        clean_docx = accept_all_revisions(redlined_docx)
        # A clean (accepted) doc deliberately has Track Changes OFF, so the redline
        # health check doesn't apply; assert ZIP + XML integrity directly, then load
        # under python-docx as the "opens without corruption" gate.
        with ZipFile(BytesIO(clean_docx)) as archive:
            self.assertIsNone(archive.testzip())
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))
        from docx import Document as _Document

        document = _Document(BytesIO(clean_docx))
        self.assertTrue(any(edited_text == p.text for p in document.paragraphs))

        _settings_root, document_root, _document_xml = docx_xml_roots(clean_docx)
        formatted = next(
            p
            for p in document_root.findall(".//w:p", W_NS)
            if "".join(t.text or "" for t in p.findall(".//w:t", W_NS)) == edited_text
        )
        val_attr = f"{{{W_NS['w']}}}val"

        def run_for(word):
            return next(
                r
                for r in formatted.findall("w:r", W_NS)
                if "".join(t.text or "" for t in r.findall("w:t", W_NS)) == word
            )

        underline_rpr = run_for("Private").find("w:rPr", W_NS)
        self.assertIsNotNone(underline_rpr.find("w:u", W_NS))
        self.assertEqual(underline_rpr.find("w:color", W_NS).get(val_attr), "FF0000")
        strike_rpr = run_for("FIRST PARTY").find("w:rPr", W_NS)
        self.assertIsNotNone(strike_rpr.find("w:strike", W_NS))
        self.assertEqual(strike_rpr.find("w:highlight", W_NS).get(val_attr), "cyan")

        # Schema order on every rPr in the rebuilt paragraph.
        canonical = ["rFonts", "b", "bCs", "i", "iCs", "strike", "dstrike",
                     "color", "sz", "szCs", "highlight", "u", "rPrChange"]
        order_index = {name: i for i, name in enumerate(canonical)}
        for rpr in formatted.findall(".//w:rPr", W_NS):
            local_names = [child.tag.split("}", 1)[-1] for child in list(rpr)]
            ranked = [order_index[name] for name in local_names if name in order_index]
            self.assertEqual(ranked, sorted(ranked), f"rPr out of schema order: {local_names}")

    # ------------------------------------------------------------------ #
    # Category A: multi-edit-per-clause export + span-level strikes + the
    # same-<w:p> collision fail-safe (U3) and coverage-gate reconciliation (U4).
    # ------------------------------------------------------------------ #

    def test_category_a_multi_edit_per_clause_applies_every_edit(self):
        # A single Category-A clause emits MULTIPLE edits, one per defective span,
        # each targeting a DISTINCT paragraph. All must land as tracked changes --
        # the multi-edit application loop already supports this; this pins it.
        para_one = "The Receiving Party shall not solicit any employee of the Disclosing Party."
        para_two = "The Receiving Party shall pay a penalty of $50,000 for each breach."
        # Span-derived replacements (each strikes a prohibited span and rewrites it):
        replace_one = "The Receiving Party shall protect the confidential information of the Disclosing Party."
        replace_two = "The Receiving Party shall be liable for actual damages for each breach."
        source_docx = make_source_docx([para_one, para_two])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": para_one},
                {"id": "p2", "index": 2, "source_index": 2, "text": para_two},
            ],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": para_one,
                    "replacement_text": replace_one,
                },
                {
                    "id": "r2",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p2",
                    "paragraph_index": 2,
                    "source_index": 2,
                    "original_text": para_two,
                    "replacement_text": replace_two,
                },
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        paragraphs = document_root.findall(".//w:body/w:p", W_NS)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in paragraphs
        ]
        self.assertIn((para_one, replace_one), states)
        self.assertIn((para_two, replace_two), states)
        assert_track_changes_contract(self, redlined_docx, review_result["redline_edits"])

    def test_category_a_span_replace_renders_word_level_strike_not_whole_paragraph(self):
        # A span-derived replace_paragraph (replacement = paragraph minus a prohibited
        # span, no internal newline) MUST diff to a surgical sub-span strike, NOT a
        # whole-paragraph del/ins. This is the cheap span mechanism: the export reuses
        # the audited word-diff path, so only the struck words carry w:del markup.
        original = (
            "The Receiving Party shall keep the information confidential and shall not "
            "solicit any employee of the Disclosing Party for two years."
        )
        # Strike only the prohibited non-solicit span; the rest survives verbatim.
        replacement = (
            "The Receiving Party shall keep the information confidential for two years."
        )
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": replacement,
                }
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        target = document_root.findall(".//w:body/w:p", W_NS)[0]
        # The deleted markup carries ONLY the struck span's words, not the whole
        # paragraph -- proving a surgical strike, not a whole-paragraph del.
        deleted_text = revision_text_for_state(
            next(child for child in target if child.tag.endswith("}del")), accepted=False
        ) if any(child.tag.endswith("}del") for child in target) else "".join(
            revision_text_for_state(d, accepted=False)
            for d in target.findall("w:del", W_NS)
        )
        self.assertIn("solicit", deleted_text)
        self.assertNotIn("keep", deleted_text)
        self.assertNotIn("confidential", deleted_text)
        # Some original words survive OUTSIDE any revision run (untracked context),
        # which a whole-paragraph del/ins would not leave.
        plain_runs = [
            r for r in target.findall("w:r", W_NS)
            if not any(
                parent.tag.endswith("}del") or parent.tag.endswith("}ins")
                for parent in [target]
            )
        ]
        self.assertTrue(plain_runs, "span replace should leave untracked context runs")
        # Accepted view == the replacement; rejected view == the original.
        self.assertEqual(revision_text_for_state(target, accepted=True), replacement)
        self.assertEqual(revision_text_for_state(target, accepted=False), original)

    def test_category_a_same_paragraph_collision_fails_safe(self):
        # Two NON-INSERTION edits of the SAME clause resolving to ONE physical <w:p>
        # is a residual uncoalesced collision (the engine should have folded the
        # clause's same-paragraph spans into one replace). Applying both would have
        # the second clobber the first -- silent redline loss. The export must fail
        # SAFE (raise) rather than corrupt.
        original = "The Receiving Party shall not solicit employees and shall not compete."
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": "The Receiving Party shall not compete.",
                },
                {
                    "id": "r2",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": "The Receiving Party shall not solicit employees.",
                },
            ],
        }
        with self.assertRaises(DocxExportError) as caught:
            build_source_redline_docx(source_docx, review_result)
        self.assertIn("coalesced", str(caught.exception).lower())

    def test_category_a_distinct_clause_alternatives_on_same_paragraph_still_land(self):
        # Regression guard: two edits from DIFFERENT clauses presenting ALTERNATIVE
        # whole-paragraph replacements of one combined sentence is NOT a collision --
        # each lands on its own tracked paragraph (the historical behavior). The
        # same-clause collision guard must NOT trip on distinct clause_ids.
        original = "This Agreement is governed by California. Confidentiality survives for seven years."
        governing_law_replacement = "This Agreement is governed by the laws of England and Wales."
        term_replacement = "Confidentiality survives for a fixed period of up to five years."
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "paragraphs": [{"id": "p1", "index": 1, "source_index": 1, "text": original}],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "governing_law",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": governing_law_replacement,
                },
                {
                    "id": "r2",
                    "clause_id": "term_and_survival",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": term_replacement,
                },
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        states = [
            (
                revision_text_for_state(paragraph, accepted=False),
                revision_text_for_state(paragraph, accepted=True),
            )
            for paragraph in document_root.findall(".//w:body/w:p", W_NS)
        ]
        self.assertIn((original, governing_law_replacement), states)
        self.assertIn((original, term_replacement), states)

    def test_category_a_coverage_gate_reconciles_multiple_strikes(self):
        # The content-coverage gate reconciles MULTIPLE span-derived strike replaces,
        # one per distinct paragraph, in a single export.
        para_one = "The Receiving Party shall not solicit any employee for two years."
        para_two = "The Receiving Party shall not engage in any competing business."
        replace_one = "The Receiving Party shall not for two years."
        replace_two = "The Receiving Party shall."
        source_text = f"{para_one}\n\n{para_two}"
        source_docx = make_source_docx([para_one, para_two])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "extracted_text": source_text,
            "paragraphs": [
                {"id": "p1", "index": 1, "source_index": 1, "text": para_one},
                {"id": "p2", "index": 2, "source_index": 2, "text": para_two},
            ],
            "clauses": [],
            "redline_edits": [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p1",
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": para_one,
                    "replacement_text": replace_one,
                },
                {
                    "id": "r2",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_id": "p2",
                    "paragraph_index": 2,
                    "source_index": 2,
                    "original_text": para_two,
                    "replacement_text": replace_two,
                },
            ],
        }
        redlined_docx = build_source_redline_docx(source_docx, review_result)
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=review_result["redline_edits"],
            ),
            [],
        )

    def test_category_a_coverage_gate_catches_dropped_redline(self):
        # Fail-closed preserved: if an expected redline never lands in the export
        # bytes, the coverage gate must report it (never silently pass a drop).
        para_one = "The Receiving Party shall not solicit any employee for two years."
        para_two = "The Receiving Party shall not engage in any competing business."
        replace_one = "The Receiving Party shall not for two years."
        replace_two = "The Receiving Party shall."
        source_text = f"{para_one}\n\n{para_two}"
        # Build an export where ONLY r1 landed (r2 was dropped): the second paragraph
        # is the untouched original.
        source_docx = make_source_docx([para_one, para_two])
        landed_only_r1 = build_source_redline_docx(
            source_docx,
            {
                "overall_status": "does_not_meet_requirements",
                "checked_at": "2026-06-01T00:00:00+00:00",
                "extracted_text": source_text,
                "paragraphs": [
                    {"id": "p1", "index": 1, "source_index": 1, "text": para_one},
                    {"id": "p2", "index": 2, "source_index": 2, "text": para_two},
                ],
                "clauses": [],
                "redline_edits": [
                    {
                        "id": "r1",
                        "clause_id": "non_circumvention",
                        "action": REDLINE_REPLACE_PARAGRAPH,
                        "paragraph_id": "p1",
                        "paragraph_index": 1,
                        "source_index": 1,
                        "original_text": para_one,
                        "replacement_text": replace_one,
                    }
                ],
            },
        )
        # The gate is told BOTH r1 and r2 were expected; r2 never landed.
        expected_both = [
            {
                "id": "r1",
                "clause_id": "non_circumvention",
                "action": REDLINE_REPLACE_PARAGRAPH,
                "paragraph_id": "p1",
                "paragraph_index": 1,
                "source_index": 1,
                "original_text": para_one,
                "replacement_text": replace_one,
            },
            {
                "id": "r2",
                "clause_id": "non_circumvention",
                "action": REDLINE_REPLACE_PARAGRAPH,
                "paragraph_id": "p2",
                "paragraph_index": 2,
                "source_index": 2,
                "original_text": para_two,
                "replacement_text": replace_two,
            },
        ]
        errors = verify_export_content_coverage(
            landed_only_r1,
            source_text,
            expected_redline_edits=expected_both,
        )
        self.assertTrue(errors, "coverage gate must flag the dropped r2 redline")

    def test_category_a_coverage_gate_fails_closed_on_same_paragraph_collision(self):
        # U4 fail-closed: two destructive (replace/delete) edits keyed to the SAME
        # expected source paragraph means the clause's edits were NOT coalesced.
        # The prior behavior silently overwrote the expectation with the second edit,
        # which could pass an export that dropped the first. The gate must now report
        # the collision rather than reconcile against one survivor.
        original = "The Receiving Party shall not solicit employees and shall not compete."
        errors = docx_health._expected_accepted_source_paragraphs(
            [docx_health._normalize_export_text(original)],
            [
                {
                    "id": "r1",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": "The Receiving Party shall not compete.",
                },
                {
                    "id": "r2",
                    "clause_id": "non_circumvention",
                    "action": REDLINE_REPLACE_PARAGRAPH,
                    "paragraph_index": 1,
                    "source_index": 1,
                    "original_text": original,
                    "replacement_text": "The Receiving Party shall not solicit employees.",
                },
            ],
        )[1]
        self.assertTrue(errors, "two destructive edits on one paragraph must fail closed")
        self.assertIn("coalesced", " ".join(errors).lower())

    def test_category_a_pdf_coverage_no_redline_case_unchanged(self):
        # PDF-source coverage path is unchanged for the no-redline case: an empty /
        # no-op expected list is trivially covered (no false positive).
        from nda_automation.docx_health import verify_pdf_reconstruction_redline_coverage

        empty_docx = make_source_docx(["Some reconstructed body paragraph."])
        self.assertEqual(verify_pdf_reconstruction_redline_coverage(empty_docx, []), [])
        self.assertEqual(verify_pdf_reconstruction_redline_coverage(empty_docx, None), [])

    def test_category_a_cross_track_multi_span_coalesce_export_coverage_end_to_end(self):
        # CROSS-TRACK INTEGRATION GUARD (engine <-> export discriminator alignment).
        #
        # The load-bearing Category-A contract: a clause emits SEVERAL span edits on
        # ONE paragraph; the ENGINE (clause_outcomes) must COALESCE them into a SINGLE
        # replace_paragraph keyed on the SAME discriminator (clause_id, paragraph_id)
        # that the EXPORT's residual-collision guard keys on (clause_id per <w:p>). If
        # the engine coalesces correctly, exactly ONE edit per (clause, paragraph)
        # reaches the export, the collision guard is NOT tripped, both struck spans
        # land as tracked changes, and the coverage gate reconciles 1:1 and PASSES.
        #
        # This drives the REAL engine entry point (build_redline_edits) starting from
        # span edits shaped EXACTLY as ai_assessment_contract lowers strike_span ->
        # replace_paragraph (carrying span_action / span_anchor_quote), so it exercises
        # the engine->export seam, not a hand-built redline_edits list.
        from nda_automation.clause_outcomes import build_redline_edits

        original = (
            "The Receiving Party shall not solicit any employee of the Disclosing "
            "Party and shall not compete with the Disclosing Party for two years."
        )
        # Two prohibited restraints in ONE paragraph -> two lowered strike spans.
        anchor_one = "shall not solicit any employee of the Disclosing Party and "
        anchor_two = "shall not compete with the Disclosing Party "
        paragraph = {"id": "p1", "index": 1, "source_index": 1, "text": original}
        paragraphs = [paragraph]
        # Each lowered strike edit as the contract emits it: action already lowered to
        # replace_paragraph, ``text`` reflecting ONLY this single span's cut, plus the
        # span provenance the coalescer re-composes from.
        from nda_automation.ai_assessment_contract import apply_span

        lowered_one = apply_span(original, anchor_one, "")
        lowered_two = apply_span(original, anchor_two, "")
        self.assertIsNotNone(lowered_one)
        self.assertIsNotNone(lowered_two)
        clause = {
            "id": "non_circumvention",
            "name": "Non-circumvention",
            "status": "check",
            "issue_type": "present_but_wrong",
            "matched_paragraph_ids": ["p1"],
            "fallback": {"redline_action": "delete_paragraph"},
            "what_to_fix": "Strike the prohibited restraints.",
            "reason": "Prohibited restraints present.",
            "proposed_edits": [
                {
                    "action": "replace_paragraph",
                    "paragraph_id": "p1",
                    "text": lowered_one,
                    "span_action": "strike_span",
                    "span_anchor_quote": anchor_one,
                },
                {
                    "action": "replace_paragraph",
                    "paragraph_id": "p1",
                    "text": lowered_two,
                    "span_action": "strike_span",
                    "span_anchor_quote": anchor_two,
                },
            ],
        }

        # ENGINE: the two same-paragraph spans MUST coalesce to ONE replace_paragraph.
        edits = build_redline_edits([clause], paragraphs)
        self.assertEqual(len(edits), 1, "engine must coalesce same-paragraph spans to one edit")
        coalesced = edits[0]
        self.assertEqual(coalesced["action"], REDLINE_REPLACE_PARAGRAPH)
        self.assertEqual(coalesced["clause_id"], "non_circumvention")
        self.assertEqual(coalesced["paragraph_id"], "p1")
        self.assertEqual(coalesced.get("source_index"), 1)
        self.assertEqual(coalesced.get("paragraph_index"), 1)
        # The coalesced replacement reflects BOTH cuts (neither span clobbered).
        self.assertNotIn("solicit", coalesced["replacement_text"])
        self.assertNotIn("compete", coalesced["replacement_text"])

        # EXPORT: the single coalesced edit applies; the residual-collision guard
        # (keyed on clause_id per <w:p>) is NOT tripped because only one edit reaches
        # the export for this (clause, paragraph).
        source_text = original
        source_docx = make_source_docx([original])
        review_result = {
            "overall_status": "does_not_meet_requirements",
            "checked_at": "2026-06-01T00:00:00+00:00",
            "extracted_text": source_text,
            "paragraphs": paragraphs,
            "clauses": [],
            "redline_edits": edits,
        }
        redlined_docx = build_source_redline_docx(source_docx, review_result)
        assert_docx_package_healthy(self, redlined_docx)
        _settings_root, document_root, _document_xml = docx_xml_roots(redlined_docx)
        target = document_root.findall(".//w:body/w:p", W_NS)[0]
        # BOTH prohibited spans struck (word-level del), surviving context untracked.
        deleted_text = "".join(
            revision_text_for_state(d, accepted=False) for d in target.findall("w:del", W_NS)
        )
        self.assertIn("solicit", deleted_text)
        self.assertIn("compete", deleted_text)
        # Rejected view == original; accepted view == both cuts applied.
        self.assertEqual(revision_text_for_state(target, accepted=False), original)
        self.assertEqual(revision_text_for_state(target, accepted=True), coalesced["replacement_text"])

        # COVERAGE GATE: the single coalesced replace reconciles 1:1 and PASSES.
        self.assertEqual(
            verify_export_content_coverage(
                redlined_docx,
                source_text,
                expected_redline_edits=edits,
            ),
            [],
        )


if __name__ == "__main__":
    unittest.main()
