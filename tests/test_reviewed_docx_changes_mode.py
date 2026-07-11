"""Ticket 1 + the G2 regression for the reviewed-docx pipeline.

* ``GET /api/matters/<id>/reviewed-docx?changes=tracked|accepted`` -- tracked keeps
  the w:ins/w:del revision markup; accepted flattens it (accept-all-revisions);
  owner-scoped 404; bad changes value -> 400.
* G2 (Approach C): a PDF-source matter + a redline serves a NON-EMPTY reviewed/working
  DOCX whose redline markup is present (anchored by index into the working DOCX),
  rather than the old 0/29 silently-dropped-redline class.
"""
from __future__ import annotations

from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document

from nda_automation import (
    approval,
    artifact_registry,
    artifact_service,
    matter_document_artifacts,
    matter_store,
    pdf_ingest_conversion,
)
from nda_automation.review_engine import review_nda_with_active_engine
from nda_automation.review_result_contract import extracted_text_from_paragraphs
from nda_automation.routes import approval as approval_routes
from nda_automation.triage import triage_review_result

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

NDA_PARAGRAPHS = [
    "This Mutual Non-Disclosure Agreement is entered into by both parties.",
    "Each party agrees to keep the other party's Confidential Information secret.",
    "This Agreement shall be governed by the laws of the State of Mars.",
    "The Receiving Party shall not solicit any employees of the Disclosing Party.",
]


def _docx(paragraphs) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _docx_with_blanks(paragraphs) -> bytes:
    """A COMPLETE DOCX package whose body carries genuinely blank paragraphs for each
    empty string (python-docx ``add_paragraph("")``), so the reviewed-docx export can
    read it as a valid redline source while still exercising the blank-paragraph shape."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _has_tracked_markup(docx_bytes: bytes) -> bool:
    with ZipFile(BytesIO(docx_bytes)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    return "w:ins" in xml or "w:del" in xml


class _FakeHandler:
    def __init__(self, *, current_user_id, path=""):
        self.current_user_id = current_user_id
        self.current_user = {"id": current_user_id} if current_user_id else None
        self.path = path
        self.status = None
        self.json = None
        self.download = None
        self.download_headers = None

    def _send_json(self, payload, status=200, headers=None, *, send_body=True):
        self.status = status
        self.json = payload

    def _send_download(self, data, filename, content_type, headers=None, *, send_body=True):
        self.status = 200
        self.download = {"data": data, "filename": filename, "content_type": content_type}
        self.download_headers = headers or {}


def _seed_approved_matter_with_redline(*, source_docx=True, working_docx=False):
    """Create an APPROVED matter whose flagged clause has an accepted redline.

    DOCX source: the source DOCX itself carries the redline's anchor paragraph.
    PDF source (+working_docx): the matter is a PDF whose reconstructed working DOCX
    carries the anchor paragraph and whose review paragraph is re-keyed by index.
    """
    review_result = review_nda_with_active_engine("\n\n".join(NDA_PARAGRAPHS))
    flagged = [
        str(clause.get("id"))
        for clause in review_result.get("clauses", [])
        if clause.get("decision") in ("fail", "review")
    ]
    target_clause = flagged[0]
    review_result["redline_edits"] = [
        {
            "id": "seeded-redline",
            "clause_id": target_clause,
            "paragraph_id": "p3",
            "source_index": 3,
            "action": "replace_paragraph",
            "original_text": NDA_PARAGRAPHS[2],
            "replacement_text": NDA_PARAGRAPHS[2] + " (amended)",
        }
    ]
    document_bytes = _docx(NDA_PARAGRAPHS) if source_docx else b"%PDF-1.7\nsource pdf\n%%EOF\n"
    matter = matter_store.create_matter(
        source_filename="mutual-nda.docx" if source_docx else "mutual-nda.pdf",
        document_bytes=document_bytes,
        extracted_text="\n\n".join(NDA_PARAGRAPHS),
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_id = matter["id"]
    if working_docx:
        # Approach C: the reconstructed working DOCX carries the body the redline
        # anchors into (source_index 3 == NDA_PARAGRAPHS[2]).
        artifact_service.register_working_docx(
            matter_id, _docx(NDA_PARAGRAPHS), owner_user_id="owner-1"
        )
    # Accept the flagged clause so its redline is exported, then approve.
    matter_store.set_clause_reviewer_decision(
        matter_id,
        target_clause,
        approval.normalize_reviewer_decision({"action": "accept"}, actor="reviewer"),
    )
    matter_store.update_matter_fields(matter_id, {"status": approval.MATTER_STATUS_APPROVED})
    return matter_id


# --------------------------------------------------------------------------- #
# Ticket 1: changes=tracked|accepted
# --------------------------------------------------------------------------- #
def test_reviewed_docx_tracked_keeps_revision_markup():
    matter_id = _seed_approved_matter_with_redline(source_docx=True)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes=tracked",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 200
    assert handler.download["content_type"].startswith("application/vnd.openxmlformats")
    assert _has_tracked_markup(handler.download["data"]) is True
    assert handler.download_headers.get("X-Reviewed-Changes") == "tracked"


def test_reviewed_docx_accepted_flattens_revision_markup():
    matter_id = _seed_approved_matter_with_redline(source_docx=True)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes=accepted",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 200
    # Accepted: the tracked w:ins/w:del are flattened away.
    assert _has_tracked_markup(handler.download["data"]) is False
    assert handler.download_headers.get("X-Reviewed-Changes") == "accepted"
    assert handler.download["filename"].endswith("-accepted.docx")


def test_reviewed_docx_defaults_to_tracked_when_no_query():
    matter_id = _seed_approved_matter_with_redline(source_docx=True)
    handler = _FakeHandler(current_user_id="owner-1", path=f"/api/matters/{matter_id}/reviewed-docx")
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 200
    assert _has_tracked_markup(handler.download["data"]) is True
    assert handler.download_headers.get("X-Reviewed-Changes") == "tracked"


def test_reviewed_docx_rejects_unknown_changes_value():
    matter_id = _seed_approved_matter_with_redline(source_docx=True)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes=bogus",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 400
    assert handler.download is None


def test_reviewed_docx_owner_scoped_404_on_mismatch():
    matter_id = _seed_approved_matter_with_redline(source_docx=True)
    handler = _FakeHandler(
        current_user_id="attacker@example.com",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes=tracked",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 404
    assert handler.download is None


def test_reviewed_docx_preview_when_reviewed_but_not_approved():
    # A reviewed-but-unapproved matter (status "in_review", review_result present)
    # serves the faithful redline as a PREVIEW (200) without minting the durable
    # role="reviewed" artifact -- approval is what registers it.
    review_result = review_nda_with_active_engine("\n\n".join(NDA_PARAGRAPHS))
    matter = matter_store.create_matter(
        source_filename="mutual-nda.docx",
        document_bytes=_docx(NDA_PARAGRAPHS),
        extracted_text="\n\n".join(NDA_PARAGRAPHS),
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_store.update_matter_fields(matter["id"], {"status": "in_review"})
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter['id']}/reviewed-docx",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter['id']}/reviewed-docx")
    assert handler.status == 200, handler.json
    assert handler.download is not None
    # Preview: no durable reviewed artifact registered.
    stored = matter_store.get_matter(matter["id"], owner_user_id="owner-1")
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_REVIEWED) is None
    assert (handler.download_headers or {}).get("X-Reviewed-Artifact-ID") is None


def test_reviewed_docx_409_when_no_completed_review():
    # No review_result on the matter -> nothing reviewed to serve -> 409.
    review_result = review_nda_with_active_engine("\n\n".join(NDA_PARAGRAPHS))
    matter = matter_store.create_matter(
        source_filename="mutual-nda.docx",
        document_bytes=_docx(NDA_PARAGRAPHS),
        extracted_text="\n\n".join(NDA_PARAGRAPHS),
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_store.update_matter_fields(matter["id"], {"status": "in_review"})
    with matter_store._locked_store():
        record = matter_store._load_matter_record_by_id(matter["id"])
        record.pop("review_result", None)
        matter_store._save_matter_record(record)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter['id']}/reviewed-docx",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter['id']}/reviewed-docx")
    assert handler.status == 409
    assert handler.download is None


# --------------------------------------------------------------------------- #
# G2 regression: PDF-source matter + redline -> anchored, non-empty reviewed DOCX
# --------------------------------------------------------------------------- #
def test_g2_pdf_source_with_working_docx_serves_anchored_reviewed_docx():
    matter_id = _seed_approved_matter_with_redline(source_docx=False, working_docx=True)
    # Sanity: the matter IS a PDF source that carries a working DOCX artifact.
    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")
    assert stored["source_filename"].endswith(".pdf")
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is not None

    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes=tracked",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 200, handler.json
    data = handler.download["data"]
    # NON-EMPTY and a real DOCX (the redline did not silently vanish: the 0/29 class).
    assert data and data[:2] == b"PK"
    # The redline markup is PRESENT -- the accepted change anchored into the working
    # DOCX body by index rather than being dropped.
    assert _has_tracked_markup(data) is True


def test_p0_converted_pdf_with_blank_paragraph_before_redline_exports():
    # P0 regression: a converted-PDF working DOCX whose body carries BLANK <w:p> (the
    # spacing paragraphs pdf2docx routinely emits) before/around a redlined clause must
    # still EXPORT. The blank-counting source_index (export anchoring) and the
    # blank-stripped paragraph_index (coverage gate) are reconciled by the review
    # pipeline, so the export does not false-block on the common blank-paragraph shape.
    clauses = [
        "This Mutual Non-Disclosure Agreement is entered into by both parties on the date below.",
        "Each party agrees to keep the other party's Confidential Information strictly secret.",
        "This Agreement shall be governed by the laws of the State of Mars and its courts.",
        "The Receiving Party shall not solicit any employees of the Disclosing Party for two years.",
    ]
    # LEADING blank + an INTERIOR blank: the reconstructed body the export anchors into.
    recon = ["", clauses[0], clauses[1], "", clauses[2], clauses[3]]
    working_docx = _docx_with_blanks(recon)
    pypdf = [
        {"id": f"p{i + 1}", "text": text, "source_index": i + 1, "source_part": "pdf"}
        for i, text in enumerate(clauses)
    ]
    body_index = pdf_ingest_conversion.reconstructed_body_index(working_docx)
    mapped, _m, _u = pdf_ingest_conversion.map_paragraphs_to_reconstruction(pypdf, body_index)
    # The re-keyed source_index lives in the blank-COUNTING space (2,3,5,6).
    assert [p["source_index"] for p in mapped] == [2, 3, 5, 6]

    extracted_text = extracted_text_from_paragraphs(mapped)
    review_result = review_nda_with_active_engine(extracted_text, paragraphs=mapped)
    review_result["extracted_text"] = extracted_text

    # The stub AI engine flags clauses but emits no redlines, so inject ONE redline on
    # the governing-law clause keyed to its reviewed paragraph. The redline carries the
    # blank-counting source_index (5 -> the 5th physical <w:p>, "...laws of Mars...")
    # AND the blank-stripped paragraph_index the gate keys on -- exactly the two-index
    # shape a real reviewed redline carries, so this exercises the P0 reconciliation.
    govlaw_paragraph = next(
        paragraph
        for paragraph in review_result.get("paragraphs", [])
        if "State of Mars" in str(paragraph.get("text") or "")
    )
    review_result["redline_edits"] = [
        {
            "id": "seeded-govlaw-redline",
            "clause_id": "governing_law",
            "paragraph_id": govlaw_paragraph["id"],
            "paragraph_index": govlaw_paragraph.get("index"),
            "source_index": govlaw_paragraph.get("source_index"),
            "action": "replace_paragraph",
            "original_text": govlaw_paragraph["text"],
            "replacement_text": "This Agreement shall be governed by the laws of England and Wales.",
        }
    ]
    assert govlaw_paragraph.get("source_index") == 5  # blank-counting (leading+interior blank)

    matter = matter_store.create_matter(
        source_filename="inbound.pdf",
        document_bytes=b"%PDF-1.7\nx\n%%EOF\n",
        extracted_text=extracted_text,
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_id = matter["id"]
    artifact_service.register_working_docx(matter_id, working_docx, owner_user_id="owner-1")
    flagged = [
        str(clause.get("id"))
        for clause in review_result.get("clauses", [])
        if clause.get("decision") in ("fail", "review")
    ]
    assert flagged, "fixture should flag at least one clause to redline"
    for clause_id in flagged:
        matter_store.set_clause_reviewer_decision(
            matter_id,
            clause_id,
            approval.normalize_reviewer_decision({"action": "accept"}, actor="reviewer"),
        )
    matter_store.update_matter_fields(matter_id, {"status": approval.MATTER_STATUS_APPROVED})

    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")
    # The export does NOT raise/block (the blank-paragraph false-block class is closed).
    reviewed = matter_document_artifacts.build_reviewed_docx(matter_id, stored, owner_user_id="owner-1")
    data = reviewed.export.data
    assert data and data[:2] == b"PK"
    # And the accepted redline actually anchored into the working DOCX body.
    assert _has_tracked_markup(data) is True


# --------------------------------------------------------------------------- #
# 500-handling regression: the reviewed-docx route must NEVER leak a bare 500.
#
# After the gate relaxation (c114b5a9) a reviewed-but-unapproved (esp. PDF-source)
# matter reaches the redline producer chain that the old 409 gate short-circuited.
# Several exceptions could escape the handler as an unhandled bare HTTP 500:
#   * PdfDocxReconstructionBusy (a RuntimeError) -> must map to a retryable 503.
#   * MatterSourceTextChangedError -> must map to 409 (parity with the export route),
#     not the 400 the generic DocxExportError clause used to give it.
#   * Any unexpected exception -> must be caught, logged, and returned as a clean
#     handled error, never a bare stack-escape 500.
# We inject these at the route's call site (matter_document_artifacts.build_reviewed_docx)
# so the route's except-chain + top-level guard are exercised directly.
# --------------------------------------------------------------------------- #
from nda_automation import pdf_docx_reconstruction  # noqa: E402
from nda_automation import redline_export_service  # noqa: E402


def _seed_reviewed_unapproved_pdf_matter():
    """A reviewed-but-UNAPPROVED PDF-source matter (status in_review, review_result
    present) -- the exact class the relaxed gate now lets reach the redline producer."""
    review_result = review_nda_with_active_engine("\n\n".join(NDA_PARAGRAPHS))
    matter = matter_store.create_matter(
        source_filename="mutual-nda.pdf",
        document_bytes=b"%PDF-1.7\nsource pdf\n%%EOF\n",
        extracted_text="\n\n".join(NDA_PARAGRAPHS),
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_store.update_matter_fields(matter["id"], {"status": "in_review"})
    return matter["id"]


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_reconstruction_busy_returns_503_not_500(monkeypatch, changes):
    matter_id = _seed_reviewed_unapproved_pdf_matter()

    def _raise_busy(*_args, **_kwargs):
        raise pdf_docx_reconstruction.PdfDocxReconstructionBusy()

    monkeypatch.setattr(approval_routes.matter_document_artifacts, "build_reviewed_docx", _raise_busy)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    # Retryable 503 with Retry-After -- NOT a bare 500.
    assert handler.status == 503, handler.json
    assert handler.download is None
    assert "error" in (handler.json or {})


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_source_text_changed_returns_409_matching_export_route(monkeypatch, changes):
    matter_id = _seed_reviewed_unapproved_pdf_matter()

    def _raise_source_changed(*_args, **_kwargs):
        raise redline_export_service.MatterSourceTextChangedError("source text drifted")

    monkeypatch.setattr(approval_routes.matter_document_artifacts, "build_reviewed_docx", _raise_source_changed)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    # 409 (stale-state conflict) -- matches handle_review_docx_export, NOT 400.
    assert handler.status == 409, handler.json
    assert handler.download is None


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_unexpected_exception_is_handled_and_logged(monkeypatch, caplog, changes):
    matter_id = _seed_reviewed_unapproved_pdf_matter()

    def _raise_unexpected(*_args, **_kwargs):
        raise KeyError("unexpected internal key")

    monkeypatch.setattr(approval_routes.matter_document_artifacts, "build_reviewed_docx", _raise_unexpected)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    with caplog.at_level("ERROR"):
        approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    # Handled clean error with a JSON body -- the FE gets an actionable status, never
    # a bare unhandled stack-escape. (500 here, but the point is "handled + logged".)
    assert handler.status == 500, handler.json
    assert handler.download is None
    assert "error" in (handler.json or {})
    # And it was LOGGED with the matter id + exception type for diagnosis.
    assert any(matter_id in record.getMessage() for record in caplog.records)
    assert any("KeyError" in record.getMessage() for record in caplog.records)


def _seed_reviewed_unapproved_docx_matter():
    """A reviewed-but-UNAPPROVED *native DOCX* matter -- the control for the producer-
    boundary translation (only PDF-source failures get translated)."""
    review_result = review_nda_with_active_engine("\n\n".join(NDA_PARAGRAPHS))
    matter = matter_store.create_matter(
        source_filename="mutual-nda.docx",
        document_bytes=_docx(NDA_PARAGRAPHS),
        extracted_text="\n\n".join(NDA_PARAGRAPHS),
        review_result=review_result,
        triage=triage_review_result(review_result),
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    matter_store.update_matter_fields(matter["id"], {"status": "in_review"})
    return matter["id"]


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_pdf_producer_unclassified_error_returns_503_not_500(monkeypatch, changes):
    """The live regression: a PDF-source, reviewed-but-unapproved matter whose redline
    PRODUCER raises an UNCLASSIFIED exception (a raw KeyError/TypeError deep in the
    reconstruction-and-anchor path, NOT a DocxExportError) must surface as a CLEAN
    classified 503 (PdfSourceRedlineUnavailableError + the source-PDF annotation
    recovery payload) -- never the generic catch-all 500."""
    matter_id = _seed_reviewed_unapproved_pdf_matter()

    def _raise_unclassified(*_args, **_kwargs):
        # A raw exception that is NOT a DocxExportError -- exactly the class that fell
        # through to the generic-500 catch-all before the producer-boundary translation.
        raise KeyError("source_index")

    # Patch the inner producer so the translation wrapper in build_matter_redline runs.
    monkeypatch.setattr(redline_export_service, "_build_redline_export", _raise_unclassified)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    # Classified 503 (PdfSourceRedlineUnavailableError), NOT a generic 500.
    assert handler.status == 503, handler.json
    assert handler.download is None
    body = handler.json or {}
    assert "error" in body
    # The typed PDF-source payload: recovery points at the marked-up source PDF.
    assert body.get("reason") == "reconstruction_failed"
    assert (body.get("recovery") or {}).get("path") == "annotated_pdf"


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_native_docx_producer_unclassified_error_stays_clean_500(monkeypatch, changes):
    """The safety net: the SAME unclassified producer exception for a NATIVE DOCX matter
    is NOT translated (it is not a PDF reconstruction failure) -- it still reaches the
    handler's top-level guard and returns a LOGGED clean 500."""
    matter_id = _seed_reviewed_unapproved_docx_matter()

    def _raise_unclassified(*_args, **_kwargs):
        raise KeyError("source_index")

    monkeypatch.setattr(redline_export_service, "_build_redline_export", _raise_unclassified)
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    assert handler.status == 500, handler.json
    assert handler.download is None
    assert "error" in (handler.json or {})


# --------------------------------------------------------------------------- #
# The LIVE prod regression (matter_a014c4ffeb07 "Pismo UK 2"): a PDF-source matter
# whose Approach-C working-DOCX redline reconstruction covers only N of M source
# paragraphs. The strong DOCX content-coverage gate CORRECTLY rejects the dropped
# content, raising DocxOpenHealthError(kind="content_coverage"). Before this fix that
# typed error mapped to a generic 500 ("integrity check (approval)" -> 500). It must
# now surface as a CLEAN classified 503 (PdfSourceRedlineUnavailableError + the source-
# PDF annotation recovery payload) so the FE faithful-fallback fires instead of a crash.
#
# These tests exercise the REAL coverage-check exception: render_source_redline_package
# is patched to return a DocxPackageRenderResult whose ``content_errors`` carry the exact
# prod "N vs M" sequence-mismatch detail, so the REAL _raise_for_package_result raises the
# REAL DocxOpenHealthError and the producer-boundary translation runs end-to-end. (NOT a
# synthetic KeyError -- the prior fix chased that wrong path.)
# --------------------------------------------------------------------------- #
from nda_automation import docx_package_renderer  # noqa: E402

# The exact prod log detail (33 of 114 paragraphs covered).
_PROD_COVERAGE_ERROR = (
    "Exported accepted-change paragraph sequence does not match the expected source/redline "
    "sequence (33 paragraph(s); expected 114). The redline may have misplaced, duplicated, "
    "or dropped source content."
)


def _coverage_failing_render(*_args, **_kwargs):
    """Stand-in for render_source_redline_package that reproduces the prod content-
    coverage shortfall: a well-formed package (no health errors) whose content gate
    reports the 33-vs-114 sequence mismatch. _raise_for_package_result then raises the
    REAL DocxOpenHealthError(kind="content_coverage")."""
    return docx_package_renderer.DocxPackageRenderResult(
        data=_docx(NDA_PARAGRAPHS),
        health_errors=[],
        content_errors=[_PROD_COVERAGE_ERROR],
    )


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_pdf_coverage_failure_returns_503_not_500(monkeypatch, changes):
    """LIVE regression: a PDF-source matter whose redline reconstruction FAILS the
    content-coverage gate (33 vs 114) must return a CLEAN 503, NOT a 500 -- for both
    changes=tracked and changes=accepted."""
    matter_id = _seed_approved_matter_with_redline(source_docx=False, working_docx=True)
    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")
    assert stored["source_filename"].endswith(".pdf")  # original upload is a PDF

    # Force the REAL coverage gate to fail (the prod 33-vs-114 mismatch).
    monkeypatch.setattr(
        redline_export_service.docx_package_renderer,
        "render_source_redline_package",
        _coverage_failing_render,
    )
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")

    # CLEAN classified 503 (PdfSourceRedlineUnavailableError), NOT a generic 500.
    assert handler.status == 503, handler.json
    assert handler.download is None
    body = handler.json or {}
    assert "error" in body
    # The coverage-shortfall reason + the source-PDF annotation recovery path.
    assert body.get("reason") == "reconstruction_coverage_shortfall"
    assert (body.get("recovery") or {}).get("path") == "annotated_pdf"


def test_reviewed_docx_pdf_coverage_failure_raises_typed_error_at_producer(monkeypatch):
    """White-box: build_matter_redline translates the REAL content-coverage
    DocxOpenHealthError into PdfSourceRedlineUnavailableError for a PDF-source matter,
    AND the underlying gate exception is preserved as the cause."""
    matter_id = _seed_approved_matter_with_redline(source_docx=False, working_docx=True)
    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")

    monkeypatch.setattr(
        redline_export_service.docx_package_renderer,
        "render_source_redline_package",
        _coverage_failing_render,
    )
    with pytest.raises(redline_export_service.PdfSourceRedlineUnavailableError) as caught:
        redline_export_service.build_matter_redline(
            matter_id,
            approval.reviewed_docx_payload(stored),
            persist=False,
            owner_user_id="owner-1",
        )

    assert caught.value.status == 503
    assert caught.value.reason == "reconstruction_coverage_shortfall"
    # The underlying coverage gate exception is preserved as the cause (real path,
    # NOT a synthetic raise): a REAL DocxOpenHealthError(kind="content_coverage").
    cause = caught.value.__cause__
    assert isinstance(cause, redline_export_service.DocxOpenHealthError)
    assert cause.kind == "content_coverage"
    assert any("expected 114" in d for d in cause.details)


@pytest.mark.parametrize("changes", ["tracked", "accepted"])
def test_reviewed_docx_native_docx_coverage_failure_stays_500(monkeypatch, changes):
    """Control: the SAME content-coverage failure on a NATIVE DOCX matter is NOT
    translated (it is not a PDF reconstruction) -- it stays a logged, leak-free 500.
    This proves the coverage gate is NOT weakened; only PDF-source falls back to 503."""
    matter_id = _seed_approved_matter_with_redline(source_docx=True, working_docx=False)
    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")
    assert stored["source_filename"].endswith(".docx")

    monkeypatch.setattr(
        redline_export_service.docx_package_renderer,
        "render_source_redline_package",
        _coverage_failing_render,
    )
    handler = _FakeHandler(
        current_user_id="owner-1",
        path=f"/api/matters/{matter_id}/reviewed-docx?changes={changes}",
    )
    approval_routes.handle_matter_reviewed_docx(handler, f"/api/matters/{matter_id}/reviewed-docx")
    # Native DOCX coverage failure stays a clean 500 (the gate is intact, leak-free body).
    assert handler.status == 500, handler.json
    assert handler.download is None
    body = handler.json or {}
    assert body.get("error") == redline_export_service.DOCX_HEALTH_CLIENT_MESSAGE


# --------------------------------------------------------------------------- #
# Approach-C reviewed download filename: the substituted working DOCX bytes carry
# the redline anchor, but the DOWNLOAD name must come from the ORIGINAL source NDA
# name -- not the internal working artifact name (e.g. "02_working_docx.docx"),
# which produced an unrecognizable "..working_docx-redlined.docx" download.
# --------------------------------------------------------------------------- #
def test_approach_c_reviewed_filename_preserves_original_source_name():
    matter_id = _seed_approved_matter_with_redline(source_docx=False, working_docx=True)
    stored = matter_store.get_matter(matter_id, owner_user_id="owner-1")
    # Original upload is a PDF; a working DOCX artifact was registered for anchoring.
    assert stored["source_filename"] == "mutual-nda.pdf"
    working = artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING)
    assert working is not None
    # The internal working artifact name is the lifecycle stage name, NOT the NDA name.
    assert "working" in (working.name or "")

    export = redline_export_service.build_matter_redline(
        matter_id,
        approval.reviewed_docx_payload(stored),
        persist=False,
        owner_user_id="owner-1",
    )

    # The download is named from the ORIGINAL source NDA ("mutual-nda.pdf" ->
    # "mutual-nda-redlined.docx"), keeping the redlined suffix -- recognizable, and
    # never leaking the internal working artifact name.
    assert export.filename == "mutual-nda-redlined.docx"
    assert "working" not in export.filename


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
