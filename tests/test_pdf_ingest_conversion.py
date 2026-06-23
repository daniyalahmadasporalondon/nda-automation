"""Approach C: PDF→DOCX-at-ingest conversion + the index re-keying it produces.

Covers the module unit (convert_pdf_matter_to_docx) and the ingest wiring in
create_matter_from_document (happy path persists a working artifact + re-keyed
paragraphs; fail-open keeps the legacy PDF matter when conversion is unavailable).
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from nda_automation import (
    artifact_registry,
    ingestion_service,
    pdf_docx_reconstruction,
    pdf_ingest_conversion,
)
from nda_automation.matter_repository import InMemoryMatterRepository

PDF_BYTES = b"%PDF-1.7\nfake pdf\n%%EOF\n"

# The clause text the reconstructed DOCX body carries, in document order.
RECON_PARAGRAPHS = [
    "This Mutual Non-Disclosure Agreement is entered into by the parties.",
    "Confidential Information shall be kept strictly confidential by the receiving party.",
    "This Agreement shall be governed by the laws of England and Wales.",
]


def make_docx(paragraphs) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_docx_with_blanks(paragraphs) -> bytes:
    """A COMPLETE DOCX package (so reconstruct's part-validation passes) where an
    empty string yields a genuinely blank body ``<w:p>`` -- the spacing paragraphs
    pdf2docx routinely emits. python-docx ``add_paragraph("")`` produces exactly such
    a blank body paragraph (the canonical walker counts it; its text normalizes to '')."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class _StubConverter:
    """A pdf2docx stand-in that writes a fixed reconstructed DOCX body."""

    name = "stub-pdf2docx"

    def __init__(self, paragraphs):
        self._paragraphs = paragraphs

    def is_available(self):
        return True

    def convert_pdf_to_docx(self, source_path: Path, output_path: Path):
        output_path.write_bytes(make_docx(self._paragraphs))


def _pypdf_paragraphs():
    """pypdf review paragraphs as the PDF extractor mints them: source_part='pdf'
    and a source_index in the pypdf index space (NOT the reconstructed space)."""
    return [
        {"id": "p1", "text": RECON_PARAGRAPHS[0], "source_index": 1, "source_part": "pdf"},
        {"id": "p2", "text": RECON_PARAGRAPHS[1], "source_index": 2, "source_part": "pdf"},
        {"id": "p3", "text": RECON_PARAGRAPHS[2], "source_index": 3, "source_part": "pdf"},
    ]


# --------------------------------------------------------------------------- #
# Module unit: convert_pdf_matter_to_docx
# --------------------------------------------------------------------------- #
def test_convert_rekeys_paragraphs_to_reconstructed_index_and_drops_pdf_marker():
    converter = _StubConverter(RECON_PARAGRAPHS)
    working = pdf_ingest_conversion.convert_pdf_matter_to_docx(
        PDF_BYTES, "inbound.pdf", _pypdf_paragraphs(), converter=converter
    )
    assert working.mapped_count == 3
    assert working.unmapped_count == 0
    # Every mapped paragraph dropped the PDF marker and now anchors by index into the
    # reconstructed body (1-based canonical body-paragraph numbering).
    for index, paragraph in enumerate(working.paragraphs, start=1):
        assert "source_part" not in paragraph
        assert paragraph["source_index"] == index
    # The reconstructed bytes are a real DOCX whose body matches the reconstruction.
    body = pdf_ingest_conversion.reconstructed_body_index(working.docx_bytes)
    assert [text for (_idx, text, _norm) in body] == RECON_PARAGRAPHS


def test_convert_keeps_marker_when_paragraph_unplaceable():
    converter = _StubConverter(RECON_PARAGRAPHS)
    paragraphs = _pypdf_paragraphs()
    paragraphs.append(
        {"id": "p4", "text": "A clause that the reconstruction never contains at all.",
         "source_index": 4, "source_part": "pdf"}
    )
    working = pdf_ingest_conversion.convert_pdf_matter_to_docx(
        PDF_BYTES, "inbound.pdf", paragraphs, converter=converter
    )
    assert working.mapped_count == 3
    assert working.unmapped_count == 1
    # The unplaceable paragraph KEEPS its PDF marker so the fail-closed text anchor
    # path still guards it (never a silent drop).
    assert working.paragraphs[-1].get("source_part") == "pdf"


class _RawDocxConverter:
    """A pdf2docx stand-in that emits a raw-XML reconstructed DOCX (supports blanks)."""

    name = "stub-raw-pdf2docx"

    def __init__(self, paragraphs):
        self._paragraphs = paragraphs

    def is_available(self):
        return True

    def convert_pdf_to_docx(self, source_path: Path, output_path: Path):
        output_path.write_bytes(make_docx_with_blanks(self._paragraphs))


def test_convert_rekeys_across_blank_reconstructed_paragraphs():
    # pdf2docx emits blank spacing <w:p>; the canonical body index COUNTS them, so the
    # re-keyed source_index lands in the blank-counting space (2,3,5,6 -- skipping the
    # leading + interior blanks). This is the index space the EXPORT physically anchors
    # into, which must include the blanks.
    recon = ["", RECON_PARAGRAPHS[0], RECON_PARAGRAPHS[1], "", RECON_PARAGRAPHS[2]]
    converter = _RawDocxConverter(recon)
    working = pdf_ingest_conversion.convert_pdf_matter_to_docx(
        PDF_BYTES, "inbound.pdf", _pypdf_paragraphs(), converter=converter
    )
    assert working.mapped_count == 3
    assert working.unmapped_count == 0
    assert [p["source_index"] for p in working.paragraphs] == [2, 3, 5]


def test_convert_merged_fragments_share_one_source_index_no_collision():
    # pdf2docx MERGED two pypdf fragments of one clause into a single reconstructed
    # paragraph. Both fragments must map to THAT paragraph's source_index; the NEXT
    # clause must keep its own distinct index (no collision -> strict export anchors).
    frag_a = "Each party agrees to keep the other party's Confidential Information"
    frag_b = "and to use it only for the stated purpose of this Agreement at all times"
    merged = frag_a + " " + frag_b
    next_clause = "This Agreement shall be governed by the laws of England and Wales."
    recon = ["This Mutual NDA is entered into by both parties today.", merged, next_clause]
    converter = _RawDocxConverter(recon)
    pypdf = [
        {"id": "p1", "text": recon[0], "source_index": 1, "source_part": "pdf"},
        {"id": "p2", "text": frag_a, "source_index": 2, "source_part": "pdf"},
        {"id": "p3", "text": frag_b, "source_index": 3, "source_part": "pdf"},
        {"id": "p4", "text": next_clause, "source_index": 4, "source_part": "pdf"},
    ]
    working = pdf_ingest_conversion.convert_pdf_matter_to_docx(
        PDF_BYTES, "inbound.pdf", pypdf, converter=converter
    )
    indexes = [p.get("source_index") for p in working.paragraphs]
    parts = [p.get("source_part") for p in working.paragraphs]
    # Both merged fragments map to the merged paragraph (source_index 2); next clause = 3.
    assert indexes == [1, 2, 2, 3]
    # All four mapped -- no fragment left with the divergent pdf marker.
    assert parts == [None, None, None, None]
    # No fragment collided onto the NEXT clause's index.
    assert indexes[3] != indexes[2]
    assert working.unmapped_count == 0


class _UnavailableConverter:
    name = "stub-unavailable"

    def is_available(self):
        return False

    def convert_pdf_to_docx(self, source_path: Path, output_path: Path):
        raise AssertionError("unavailable converter should not be invoked")


def test_convert_propagates_unavailable_engine_error():
    with pytest.raises(pdf_docx_reconstruction.PdfDocxReconstructionError):
        pdf_ingest_conversion.convert_pdf_matter_to_docx(
            PDF_BYTES, "inbound.pdf", _pypdf_paragraphs(),
            converter=_UnavailableConverter(),
        )


# --------------------------------------------------------------------------- #
# Ingest wiring: create_matter_from_document
# --------------------------------------------------------------------------- #
def test_ingest_pdf_persists_working_docx_and_rekeyed_paragraphs(monkeypatch):
    repo = InMemoryMatterRepository()
    converter = _StubConverter(RECON_PARAGRAPHS)

    monkeypatch.setattr(
        ingestion_service, "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )
    # Route the conversion through the stub converter.
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )

    matter = ingestion_service.create_matter_from_document(
        filename="inbound.pdf",
        document_bytes=PDF_BYTES,
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
        repository=repo,
        drive_sync_runner=lambda func: None,
    )

    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")
    artifact = artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING)
    assert artifact is not None
    working_bytes = repo.get_artifact_document(artifact.stored_filename)
    assert working_bytes and working_bytes[:2] == b"PK"  # a real (zip) DOCX
    rekeyed = stored.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD)
    assert isinstance(rekeyed, list) and len(rekeyed) == 3
    assert all("source_part" not in paragraph for paragraph in rekeyed)
    # The matter stays "Not Reviewed": the conversion seeds NO review_result, so the
    # intentional deferred-review state is preserved.
    assert stored.get("review_result") in (None, {})


def test_ingest_pdf_fails_open_when_conversion_unavailable(monkeypatch):
    repo = InMemoryMatterRepository()
    monkeypatch.setattr(
        ingestion_service, "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )

    def _boom(*_args, **_kwargs):
        raise pdf_docx_reconstruction.PdfDocxReconstructionUnavailableError("no engine")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _boom
    )

    matter = ingestion_service.create_matter_from_document(
        filename="inbound.pdf",
        document_bytes=PDF_BYTES,
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
        repository=repo,
        drive_sync_runner=lambda func: None,
    )

    # Ingest still succeeded; the matter is the legacy un-converted PDF with NO
    # working artifact and NO re-keyed paragraphs.
    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is None
    assert stored.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD) is None


def test_ingest_artifact_failure_rolls_back_no_half_persist(monkeypatch):
    # Half-persist guard: if the working-artifact registration fails AFTER the re-keyed
    # paragraphs are written, the paragraphs are rolled back so the matter never holds
    # the divergent state (artifact present but no paragraphs, or paragraphs present but
    # no artifact). The matter falls cleanly to the legacy un-converted PDF.
    repo = InMemoryMatterRepository()
    converter = _StubConverter(RECON_PARAGRAPHS)
    monkeypatch.setattr(
        ingestion_service, "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )

    def _boom(*_args, **_kwargs):
        raise RuntimeError("artifact store write failed")

    monkeypatch.setattr(ingestion_service.artifact_service, "register_working_docx", _boom)

    matter = ingestion_service.create_matter_from_document(
        filename="inbound.pdf",
        document_bytes=PDF_BYTES,
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
        repository=repo,
        drive_sync_runner=lambda func: None,
    )

    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")
    # No working artifact (registration failed) AND no orphan re-keyed paragraphs.
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is None
    assert stored.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD) is None
    # Sanity: the field was actually written then rolled back (not just never set) --
    # the returned matter reflects the rolled-back state, not the divergent one.
    assert matter.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD) is None


def test_ingest_docx_does_not_convert(monkeypatch):
    repo = InMemoryMatterRepository()
    docx_bytes = make_docx(RECON_PARAGRAPHS)
    calls = {"n": 0}

    def _should_not_run(*_args, **_kwargs):
        calls["n"] += 1
        raise AssertionError("DOCX matters must not be reconstructed")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _should_not_run
    )
    matter = ingestion_service.create_matter_from_document(
        filename="native.docx",
        document_bytes=docx_bytes,
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
        repository=repo,
        drive_sync_runner=lambda func: None,
    )
    assert calls["n"] == 0
    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is None


# --------------------------------------------------------------------------- #
# Empty-body guard: a scanned / text-empty reconstruction must NOT convert
# --------------------------------------------------------------------------- #
class _EmptyBodyConverter:
    """A pdf2docx stand-in whose reconstructed DOCX has NO non-empty body text --
    exactly what a scanned / image-only / text-empty PDF reconstructs to (a valid
    DOCX package whose body is blank/spacing paragraphs only)."""

    name = "stub-empty-body"

    def is_available(self):
        return True

    def convert_pdf_to_docx(self, source_path: Path, output_path: Path):
        # A structurally-valid DOCX (passes the 4-required-parts check) whose body is
        # only blank paragraphs -> no anchorable text.
        output_path.write_bytes(make_docx_with_blanks(["", "", ""]))


def test_convert_refuses_when_reconstructed_body_has_no_text():
    # P0 empty-body guard: a scanned/text-empty PDF reconstructs to a valid-but-empty
    # DOCX. convert_pdf_matter_to_docx must REFUSE (raise the reconstruction-failed
    # error) rather than return a useless empty working document, so neither ingest nor
    # the retro path registers a "ready" working artifact with nothing to anchor to.
    with pytest.raises(pdf_docx_reconstruction.PdfDocxReconstructionFailedError):
        pdf_ingest_conversion.convert_pdf_matter_to_docx(
            PDF_BYTES, "scanned.pdf", _pypdf_paragraphs(),
            converter=_EmptyBodyConverter(),
        )


def test_convert_refuses_when_no_paragraph_maps():
    # The reconstructed body HAS text, but none of the pypdf review paragraphs align to
    # it (mapped_count == 0). That working DOCX is unanchorable for THIS matter's
    # review, so the guard refuses it too (fail-open keeps the PDF source).
    converter = _StubConverter(["Totally unrelated reconstructed prose number one.",
                               "And a second unrelated reconstructed sentence here."])
    pypdf = [
        {"id": "p1", "text": "Mismatched clause alpha bravo charlie delta echo.",
         "source_index": 1, "source_part": "pdf"},
        {"id": "p2", "text": "Mismatched clause foxtrot golf hotel india juliet.",
         "source_index": 2, "source_part": "pdf"},
    ]
    with pytest.raises(pdf_docx_reconstruction.PdfDocxReconstructionFailedError):
        pdf_ingest_conversion.convert_pdf_matter_to_docx(
            PDF_BYTES, "inbound.pdf", pypdf, converter=converter
        )


def test_ingest_pdf_empty_body_fails_open_no_working_artifact(monkeypatch):
    # End-to-end: a scanned/text-empty PDF must ingest fine but stay a legacy PDF matter
    # (no working artifact, no re-keyed paragraphs) so the FE keeps the page-image view.
    repo = InMemoryMatterRepository()
    converter = _EmptyBodyConverter()
    monkeypatch.setattr(
        ingestion_service, "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )
    matter = ingestion_service.create_matter_from_document(
        filename="scanned.pdf",
        document_bytes=PDF_BYTES,
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
        repository=repo,
        drive_sync_runner=lambda func: None,
    )
    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is None
    assert stored.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD) is None
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(stored) is False


# --------------------------------------------------------------------------- #
# Retro-conversion: an already-stored pre-Approach-C PDF matter gains anchors
# --------------------------------------------------------------------------- #
def _store_legacy_pdf_matter(repo, owner_user_id="owner-1"):
    """Create a PDF matter as it looked BEFORE Approach C shipped: a role="original"
    PDF artifact (via backfill) + the raw pypdf review paragraphs (source_part='pdf')
    on review_result, and NO working DOCX."""
    from nda_automation import artifact_service

    review_result = {
        "source": {"type": "pdf", "filename": "legacy.pdf"},
        "paragraphs": _pypdf_paragraphs(),
        "extracted_text": "\n\n".join(RECON_PARAGRAPHS),
    }
    matter = repo.create_matter(
        source_filename="legacy.pdf",
        document_bytes=PDF_BYTES,
        extracted_text="\n\n".join(RECON_PARAGRAPHS),
        review_result=review_result,
        triage={},
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id=owner_user_id,
    )
    # Backfill so the matter carries the role="original" artifact a real legacy matter
    # has (the retro path reads original bytes; backfill reuses stored_filename bytes).
    artifact_service.backfill_matter(matter, repository=repo, owner_user_id=owner_user_id)
    return repo.get_matter(matter["id"], owner_user_id=owner_user_id)


def _store_never_reviewed_pdf_matter(repo, owner_user_id="owner-1"):
    """A PDF matter whose every prior review FAILED: review_result is None, so it has
    NO stored review paragraphs at all -- the Pismo case. role="original" PDF artifact
    (via backfill) + extracted_text, but no working DOCX and no review_result."""
    from nda_automation import artifact_service

    matter = repo.create_matter(
        source_filename="pismo.pdf",
        document_bytes=PDF_BYTES,
        extracted_text="\n\n".join(RECON_PARAGRAPHS),
        review_result=None,
        triage={},
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id=owner_user_id,
    )
    artifact_service.backfill_matter(matter, repository=repo, owner_user_id=owner_user_id)
    return repo.get_matter(matter["id"], owner_user_id=owner_user_id)


def test_retro_convert_never_reviewed_pdf_reextracts_paragraphs_and_gains_working_docx(monkeypatch):
    # REGRESSION (Pismo): a PDF matter with NO prior successful review (review_result is
    # None -> no stored review paragraphs) must STILL gain a working DOCX. Previously the
    # retro conversion skipped with "no stored review paragraphs" and the on-demand review
    # never produced a working DOCX for such a matter. The fix re-extracts the pypdf
    # paragraphs from the PDF bytes when none are stored.
    repo = InMemoryMatterRepository()
    stored = _store_never_reviewed_pdf_matter(repo)
    assert stored.get("review_result") is None
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(stored) is False

    # Stub the PDF re-extraction (PDF_BYTES is a fake header) to return the pypdf
    # paragraphs a real text-based PDF extractor would mint.
    monkeypatch.setattr(
        ingestion_service,
        "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )
    converter = _StubConverter(RECON_PARAGRAPHS)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )

    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")

    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    # The conversion happened in THIS pass despite no prior review: working DOCX present.
    assert matter_has_working_docx(refreshed) is True
    working = artifact_registry.latest_artifact_for_role(refreshed, artifact_registry.ROLE_WORKING)
    assert working is not None
    rekeyed = refreshed.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD)
    assert isinstance(rekeyed, list) and len(rekeyed) == 3
    assert all("source_part" not in paragraph for paragraph in rekeyed)


def test_retro_convert_never_reviewed_pdf_fail_open_when_reextraction_empty(monkeypatch):
    # A scanned/image-only never-reviewed PDF re-extracts to NO paragraphs: the
    # conversion fails open (no working DOCX), leaving the page-image view.
    repo = InMemoryMatterRepository()
    stored = _store_never_reviewed_pdf_matter(repo)
    monkeypatch.setattr(
        ingestion_service,
        "extract_document",
        lambda filename, document_bytes: ("pdf", [], None),
    )

    def _should_not_run(*_args, **_kwargs):
        raise AssertionError("conversion must not run with no re-extracted paragraphs")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _should_not_run
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(refreshed) is False


def _stub_ai_review_result(text, paragraphs=None):
    """A minimal ai_first review result the persist layer accepts."""
    return {
        "extracted_text": text,
        "paragraphs": list(paragraphs or []),
        "clauses": [],
        "review_state": {"state": "pass"},
        "active_review_engine": {"executed_engine": "ai_first", "engine": "ai_first"},
        "ai_first_review": {"status": "completed"},
        "source": {"type": "pdf"},
    }


def test_on_demand_review_of_never_reviewed_pdf_yields_working_docx_in_same_pass(monkeypatch):
    """END-TO-END (Pismo): a PDF matter with NO prior successful review, reviewed
    on-demand, must END the pass with a working DOCX present -- the conversion happens
    in the SAME review pass, not the next one and not never."""
    from nda_automation import telemetry

    repo = InMemoryMatterRepository()
    stored = _store_never_reviewed_pdf_matter(repo)
    matter_id = stored["id"]
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(stored) is False

    # Re-extraction (PDF_BYTES is a fake header) returns realistic pypdf paragraphs.
    monkeypatch.setattr(
        ingestion_service,
        "extract_document",
        lambda filename, document_bytes: ("pdf", _pypdf_paragraphs(), None),
    )
    converter = _StubConverter(RECON_PARAGRAPHS)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )

    engine_calls: list[object] = []

    def _engine(text, *, paragraphs=None, **_kwargs):
        engine_calls.append(paragraphs)
        return _stub_ai_review_result(text, paragraphs)

    ingestion_service._perform_inbound_ai_review_body(
        matter_id,
        repository=repo,
        owner_user_id="owner-1",
        review_engine_func=_engine,
        telemetry=telemetry,
        is_on_demand=True,
    )

    refreshed = repo.get_matter(matter_id, owner_user_id="owner-1")
    # The conversion ran in THIS pass: working DOCX present, anchoring ready next load.
    assert matter_has_working_docx(refreshed) is True
    assert refreshed.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD)
    # The review consumed the RE-KEYED working paragraphs (conversion ran before review).
    assert engine_calls and engine_calls[0]
    assert all("source_part" not in p for p in engine_calls[0])
    # The review itself completed (status stamped completed).
    assert str(refreshed.get("review_status") or "") == "completed"


def test_on_demand_review_completes_fail_open_when_conversion_hangs(monkeypatch):
    """The review must ALWAYS complete even if the retro conversion hangs/raises: the
    guarded wrapper abandons it (fail-open) and the review proceeds on the PDF source."""
    from nda_automation import telemetry

    repo = InMemoryMatterRepository()
    stored = _store_never_reviewed_pdf_matter(repo)
    matter_id = stored["id"]

    def _hang(*_args, **_kwargs):
        import time
        time.sleep(30)
        raise AssertionError("should have been abandoned")

    monkeypatch.setattr(ingestion_service, "retro_convert_pdf_matter", _hang)
    # Tighten the wall-clock budget so the test is fast.
    monkeypatch.setattr(ingestion_service, "RETRO_PDF_CONVERT_WALL_CLOCK_SECONDS", 0.5)

    def _engine(text, *, paragraphs=None, **_kwargs):
        return _stub_ai_review_result(text, paragraphs)

    import time
    started = time.monotonic()
    ingestion_service._perform_inbound_ai_review_body(
        matter_id,
        repository=repo,
        owner_user_id="owner-1",
        review_engine_func=_engine,
        telemetry=telemetry,
        is_on_demand=True,
    )
    elapsed = time.monotonic() - started

    assert elapsed < 10.0, "a hung conversion must NOT stall the review"
    refreshed = repo.get_matter(matter_id, owner_user_id="owner-1")
    # Review still completed (fail-open); the working DOCX was abandoned (PDF source).
    assert str(refreshed.get("review_status") or "") == "completed"
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(refreshed) is False


def test_retro_convert_legacy_pdf_gains_working_docx_and_rekeyed_paragraphs(monkeypatch):
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    # Sanity: it starts as a legacy PDF matter with a role="original" PDF artifact and
    # NO working DOCX (the dead-anchor state).
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_ORIGINAL) is not None
    assert artifact_registry.latest_artifact_for_role(stored, artifact_registry.ROLE_WORKING) is None

    converter = _StubConverter(RECON_PARAGRAPHS)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )

    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")

    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    from nda_automation.matter_render_job import matter_has_working_docx
    # The matter now HAS a working DOCX -> working_docx_ready True.
    assert matter_has_working_docx(refreshed) is True
    working = artifact_registry.latest_artifact_for_role(refreshed, artifact_registry.ROLE_WORKING)
    assert working is not None
    working_bytes = repo.get_artifact_document(working.stored_filename)
    assert working_bytes and working_bytes[:2] == b"PK"  # a real (zip) DOCX
    # Re-keyed paragraphs are persisted, PDF marker dropped, anchored by index.
    rekeyed = refreshed.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD)
    assert isinstance(rekeyed, list) and len(rekeyed) == 3
    assert all("source_part" not in paragraph for paragraph in rekeyed)
    for index, paragraph in enumerate(rekeyed, start=1):
        assert paragraph["source_index"] == index


def test_retro_convert_is_idempotent_noop_when_already_converted(monkeypatch):
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    converter = _StubConverter(RECON_PARAGRAPHS)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    converted = repo.get_matter(stored["id"], owner_user_id="owner-1")

    # Second run must be a no-op: the converter is never called again.
    def _should_not_run(*_args, **_kwargs):
        raise AssertionError("retro conversion must be a no-op once a working DOCX exists")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _should_not_run
    )
    ingestion_service.retro_convert_pdf_matter(converted, repository=repo, owner_user_id="owner-1")


def test_retro_convert_skips_docx_source_matter(monkeypatch):
    repo = InMemoryMatterRepository()
    matter = repo.create_matter(
        source_filename="native.docx",
        document_bytes=make_docx(RECON_PARAGRAPHS),
        extracted_text="\n\n".join(RECON_PARAGRAPHS),
        review_result={"source": {"type": "docx"}, "paragraphs": []},
        triage={},
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    stored = repo.get_matter(matter["id"], owner_user_id="owner-1")

    def _should_not_run(*_args, **_kwargs):
        raise AssertionError("DOCX matters must not be retro-converted")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _should_not_run
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(matter["id"], owner_user_id="owner-1")
    assert artifact_registry.latest_artifact_for_role(refreshed, artifact_registry.ROLE_WORKING) is None


def test_retro_convert_fail_open_when_empty_body(monkeypatch):
    # A scanned/text-empty legacy PDF matter: retro conversion refuses (empty body) and
    # leaves it on the page-image view (no working artifact) -- never a useless empty DOCX.
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    converter = _EmptyBodyConverter()
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(refreshed) is False
    assert refreshed.get(ingestion_service.WORKING_DOCX_PARAGRAPHS_FIELD) is None


def test_retro_convert_guarded_abandons_a_slow_conversion_fail_open(monkeypatch):
    # REGRESSION: a slow/hung pdf2docx conversion must NEVER stall the review. The
    # guarded wrapper enforces an OUTER wall-clock budget; on timeout it ABANDONS the
    # conversion and returns the un-converted matter (fail-open) so the review proceeds.
    import time

    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)

    def _hang(*_args, **_kwargs):
        time.sleep(30)  # far beyond the test budget below
        raise AssertionError("should have been abandoned before completing")

    monkeypatch.setattr(ingestion_service, "retro_convert_pdf_matter", _hang)

    started = time.monotonic()
    result = ingestion_service.retro_convert_pdf_matter_guarded(
        stored, repository=repo, owner_user_id="owner-1", timeout_seconds=0.5
    )
    elapsed = time.monotonic() - started

    assert elapsed < 5.0, "guarded conversion must abandon a hung conversion fast"
    # Fail-open: returns the matter as-passed (un-converted) so the review proceeds.
    assert result is stored
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(result) is False


def test_retro_convert_guarded_returns_converted_matter_on_success(monkeypatch):
    # On a fast successful conversion the guarded wrapper returns the converted matter.
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)

    sentinel = {"id": stored["id"], "converted": True}
    monkeypatch.setattr(
        ingestion_service, "retro_convert_pdf_matter", lambda *a, **k: sentinel
    )

    result = ingestion_service.retro_convert_pdf_matter_guarded(
        stored, repository=repo, owner_user_id="owner-1", timeout_seconds=5.0
    )
    assert result is sentinel


def test_retro_convert_guarded_fail_open_when_conversion_raises(monkeypatch):
    # A raising conversion is swallowed; the un-converted matter is returned.
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)

    def _boom(*_args, **_kwargs):
        raise RuntimeError("pdf2docx exploded")

    monkeypatch.setattr(ingestion_service, "retro_convert_pdf_matter", _boom)

    result = ingestion_service.retro_convert_pdf_matter_guarded(
        stored, repository=repo, owner_user_id="owner-1", timeout_seconds=5.0
    )
    assert result is stored


# --------------------------------------------------------------------------- #
# Observability: working_docx_status outcome signal (durable, surfaced)
# --------------------------------------------------------------------------- #
STATUS_FIELD = ingestion_service.WORKING_DOCX_STATUS_FIELD


def test_retro_convert_sets_status_converted_on_success(monkeypatch):
    # A successful retro conversion persists working_docx_status="converted" so the next
    # re-run / an operator can SEE the outcome rather than inferring it from presence.
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    converter = _StubConverter(RECON_PARAGRAPHS)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    assert refreshed.get(STATUS_FIELD) == ingestion_service.WORKING_DOCX_STATUS_CONVERTED


def test_retro_convert_sets_status_failed_on_raising_conversion(monkeypatch):
    # A reconstruction that raises a NON-empty-body error records status="failed" (with
    # the exception class as the reason) and STILL fails open (matter unchanged).
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)

    def _boom(pdf_bytes, source_filename, paragraphs, **_):
        raise RuntimeError("pdf2docx exploded")

    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx", _boom
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    from nda_automation.matter_render_job import matter_has_working_docx

    assert matter_has_working_docx(refreshed) is False  # review still completes (fail-open)
    assert refreshed.get(STATUS_FIELD) == ingestion_service.WORKING_DOCX_STATUS_FAILED
    assert "RuntimeError" in str(refreshed.get(STATUS_FIELD + "_reason") or "")


def test_retro_convert_sets_status_empty_body_on_scanned_pdf(monkeypatch):
    # A scanned / text-empty PDF (reconstructs to a blank-body DOCX) records
    # status="empty_body" -- the durable signal that there was nothing to anchor.
    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    converter = _EmptyBodyConverter()
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion, "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paragraphs, **_: real_convert(
            pdf_bytes, source_filename, paragraphs, converter=converter
        ),
    )
    ingestion_service.retro_convert_pdf_matter(stored, repository=repo, owner_user_id="owner-1")
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    from nda_automation.matter_render_job import matter_has_working_docx

    assert matter_has_working_docx(refreshed) is False
    assert refreshed.get(STATUS_FIELD) == ingestion_service.WORKING_DOCX_STATUS_EMPTY_BODY


def test_retro_convert_guarded_sets_status_timed_out_on_slow_conversion(monkeypatch):
    # When the OUTER wall-clock guard abandons a slow conversion, it records
    # status="timed_out" so the next re-run TELLS us the conversion was abandoned for
    # exceeding the budget (not skipped, not failed) -- the core observability fix.
    import time

    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)

    def _hang(*_args, **_kwargs):
        time.sleep(30)

    monkeypatch.setattr(ingestion_service, "retro_convert_pdf_matter", _hang)

    result = ingestion_service.retro_convert_pdf_matter_guarded(
        stored, repository=repo, owner_user_id="owner-1", timeout_seconds=0.5
    )
    # Fail-open: the review still proceeds on the un-converted matter.
    assert result is stored
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    assert refreshed.get(STATUS_FIELD) == ingestion_service.WORKING_DOCX_STATUS_TIMED_OUT


def test_retro_pdf_convert_timeout_default_is_sixty_seconds(monkeypatch):
    # Part 3: the outer guard default is raised to 60s so a multi-page table-heavy PDF
    # (pdf2docx ~30-60s) is no longer silently abandoned by a 25s guard shorter than the
    # inner 90s subprocess timeout.
    monkeypatch.delenv("NDA_RETRO_PDF_CONVERT_TIMEOUT_SECONDS", raising=False)
    assert ingestion_service._retro_pdf_convert_timeout_seconds() == 60.0
    monkeypatch.setenv("NDA_RETRO_PDF_CONVERT_TIMEOUT_SECONDS", "45")
    assert ingestion_service._retro_pdf_convert_timeout_seconds() == 45.0


# --------------------------------------------------------------------------- #
# One-time, admin-triggered, CONVERT-ONLY PDF->working-DOCX backfill.
# --------------------------------------------------------------------------- #
def _stub_converter_patch(monkeypatch, paragraphs=RECON_PARAGRAPHS):
    """Route convert_pdf_matter_to_docx through a fixed-body stub converter so the
    backfill conversions succeed deterministically without a real pdf2docx."""
    converter = _StubConverter(paragraphs)
    real_convert = pdf_ingest_conversion.convert_pdf_matter_to_docx
    monkeypatch.setattr(
        ingestion_service.pdf_ingest_conversion,
        "convert_pdf_matter_to_docx",
        lambda pdf_bytes, source_filename, paras, **_: real_convert(
            pdf_bytes, source_filename, paras, converter=converter
        ),
    )


def test_selector_picks_only_pdf_no_docx_non_emptybody_matters(monkeypatch):
    repo = InMemoryMatterRepository()
    # 1) A legacy PDF with no working DOCX -> SELECTED.
    pdf_needs = _store_legacy_pdf_matter(repo, owner_user_id="owner-1")
    # 2) A native DOCX matter -> NOT selected (not a PDF source).
    repo.create_matter(
        source_filename="native.docx",
        document_bytes=make_docx(RECON_PARAGRAPHS),
        extracted_text="\n\n".join(RECON_PARAGRAPHS),
        review_result={"source": {"type": "docx"}, "paragraphs": []},
        triage={},
        source_type="manual_upload",
        board_column="in_review",
        owner_user_id="owner-1",
    )
    # 3) A PDF previously recorded empty_body -> NOT selected (don't retry forever).
    empty_body = _store_legacy_pdf_matter(repo, owner_user_id="owner-1")
    repo.update_matter_fields(
        empty_body["id"],
        {ingestion_service.WORKING_DOCX_STATUS_FIELD: ingestion_service.WORKING_DOCX_STATUS_EMPTY_BODY},
        owner_user_id="owner-1",
    )
    # 4) A PDF that already has a working DOCX -> NOT selected (idempotent skip).
    already = _store_legacy_pdf_matter(repo, owner_user_id="owner-1")
    _stub_converter_patch(monkeypatch)
    ingestion_service.retro_convert_pdf_matter(
        repo.get_matter(already["id"], owner_user_id="owner-1"),
        repository=repo,
        owner_user_id="owner-1",
    )

    selected = ingestion_service.select_pdf_docx_backfill_matter_ids(repository=repo)
    assert pdf_needs["id"] in selected
    assert empty_body["id"] not in selected
    assert already["id"] not in selected
    assert len(selected) == 1


def test_runner_converts_a_stub_matter_and_tallies(monkeypatch):
    from nda_automation.matter_render_job import matter_has_working_docx

    repo = InMemoryMatterRepository()
    stored = _store_legacy_pdf_matter(repo)
    _stub_converter_patch(monkeypatch)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")

    tally = ingestion_service.run_pdf_docx_backfill(repository=repo)

    assert tally["converted"] == 1
    assert tally["processed"] == 1
    assert tally["total"] == 1
    refreshed = repo.get_matter(stored["id"], owner_user_id="owner-1")
    assert matter_has_working_docx(refreshed) is True


def test_runner_skips_already_converted_and_empty_body(monkeypatch):
    repo = InMemoryMatterRepository()
    _stub_converter_patch(monkeypatch)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")
    # Already-converted matter.
    already = _store_legacy_pdf_matter(repo)
    ingestion_service.retro_convert_pdf_matter(
        repo.get_matter(already["id"], owner_user_id="owner-1"),
        repository=repo,
        owner_user_id="owner-1",
    )
    # empty_body matter.
    empty_body = _store_legacy_pdf_matter(repo)
    repo.update_matter_fields(
        empty_body["id"],
        {ingestion_service.WORKING_DOCX_STATUS_FIELD: ingestion_service.WORKING_DOCX_STATUS_EMPTY_BODY},
        owner_user_id="owner-1",
    )

    tally = ingestion_service.run_pdf_docx_backfill(repository=repo)

    # Neither was a candidate, so nothing is processed (total 0).
    assert tally["total"] == 0
    assert tally["converted"] == 0
    assert tally["processed"] == 0


def test_runner_never_calls_the_review_pipeline(monkeypatch):
    """SAFETY: the backfill is CONVERT-ONLY. Spy every review entry point in the
    ingestion service and assert ZERO calls -- the backfill must never trigger / enqueue
    an AI review or hit the review pipeline."""
    repo = InMemoryMatterRepository()
    _store_legacy_pdf_matter(repo)
    _stub_converter_patch(monkeypatch)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")

    review_calls: list[str] = []

    def _spy(name):
        def _inner(*_a, **_k):
            review_calls.append(name)
            raise AssertionError(f"backfill must NEVER call the review pipeline: {name}")

        return _inner

    # Patch every plausible review trigger that exists on the ingestion service so a
    # regression that wired a review into the backfill would FAIL loudly here.
    for attr in (
        "_run_inbound_ai_review",
        "run_inbound_ai_review",
        "enqueue_review",
        "review_nda",
        "run_review",
    ):
        if hasattr(ingestion_service, attr):
            monkeypatch.setattr(ingestion_service, attr, _spy(attr), raising=False)

    tally = ingestion_service.run_pdf_docx_backfill(repository=repo)

    assert review_calls == []
    assert tally["converted"] == 1


def test_runner_fail_open_when_one_matter_conversion_raises(monkeypatch):
    """Fail-open: a per-matter conversion error is caught and the loop continues; the
    summary is still returned. One bad PDF never aborts the run."""
    repo = InMemoryMatterRepository()
    good = _store_legacy_pdf_matter(repo)
    bad = _store_legacy_pdf_matter(repo)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")

    real_guarded = ingestion_service.retro_convert_pdf_matter_guarded

    def _guarded(matter, **kwargs):
        if matter.get("id") == bad["id"]:
            raise RuntimeError("conversion exploded for this matter")
        return real_guarded(matter, **kwargs)

    _stub_converter_patch(monkeypatch)
    monkeypatch.setattr(ingestion_service, "retro_convert_pdf_matter_guarded", _guarded)

    tally = ingestion_service.run_pdf_docx_backfill(repository=repo)

    # Both matters were processed; the good one converted, the bad one counted failed.
    assert tally["processed"] == 2
    assert tally["converted"] == 1
    assert tally["failed"] == 1
    from nda_automation.matter_render_job import matter_has_working_docx
    assert matter_has_working_docx(repo.get_matter(good["id"], owner_user_id="owner-1")) is True


def test_runner_respects_limit(monkeypatch):
    repo = InMemoryMatterRepository()
    for _ in range(3):
        _store_legacy_pdf_matter(repo)
    _stub_converter_patch(monkeypatch)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")

    tally = ingestion_service.run_pdf_docx_backfill(repository=repo, limit=2)

    assert tally["total"] == 3  # full candidate count
    assert tally["processed"] == 2  # bounded by limit this run
    assert tally["converted"] == 2


def test_rerun_is_safe_and_resumes(monkeypatch):
    repo = InMemoryMatterRepository()
    for _ in range(2):
        _store_legacy_pdf_matter(repo)
    _stub_converter_patch(monkeypatch)
    monkeypatch.setenv("NDA_PDF_DOCX_BACKFILL_SLEEP_SECONDS", "0")

    first = ingestion_service.run_pdf_docx_backfill(repository=repo, limit=1)
    assert first["converted"] == 1
    # A re-run picks up exactly the remaining one and converts it -- no duplication.
    second = ingestion_service.run_pdf_docx_backfill(repository=repo)
    assert second["total"] == 1
    assert second["converted"] == 1
    # A third run has nothing left to do.
    third = ingestion_service.run_pdf_docx_backfill(repository=repo)
    assert third["total"] == 0
    assert third["converted"] == 0


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
